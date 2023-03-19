# -*- coding: UTF-8 -*-

import getpass
import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import TimeoutException
import pickle
import random
from bs4 import BeautifulSoup
import wwaformatcreate
import wwaglobal
import wwalog
import os
import docx
from distutils.dir_util import copy_tree
import shutil
import dropbox
import re
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from collections import Counter
from operator import itemgetter
from datetime import datetime


def postAuction():
	postNumber = int(input("貼文數量 (環球水族HK - 拍賣區): "))
	postNumber2 = int(input("貼文數量 (發燒友拍賣區): "))

	wwaformatcreate.formatCreate(postNumber, postNumber2)

	# check if today's document is ready
	if not os.path.isfile(wwaglobal.auctionDocumentPath):
		# no document
		print("Today's auction document not found: " + wwaglobal.auctionDocumentPath + "\nexitting...")
		wwalog.log("Today's auction document not found: " + wwaglobal.auctionDocumentPath + "\nexitting...")
		return




	# 環球水族HK - 拍賣區: https://band.us/band/78427905/
	# 發燒友拍賣區（珊瑚、海水魚、器材用品）: https://band.us/band/75420559


	options = Options()
	options.headless = False

	print("Copying Chrome profile.")
	wwalog.log("Copying Chrome profile.")
	os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")

	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
	driver.get("https://band.us/")

	"""
	#gmail_address = input("Enter your Gmail address:\t")
	#gmail_password = input("Enter your password:\t")
	f = open("account.txt", "r")
	temp = f.read().split("\n")
	f.close()
	gmail_address = temp[0]
	gmail_password = temp[1]

	# login band via Google
	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, chrome_options=options)

	url = "https://accounts.google.com/o/oauth2/v2/auth/identifier?response_type=code&client_id=297053617361-2jq02c7vlo14302ppvfvb3di2r9ca4bq.apps.googleusercontent.com&scope=email%20profile&state=ZQDGTYK3M2LNPUBCG7L4LICPUCQL5JVNCMA3QHAZFKDTJWESRW5ZIFDOUKUJXF7O3VL3QV7GGIU5A%3D%3D%3D&prompt=consent&redirect_uri=https%3A%2F%2Fauth.band.us%2Fexternal_account_login%3Ftype%3Dgoogle&flowName=GeneralOAuthFlow"
	driver.get(url)
	emailinput = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.NAME, "identifier")))
	emailinput.send_keys(gmail_address)
	time.sleep(random.randint(1, 10))
	nextbutton = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.ID, "identifierNext")))
	ActionChains(driver).move_to_element(nextbutton).click().perform()
	time.sleep(random.randint(1, 10))

	passwordinput = WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.NAME, "password")))
	passwordinput.send_keys(gmail_password)
	time.sleep(random.randint(1, 10))
	nextbutton = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.ID, "passwordNext")))
	ActionChains(driver).move_to_element(nextbutton).click().perform()
	time.sleep(random.randint(1, 10))

	#print(colored("[P] ","red")+colored(f"You logged in using {gmail_address} ... ", "blue"))


	# check if there is alert for login error
	try:
		WebDriverWait(driver, 3).until(EC.alert_is_present())

		alert = driver.switch_to.alert
		alert.accept()
		wwalog.log("alert accepted")
		loginWithGoogleBt = driver.find_element_by_partial_link_text("Log in with Google")
		loginWithGoogleBt.click()
	except TimeoutException:
		wwalog.log("no alert")
	"""

	# check if login band success
	while 1:
		try:
			bt = WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.CLASS_NAME, "bandCreate")))
			wwalog.log("login band success.")
			break
		except TimeoutException:
			pass

	#Save cookies
	#pickle.dump(driver.get_cookies(), open(cookie_file_name, "wb"))

	coralType = []
	startBid = []

	photoList = os.listdir(wwaglobal.photoPath)

	# 發燒友拍賣區
	if postNumber2 > 0:
		# 獲取最後一個post的編號 (發燒友拍賣區（珊瑚、海水魚、器材用品）)
		driver.get("https://band.us/band/75420559/")

		firstPostCode2 = -1
		while firstPostCode2 == -1:
			try:
				soup = BeautifulSoup(driver.page_source, "lxml")
				a_tags = soup.find_all('a')
				link = ""
				for tag in a_tags:
					if tag.get('href').startswith('https://band.us/band/75420559/post/'):
						link = tag.get('href')
						break
				firstPostCode2 = int(link.split('post/')[1])
			except IndexError:
				time.sleep(0.5)

		auctionDocument2 = docx.Document(wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\發燒友docx\\" + wwaglobal.today + ".docx")
		#auctionDocument2 = docx.Document(wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\發燒友doc\\2021-09-30.docx")

		coralType = []
		startBid = []

		for i in range(postNumber2):
			coralType.append(auctionDocument2.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
			startBid.append(auctionDocument2.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())
			wwalog.log("[發燒友拍賣區 Post " + str(i + 1) + "]\n" + coralType[i] + " 起標價：" + startBid[i] + "\n拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")

			time.sleep(2)

			# click create post button
			driver.find_element_by_css_selector("[class*='cPostWriteEventWrapper'][class*='_btnOpenWriteLayer']").click()

			time.sleep(2)

			# find write post input
			postTextInput = driver.find_element_by_css_selector("div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")

			# write post text
			postTextInput.send_keys(auctionDocument2.paragraphs[i].text)


			time.sleep(0.5)

			# find file input
			postFileInput = driver.find_element_by_css_selector("input[type='file'][id^=postPhotoInput_view]")

			# upload file
			print(wwaglobal.photoPath)
			print(i)
			postFileInput.send_keys(wwaglobal.photoPath + '\\' + photoList[i])
			wwalog.log("發燒友拍賣區: Uploaded file " + wwaglobal.photoPath + '\\' + photoList[i])

			# find attach button
			while 1:
				try:
					bt = WebDriverWait(driver, 1).until(
						EC.visibility_of_element_located((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
					# click attach button
					bt.click()
					break
				except TimeoutException:
					pass

			time.sleep(1)

			# find post button
			while 1:
				try:
					bt = WebDriverWait(driver, 1).until(
						EC.visibility_of_element_located((By.CSS_SELECTOR, "button.uButton._btnSubmitPost.-confirm")))
					time.sleep(1)
					# click post button
					bt.click()
					wwalog.log("Finished 發燒友拍賣區 post " + str(i + 1) + ".")
					time.sleep(2)
					break
				except TimeoutException:
					pass

			wwalog.log("\n")
	time.sleep(1)
	#獲取最後一個post的編號
	driver.get("https://band.us/band/78427905/")

	firstPostCode = -1
	while firstPostCode == -1:
		try:
			soup = BeautifulSoup(driver.page_source, "lxml")
			a_tags = soup.find_all('a')
			link = ""
			for tag in a_tags:
				if tag.get('href').startswith('https://band.us/band/78427905/post/'):
					link = tag.get('href')
					break
			firstPostCode = int(link.split('post/')[1])
		except IndexError:
			time.sleep(0.5)

	wwalog.log("firstPostCode: " + str(firstPostCode) + "\n")

	auctionDocument = docx.Document(wwaglobal.auctionDocumentPath)

	for i in range(postNumber):
		coralType.append(auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
		startBid.append(auctionDocument.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())
		wwalog.log("[Post " + str(i + 1) + "]\n" + coralType[i + postNumber2] + " 起標價：" + startBid[i + postNumber2] + "\n拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1) + "\n")

		time.sleep(2)

		# click create post button
		driver.find_element_by_css_selector("[class*='cPostWriteEventWrapper'][class*='_btnOpenWriteLayer']").click()

		time.sleep(2)

		# find write post input
		postTextInput = driver.find_element_by_css_selector("div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")

		# write post text
		postTextInput.send_keys(auctionDocument.paragraphs[i].text)

		time.sleep(0.5)

		# find file input
		postFileInput = driver.find_element_by_css_selector("input[type='file'][id^=postPhotoInput_view]")

		# upload file
		postFileInput.send_keys(wwaglobal.photoPath + '\\' + photoList[postNumber2 + i])
		wwalog.log("Uploaded file " + wwaglobal.photoPath + '\\' + photoList[postNumber2 + i])

		# find attach button
		while 1:
			try:
				bt = WebDriverWait(driver, 1).until(EC.visibility_of_element_located((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
				# click attach button
				bt.click()
				break
			except TimeoutException:
				pass

		time.sleep(1)

		# find post button
		while 1:
			try:
				bt = WebDriverWait(driver, 1).until(EC.visibility_of_element_located((By.CSS_SELECTOR, "button.uButton._btnSubmitPost.-confirm")))
				time.sleep(1)
				# click post button
				bt.click()
				wwalog.log("Finished post " + str(i + 1) + ".")
				time.sleep(2)
				break
			except TimeoutException:
				pass

		wwalog.log("\n")



	time.sleep(1)


	# open web whatsapp
	wwalog.log("Finished posting to band.\nOpening https://web.whatsapp.com/")
	print("Finished posting to band.\nOpening https://web.whatsapp.com/")
	driver.get("https://web.whatsapp.com/")

	# check if login whatsapp success
	while 1:
		try:
			WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
			wwalog.log("login whatsapp success.")
			break
		except TimeoutException:
			pass

	grpList = ['拍賣通知專區']
	#grpList = ['93258078'] #DEBUG
	counter = 0
	time.sleep(2)

	auctionDocPath = wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\AuctionDoc.docx"
	auctionDocMessage = ""

	for grpName in grpList:
		# find group
		wwalog.log("Finding group \"" + grpName + "\"")
		driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(grpName)
		#driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys('🌊環球水族 WWA🌊會員通知區')
		time.sleep(1)
		driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(Keys.RETURN)
		time.sleep(2)

		# type message
		wwalog.log("Finding input box.")
		inputbox = driver.find_element_by_css_selector("div[role='textbox'][spellcheck='true']")
		inputbox.send_keys("❗最新珊瑚拍賣❗")
		wwalog.log("\n❗最新珊瑚拍賣❗")
		auctionDocMessage += "❗最新珊瑚拍賣❗\n"
		time.sleep(0.1)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)

		# create a file for sharing post
		f = open(wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\today_auction_links.txt", "w+", encoding = "utf-8")
		for i in range(postNumber2):
			wwalog.log("Generating msg for post " + str(i + 1) + ".")
			time.sleep(0.1)
			inputbox.send_keys(coralType[i] + " 起標價：" + startBid[i])
			wwalog.log(coralType[i] + " 起標價：" + startBid[i])
			auctionDocMessage += coralType[i] + " 起標價：" + startBid[i] + "\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys("拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1))
			wwalog.log("拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1))
			auctionDocMessage += "拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			time.sleep(0.1)

			f.write("https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")

		for i in range(postNumber):
			# each post info
			wwalog.log("Generating msg for post " + str(i + postNumber2 + 1) + ".")
			time.sleep(0.1)
			inputbox.send_keys(coralType[postNumber2+i] + " 起標價：" + startBid[postNumber2+i])
			wwalog.log(coralType[postNumber2+i] + " 起標價：" + startBid[postNumber2+i])
			auctionDocMessage += coralType[postNumber2+i] + " 起標價：" + startBid[postNumber2+i] + "\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys("拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1))
			wwalog.log("拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1))
			auctionDocMessage += "拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1) + "\n\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			time.sleep(0.1)

			f.write("https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1) + "\n")

		f.close()

		# send message
		driver.find_element_by_css_selector("span[data-testid='send'][data-icon='send']").find_element_by_xpath("..").send_keys(Keys.RETURN)
		wwalog.log("Sent message.")
		time.sleep(0.5)

		# click attach button
		driver.find_element_by_css_selector("span[data-testid='clip'][data-icon='clip']").find_element_by_xpath("..").click()
		time.sleep(0.5)
		photoPath = wwaglobal.photoPath
		dirlist = os.listdir(photoPath)

		# first photo
		driver.find_element_by_css_selector("span[data-testid='attach-image'][data-icon='attach-image']").find_element_by_xpath("../input").send_keys(photoPath + '\\' + dirlist[0])
		wwalog.log("Uploaded photo " + photoPath + '\\' + dirlist[0])
		time.sleep(2)
		for path in dirlist:
			counter += 1
			if path != dirlist[0] and counter <= postNumber + postNumber2:
				# each photo
				driver.find_element_by_css_selector("input[type='file']").send_keys(photoPath + '\\' + path)
				wwalog.log("Uploaded photo " + photoPath + '\\' + path)
				time.sleep(0.5)
		time.sleep(0.5)
		# send photos
		driver.find_element_by_css_selector("span[data-testid='send'][data-icon='send']").find_element_by_xpath("..").click()
		wwalog.log("Sent photos.")
		time.sleep(5)
	# time.sleep(0.5)
	# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys('🌊環球水族🌊到貨及優惠通知專區')
	# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(Keys.RETURN)
	# time.sleep(0.5)
	# driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys("❗最新珊瑚拍賣❗")
	# driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys(Keys.RETURN)


	auctionDoc = docx.Document()
	auctionDocLines = auctionDocMessage.split("\n")
	for i in range(len(auctionDocLines)):
		auctionDoc.add_paragraph(auctionDocLines[i])
	auctionDoc.save(auctionDocPath)


	time.sleep(5)
	driver.quit()

	print("Finished.")







def endAuction():
	date = input("Date (yyyy-mm-dd): ")
	auctionFolder = wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\" + date + " Auction Coral\\"

	tempFolder = wwaglobal.personalDropboxPath + "Temporary Store for Posting"
	if not os.path.isdir(auctionFolder):
		print("Path does not exist: " + auctionFolder)
		wwalog.log("Path does not exist: " + auctionFolder)
		return

	photoFolder = auctionFolder + "Picture"
	if not os.path.isdir(photoFolder):
		print("Path does not exist: " + photoFolder)
		wwalog.log("Path does not exist: " + photoFolder)
		return

	commentDocumentPath = wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\" + date + " Auction Coral\\comment.docx"
	if not os.path.isfile(commentDocumentPath):
		print("Comment document does not exist: " + commentDocumentPath)
		wwalog.log("Comment document does not exist: " + commentDocumentPath)
		return

	wwalog.log("commentDocumentPath: " + commentDocumentPath)
	commentDocument = docx.Document(commentDocumentPath)

	"""
	# remove temp folder from dropbox root folder
	if os.path.isdir(tempFolder):
		if not wwaglobal.debug:
			shutil.rmtree(tempFolder)
			wwalog.log("Removed folder: " + tempFolder)

			# create it again
			os.mkdir(tempFolder)
			wwalog.log("Created folder: " + tempFolder)


	# move photos from date's picture folder to root folder
	wwalog.log("Trying to copy photo folder.")
	print("Trying to copy photo folder. Please wait about several minutes.")
	if not wwaglobal.debug:
		os.system("Xcopy \"" + photoFolder + "\" \"" + tempFolder + "\" /E/H/C/I/Y/Q > nul 2>&1")
	"""
	# get the dropbox shared link of those pictures
	f = open("dropbox_token.txt", "r", encoding = "utf-8")
	temp = f.read().split("\n")
	f.close()
	DROPBOX_TOKEN = temp[0]

	wwalog.log("Trying to connect Dropbox api.")
	dbx = dropbox.Dropbox(DROPBOX_TOKEN)
	wwalog.log("Connected Dropbox api.")
	dropboxApiPhotoPath = "/Temporary Store for Posting/"
	#photoList = os.listdir(tempFolder)
	photoLink = []


	options = Options()
	options.headless = False

	print("Copying Chrome profile.")
	wwalog.log("Copying Chrome profile.")
	os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")

	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)


	f = open(auctionFolder + "today_auction_links.txt", "r", encoding = "utf-8")
	temp = f.read().split("\n")
	f.close()
	for i in range(len(temp)):
		if not "http" in temp[i]:
			continue

		driver.get(temp[i])

		while 1:
			try:
				elem = driver.find_element_by_css_selector("[class*='collageItem']").find_element_by_xpath("a/img").get_attribute('src')
				break
			except:
				time.sleep(1)

		photoLink.append(elem)
		wwalog.log(photoLink[i])



	"""
	for i in range(len(photoList)):
		photoLink.append(dbx.sharing_get_shared_links(dropboxApiPhotoPath + photoList[i]).links)
		if len(photoLink[i]) == 0:
			# generate shared link
			while 1:
				try:
					dbx.sharing_create_shared_link_with_settings(dropboxApiPhotoPath + photoList[i])
					break
				except:
					wwalog.log("Path in dropbox \"" + dropboxApiPhotoPath + photoList[i] + "\" not found. Try again in 5 seconds...")
					print("Path in dropbox \"" + dropboxApiPhotoPath + photoList[i] + "\" not found. Try again in 5 seconds...")
					time.sleep(5)

		photoLink[i] = str(dbx.sharing_get_shared_links(dropboxApiPhotoPath + photoList[i]).links[0].url)
		wwalog.log("Link of photo " + str(i + 1) + ": " + photoLink[i])
	"""

	bid_price = []
	for i in range(len(commentDocument.paragraphs)):
		try:
			bid_price.append(int(commentDocument.paragraphs[i].text.split('$')[1].split('中標')[0].strip()))
		except:
			# no one bid
			bid_price.append(0)
		wwalog.log("Bid " + str(i + 1) + ": $" + str(bid_price[i]))

	f = open(auctionFolder + date + " Auction Coral CSV.csv", "r", encoding = "utf-8")
	lines = f.read().split("\n")
	f.close()
	print(lines)
	print(len(lines))
	print(bid_price)
	print(len(bid_price))
	print(photoLink)
	print(len(photoLink))
	#modified
	for i in range(1, len(lines)-1):
		if lines[i] != "":
			print(i)
			lines[i] = lines[i].replace("[auction_price]", str(bid_price[i-1])).replace("[photo_link]", photoLink[i-1])

	f = open(auctionFolder + date + " Auction Coral CSV.csv", "w+", encoding = "utf-8")
	for i in range(len(lines)):
		if lines[i] != "":
			f.write(lines[i] + "\n")
	f.close()







	driver.get("https://www.boutir.com/user_cms/upload-products")


	bt = None
	# find upload input
	while 1:
		try:
			bt = WebDriverWait(driver, 1).until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.csvCapture")))
			break
		except TimeoutException:
			pass

	time.sleep(2)

	bt.send_keys(auctionFolder + date + " Auction Coral CSV.csv")

	time.sleep(5)

	if not wwaglobal.debug:
		WebDriverWait(driver, 1).until(EC.visibility_of_element_located((By.CSS_SELECTOR, "button.btn.btn-primary"))).click()

	print("Uploading CSV...")
	wwalog.log("Uploading CSV...")

	input("When upload finished, press <enter> to continue.")
	wwalog.log("User pressed <enter> to continue.")


	f = open(auctionFolder + date + " Auction Coral CSV.csv", "r", encoding = "utf-8")
	lines = f.read().split("\n")
	f.close()

	itemLinks = []
	if wwaglobal.debug:
		for i in range(1, len(lines)):
			if lines[i] == "":
				continue
			print("link" + str(i))
			itemLinks.append("link" + str(i))
	seller_id = "5690381968343040"
	# find itemLinks
	if not wwaglobal.debug:
		for i in range(1, len(lines)):
			if lines[i] == "":
				continue
			query = lines[i].split(",")[2].strip()
			url = "https://wwa.boutir.com/?q=" + query
			driver.get(url)
			WebDriverWait(driver, 10).until(
				EC.presence_of_element_located((By.CLASS_NAME, 'search-result-title'))
			)
			driver.find_element_by_class_name('clickable').click()
			WebDriverWait(driver, 10).until(
				EC.presence_of_element_located((By.CLASS_NAME, 'addthis_button_more'))
			)
			time.sleep(1)
			item_id = driver.current_url.split("?q=")[0].split("/")[-1]

			driver.get("https://wwa.boutir.com/apis/get_cache_item_details?is_shorten_url=0&item_id=" + item_id + "&item_preview=&seller_id=" + seller_id)
			time.sleep(1)
			itemLinks.append(driver.page_source.split("\"shorten_item_url\": \"")[1].split("\"")[0])
			print(itemLinks[i-1])
			wwalog.log(itemLinks[i-1])




	linkDocument = docx.Document(auctionFolder + "AuctionDoc.docx")

	links = []

	for i in range(2, len(linkDocument.paragraphs), 3):
		links.append(linkDocument.paragraphs[i].text.split('拍賣連結：')[1].strip())

	nameList = []
	f = open("nameList.txt", "r", encoding = "utf-8")
	nameList = f.read().split("\n")
	f.close()

	for i in range(len(links)):
		wwalog.log("Opening: " + links[i])
		driver.get(links[i])
		time.sleep(5)

		commentText = []

		soup = BeautifulSoup(driver.page_source, "lxml")
		for el in soup.find_all('p', attrs={'class': 'txt _commentContent'}):
			commentText.append(el.get_text())

		names = []
		for j in range(len(commentText)):
			try:
				names.append(WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="content"]/div/section/div/div/div[5]/div/div/div[' + str(j+1) + ']/div[1]/div[1]/div[2]/button/strong'))).text)
			except:
				continue

		print(names)

		index = 0

		if nameList[i] == "No one":
			# no one bid
			continue

		for j in range(len(names) - 1, -1, -1):
			if names[j] == nameList[i]:
				index = j
				break

		#print("i = " + str(i) + "; index = " + str(index))
		bt = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, '//*[@id="content"]/div/section/div/div/div[5]/div/div/div[' + str(index+1) + ']/div[1]/div[1]/div[3]/div[3]/button[2]')))
		# click reply button
		bt.click()

		inputBox = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, 'textarea.commentWrite._use_keyup_event._messageTextArea')))
		JS_ADD_TEXT_TO_INPUT = """
			var elm = arguments[0], txt = arguments[1];
			elm.value += txt;
			elm.dispatchEvent(new Event('change'));
			"""

		elem = driver.find_element_by_css_selector('textarea.commentWrite._use_keyup_event._messageTextArea')
		text = commentDocument.paragraphs[i].text.replace("拍賣品連結：", "拍賣品連結：" + itemLinks[i])

		driver.execute_script(JS_ADD_TEXT_TO_INPUT, elem, text)

		# send
		if not wwaglobal.debug:
			inputBox.send_keys(Keys.CONTROL + Keys.RETURN)



		time.sleep(5)

	time.sleep(5)
	driver.quit()


def checkOrderList():
	options = Options()
	options.headless = False

	print("Copying Chrome profile.")
	wwalog.log("Copying Chrome profile.")
	os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")

	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
	driver.get("https://www.boutir.com/user_cms/orders/")

	elemList = []
	dateList = []
	months = set()
	googleSheetOrders = []

	while len(elemList) == 0 or len(dateList) == 0:
		time.sleep(5)
		elemList = driver.find_elements_by_css_selector("[href*='/user_cms/edit-order/']")
		dateList = driver.find_elements_by_css_selector("[class*='purchase_date']")

	"""
	for i in range(len(elemList)):
		elemList[i] = elemList[i].find_element_by_xpath("../..")

	print(elemList[0].find_element_by_xpath("td[3]/a").text)
	"""

	i = 0
	while i < len(dateList):
		if "日期" in dateList[i].text:
			dateList.pop(i)
		else:
			#print(dateList[i].text)
			try:
				months.add(int(dateList[i].text.split("/")[1]))
			except:
				pass
		i += 1

	#print(months) # {9, 10}
	#print(len(elemList))
	#print(len(dateList))
	#print(dateList[0].text)

	# dict of boolean
	wrote = {}
	temp = ""
	for i in range(len(elemList)):
		temp = str(elemList[i].text).strip(" \n")
		wrote[temp] = False
		if "已發貨" in elemList[i].find_element_by_xpath("../..").find_element_by_css_selector("[class='status']").find_element_by_xpath("div[1]/span").text:
			wrote[temp] = True
		elif "無效訂單" in elemList[i].find_element_by_xpath("../..").find_element_by_css_selector("[class='status']").find_element_by_xpath("div[1]/span").text:
			wrote[temp] = True
		elif "等待收款確認" in elemList[i].find_element_by_xpath("../..").find_element_by_css_selector("[class='status']").find_element_by_xpath("div[1]/span").text:
			wrote[temp] = True
		wwalog.log("[CMS " + str(i+1) + "]: " + temp)

	# If modifying these scopes, delete the file token.json.
	SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

	# The ID and range of a spreadsheet.
	SPREADSHEET_ID = '1VDfkinOq_SkezxMwmF4Fhi1pe0qTtYEbCxp1CQxtxbY'
	range_names = ["to_add"]
	for month in months:
		range_names.append(str(month) + '月')

	wwalog.log("Range names: " + ", ".join(range_names))

	"""Shows basic usage of the Sheets API.
	Prints values from a sample spreadsheet.
	"""
	creds = None
	# The file token.json stores the user's access and refresh tokens, and is
	# created automatically when the authorization flow completes for the first
	# time.
	if os.path.exists('token.json'):
		creds = Credentials.from_authorized_user_file('token.json', SCOPES)
	# If there are no (valid) credentials available, let the user log in.
	if not creds or not creds.valid:
		if creds and creds.expired and creds.refresh_token:
			creds.refresh(Request())
		else:
			flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
			creds = flow.run_local_server(port=0)
		# Save the credentials for the next run
		with open('token.json', 'w') as token:
			token.write(creds.to_json())

	service = build('sheets', 'v4', credentials=creds)

	# Call the Sheets API
	sheet = service.spreadsheets()

	for i in range(len(range_names)):
		result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=range_names[i], valueRenderOption="FORMULA").execute()
		values = result.get('values', [])
		#print(values)

		if not values:
			print("Google sheet " + range_names[i] + ": No data found.")
			wwalog.log("Google sheet " + range_names[i] + ": No data found.")
		else:
			for row in values:
				# column A = row[0], B = row[1], C = row[2]...
				try:
					# format: 16-digit
					if len(re.findall("\d{16}", str(row[2]))) >= 1:
						googleSheetOrders.append(re.findall("\d{16}", str(row[2]))[0])
						wwalog.log("Google sheet: " + re.findall("\d{16}", str(row[2]))[0])
					else:
						wwalog.log("Not 16-digit: " + row[2])
				except:
					wwalog.log("Row: " + str(row))

	for key in wrote.keys():
		if key in googleSheetOrders:
			wrote[key] = True

	valuesToAdd = []

	temp = ""
	counter = 0
	for key in wrote.keys():
		if not wrote[key]:
			counter += 1
			temp += str(counter) + ": " + key + "\n"

			driver.get("https://www.boutir.com/user_cms/edit-order/" + key)
			tempRow = ["", "", "'" + key]

			table = None
			while 1:
				try:
					table = driver.find_element_by_css_selector("table[class*='tool-tip-table']")
					break
				except:
					time.sleep(1)

			# check if need to add
			if "已發貨" in driver.find_element_by_css_selector("[class*='status-label']").text:
				continue
			elif "無效訂單" in driver.find_element_by_css_selector("[class*='status-label']").text:
				continue
			elif "等待收款確認" in driver.find_element_by_css_selector("[class*='status-label']").text:
				continue


			# name
			tempRow.append(table.find_element_by_xpath("tr[1]/td[2]").find_element_by_css_selector("[class*='bold-class']").text)

			# phone
			if table.find_element_by_xpath("tr[2]/td[2]").find_element_by_css_selector("[class*='bold-class']").text[4] == "1":
				continue
			tempRow.append(table.find_element_by_xpath("tr[2]/td[2]").find_element_by_css_selector("[class*='bold-class']").text[4:])

			# location
			if "門市" in table.find_element_by_xpath("tr[3]/td[2]").find_element_by_css_selector("[class*='bold-class']").text:
				tempRow.append("門市")
			elif "魚街" in table.find_element_by_xpath("tr[3]/td[2]").find_element_by_css_selector("[class*='bold-class']").text:
				tempRow.append("魚街")
			elif "送貨" in table.find_element_by_xpath("tr[3]/td[2]").find_element_by_css_selector("[class*='bold-class']").text:
				tempRow.append("送貨")
			else:
				tempRow.append("")

			# items
			tempElements = driver.find_elements_by_css_selector("[class*='item-title-cell']")
			tempItems = []
			for i in range(len(tempElements)):
				tempElements[i] = tempElements[i].find_element_by_xpath("span")
				if "CAA" in tempElements[i].text:
					tempItems.append(tempElements[i].text.split("CAA")[1].split(" ")[0])
				elif "caa" in tempElements[i].text:
					tempItems.append(tempElements[i].text.split("caa")[1].split(" ")[0])
				else:
					tempItems.append(tempElements[i].text.split(" ")[-1])

			if wwaglobal.isInt(tempItems[0]):
				# add a '#' before it
				tempRow.append("#" + " ".join(tempItems))
			else:
				tempRow.append(" ".join(tempItems))

			# money
			tempRow.append(driver.find_element_by_css_selector("[class*='dollars']").find_element_by_css_selector("[class*='bold-class']").text[4:])

			# payment method
			if "payme" in driver.find_element_by_css_selector("[class*='payment-method']").text.lower():
				tempRow.append("payme")
			else:
				tempRow.append("")

			valuesToAdd.append(tempRow)


	valuesToAdd = sorted(valuesToAdd, key=itemgetter(4))

	# add to google sheet
	body = {
		'values': valuesToAdd
	}
	value_input_option = 'USER_ENTERED'

	result = service.spreadsheets().values().append(
		spreadsheetId=SPREADSHEET_ID, range=range_names[0],
		valueInputOption=value_input_option, body=body).execute()

	if result.get('updates').get('updatedCells') == None:
		print("Google sheet: 0 cells appended.")
		wwalog.log("Google sheet: 0 cells appended.")
		print("Do not have new order.")
		wwalog.log("Do not have new order.")
	else:
		# updated
		print("Google sheet: " + str(result.get('updates').get('updatedCells')) + ' cells appended.')
		wwalog.log("Google sheet: " + str(result.get('updates').get('updatedCells')) + ' cells appended.')
		print("Have new order, please check the google sheet \"" + range_names[0] + "\" tab.")
		wwalog.log("Have new order, please check the google sheet \"" + range_names[0] + "\" tab.")




	# wwalog.log("\nOrders that did not put into Google sheet (Total " + str(counter) + "):")
	# wwalog.log(temp)
	# print("\nOrders that did not put into Google sheet (Total " + str(counter) + "):")
	# print(temp)

	duplicatedItems = []
	Counter(googleSheetOrders)
	for key in Counter(googleSheetOrders).keys():
		if Counter(googleSheetOrders)[key]>1:
			duplicatedItems.append(key)

	if len(duplicatedItems) > 0:
		print("Duplicated items in Google sheet: " + ", ".join(duplicatedItems) + ".\n")
		wwalog.log("Duplicated items in Google sheet: " + ", ".join(duplicatedItems) + ".\n")

	driver.quit()

def meetingNotification():
	# If modifying these scopes, delete the file token.json.
	SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

	# The ID and range of a spreadsheet.
	SPREADSHEET_ID = '1VDfkinOq_SkezxMwmF4Fhi1pe0qTtYEbCxp1CQxtxbY'
	range_names = str(datetime.today().month) + "月"
	creds = None
	# The file token.json stores the user's access and refresh tokens, and is
	# created automatically when the authorization flow completes for the first
	# time.
	if os.path.exists('token.json'):
		creds = Credentials.from_authorized_user_file('token.json', SCOPES)
	# If there are no (valid) credentials available, let the user log in.
	if not creds or not creds.valid:
		if creds and creds.expired and creds.refresh_token:
			creds.refresh(Request())
		else:
			flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
			creds = flow.run_local_server(port=0)
		# Save the credentials for the next run
		with open('token.json', 'w') as token:
			token.write(creds.to_json())

	service = build('sheets', 'v4', credentials=creds)
	sheet = service.spreadsheets()
	result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=range_names,
								valueRenderOption="FORMULA").execute()
	values = result.get('values', [])
	print(values)
