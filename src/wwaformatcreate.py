# -*- coding: UTF-8 -*-

import openpyxl
import pandas as pd
import docx
import wwaglobal
import wwalog
from datetime import datetime, timedelta
import os
from shutil import copy
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
import shutil
import time
import shutil
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
import time
from datetime import datetime
from datetime import timedelta
import csv
from PIL import Image, ImageTk, ImageFont, ImageDraw, ImageGrab
from pprint import pprint
import cv2
import numpy as np
import glob
import emoji
import win32clipboard as clip
import win32con
from io import BytesIO


wwaglobal.init()


def formatCreate(postNumber, postNumber2):
	wb = openpyxl.load_workbook(wwaglobal.src_path +'Auction Item.xlsx')
	ws = wb.active

	coral_code = []
	coral_name = []
	coral_size = []
	base_price = []
	per_bid = []
	fever_guy = None

	if postNumber2 > 0:
		fever_guy = 'y'
	else:
		fever_guy = 'n'

	#print(ws.max_row+1)
	#print(ws['A2'].value)
	#print(ws['B2'].value)
	for i in range(2,ws.max_row+1):
		anemone = False
		if ws['A' + str(i)].value != None:
			coral_code.append('CAA' + str(ws['A' + str(i)].value))
			print(coral_code)
		else:
			coral_code.append(str(ws['A' + str(i)].value))
		if ws['A' + str(i)].value == None:


			anemone = True
		#print(i)
		#print(anemone)
		if anemone:
			try:
				if ws['B' + str(i)].value.split('吋',1)[0] != "":
					if wwaglobal.isInt(ws['B' + str(i)].value.split('吋',1)[0][-1]):
						coral_name.append(ws['B' + str(i)].value.split('吋')[1])
					else:
						coral_name.append(ws['B' + str(i)].value)
				else:
					coral_name.append(ws['B' + str(i)].value)
			except:
				coral_name.append(ws['B' + str(i)].value)
		else:
			try:
				if ws['B' + str(i)].value.split('頭',1)[0] != "":
					if wwaglobal.isInt(ws['B' + str(i)].value.split('頭',1)[0][-1]):
						coral_name.append(ws['B' + str(i)].value.split('頭',1)[1])
					else:
						coral_name.append(ws['B' + str(i)].value)
				else:
					coral_name.append(ws['B' + str(i)].value)

			except:

				coral_name.append(ws['B' + str(i)].value)

		if anemone:
			try:
				if str(ws['B' + str(i)].value).includes('吋'):
					coral_size.append(ws['B' + str(i)].value.split('吋')[0] + '吋')
			except:
				coral_size.append(str(ws['C' + str(i)].value) + '吋')
		else:
			if ws['C' + str(i)].value == None:
				coral_size.append('可參考(frag座2.5cm)')
			else:
				coral_size.append(str(ws['C' + str(i)].value) + '頭')
		base_price.append(ws['D' + str(i)].value)
		print(f"{i=} {ws['D' + str(i)].value=}")
		if ws['D' + str(i)].value <=1000:
			per_bid.append(20)
		else: per_bid.append(50)
	#⭐️即日晚上11點完標⭐️\n
	post_format = '⭐️[tmr]號晚上9點完標⭐️\n❗請注意交收時間、地點及拍賣規則，OK先好落標❗\n品種：[coral_name] [coral_code]\n尺寸：[coral_size]\n底價：$[base_price]\n每口價：$[perBid]\n截止時間：[date] 21:00\n\n交收時間及地點：\n星期二 太子金魚街萬寧 7-7:15pm\n星期日 太子金魚街萬寧4:30-4:45pm\n星期一 新蒲崗門市 5-8pm\n星期六 新蒲崗門市 5-8pm\n標準完標時間每晚9點，如最後15分鐘有人投標，拍賣將延遲15分鐘，無限次延遲，直到最後一個投標時間超過15分鐘為之結束。\n\n如21:00完標，20:50有人投標 ，完標時間延遲至21:05；21:03再有人投標，完標時間延遲至21:18；如21:17有人投標，完標時間延遲至21:32；直到21:32仍無人競標，則21:17的投標者中標。\n\n注意事項:\n1.中標者在我哋同事提供拍賣品中標連結後，請盡早落單\n2.中標者必須在兩星期內取貨\n3.如需hold貨超過兩星期，必需先支付貨款，不包生死，生死自付，請在拍賣前確保能在特定的時間及地點取貨\n4.如中標者無合理解釋下棄標或兩星期內沒有取貨，除會被踢出Band Group，取消其拍賣資格外，環球水族保留法律追討的權利\n5.相機影相始終或多或少點都有啲色差，我哋已盡量影接近實物顏色，介意勿拍'
	fever_guy_post_format = '品種：[coral_name] [coral_code]\n尺寸：[coral_size]\n底價：$[base_price]\n每口價：$[perBid]\n截止時間：[date] 21:00\n\n交收時間及地點：\n星期二 太子金魚街萬寧 7-7:15pm\n星期日 太子金魚街萬寧4:30-4:45pm\n星期一 新蒲崗門市 5-8pm\n星期六 新蒲崗門市 5-8pm\n如最後5分鐘有人投標，拍賣將延遲5分鐘。\n如20:00結標，19:59有人投標 ，結標時間延遲為20:05。\n\n備注：中標者請在兩星期內完成交收'


	#print(post_format.replace('[coral_name]', coral_name[1]))
	post = docx.Document()
	fever_guy_post = docx.Document()
	start_from=0
	if fever_guy == 'y':
		start_from = postNumber2
		for i in range(postNumber2):
			if coral_code[i] != 'None':
				fever_guy_post.add_paragraph(fever_guy_post_format.replace('[coral_name]', coral_name[i]).replace('[coral_code]', coral_code[i]).replace('[coral_size]', coral_size[i]).replace('[base_price]', str(base_price[i])).replace('[perBid]', str(per_bid[i])).replace('[date]', (datetime.today() + timedelta(3)).strftime('%Y-%m-%d')))
			else:
				fever_guy_post.add_paragraph(fever_guy_post_format.replace('[coral_name]', coral_name[i]).replace('[coral_code]', '').replace('[coral_size]',coral_size[i]).replace('[base_price]', str(base_price[i])).replace('[perBid]', str(per_bid[i])).replace('[date]',(datetime.today() + timedelta(3)).strftime('%Y-%m-%d')))

	for i in range(start_from, len(coral_code)):
		if coral_code[i] != 'None':
			#print(coral_name)
			#print(coral_size)
			#print(base_price)
			post.add_paragraph(post_format.replace('[tmr]',(datetime.today()+timedelta(days=1)).strftime('%d')).replace('[coral_name]', coral_name[i]).replace('[coral_code]', coral_code[i]).replace('[coral_size]', coral_size[i]).replace('[base_price]', str(base_price[i])).replace('[perBid]', str(per_bid[i])).replace('[date]', (datetime.today()+timedelta(days=1)).strftime('%Y-%m-%d')))
		else:
			post.add_paragraph(post_format.replace('[tmr]',(datetime.today()+timedelta(days=1)).strftime('%d')).replace('[coral_name]', coral_name[i]).replace('[coral_code]', '').replace('[coral_size]', coral_size[i]).replace('[base_price]', str(base_price[i])).replace('[perBid]', str(per_bid[i])).replace('[date]', (datetime.today()+timedelta(days=1)).strftime('%Y-%m-%d')))

	path = wwaglobal.dropboxPath + str(datetime.today().year)+'\\Auction Coral Retial\\' + wwaglobal.today + ' Auction Coral'
	fever_guy_path = wwaglobal.dropboxPath + str(datetime.today().year)+'\\Auction Coral Retial\\發燒友docx\\'
	if os.path.isdir(path):
		# path already exists, remove it
		wwalog.log("Path \"" + path + "\" already exists, trying to remove it")
		shutil.rmtree(path)
	os.mkdir(path)
	os.mkdir(path + '\\Picture')
	post.save(path + '\\品種.docx')
	if fever_guy == 'y':
		fever_guy_post.save(fever_guy_path + wwaglobal.today + '.docx')


	# copy all pictures
	temp_path = os.listdir(path + '\\Picture')
	# for f in temp_path:
	# 	print(f)
	#shutil.copyfile('Picture', path+'\\Picture')
	os.system("Xcopy \"" + wwaglobal.src_photoPath + "\" \"" + (path + '\\Picture') + "\" /E/H/C/I/Y/Q > nul 2>&1")
	if os.path.exists(path + '\\Picture\\desktop.ini'):
		os.remove(path + '\\Picture\\desktop.ini')
	#shutil.copytree(Picture, temp_path, dirs_exist_ok=True)
	time.sleep(1)
	#print(coral_code)
	#print(coral_name)
	#print(coral_size)
	#print(base_price)


	# CSV
	counter = 0
	f = open(path + "\\" + wwaglobal.today + " Auction Coral CSV.csv", "w+", encoding = "utf-8")
	f.write("連結,商品編號,名稱,介紹,商品選項 1 - 種類,商品選項 1 - 名稱,商品選項 2 - 種類,商品選項 2 - 名稱,商品選項 3 - 種類,商品選項 3 - 名稱,商品選項相片連結,商品選項影片連結,類別,原價,折扣價,會員價,存貨,購買上限,重量 (公斤),SKU,相片網址,影片網址,Hashtags,上架狀態\n")
	if fever_guy == 'y':
		counter = postNumber2
	for i in range(counter, len(coral_code)):
		temp = ",,拍賣品 請勿購買  "
		if coral_code[i] == "None":
			# 海葵
			temp += coral_size[i] + coral_name[i] + ",,,,,,,,,,海葵 / 拍賣品,"
		else:
			if coral_size[i] == "可參考(frag座2.5cm)":
				# frag座
				temp += coral_name[i] + " " + coral_size[i] + " " + coral_code[i] + " ,,,,,,,,,,珊瑚 / 拍賣品,"
			else:
				# have size
				temp += coral_size[i] + coral_name[i] + " " + coral_code[i] + " ,,,,,,,,,,珊瑚 / 拍賣品,"

		temp += "[auction_price],,,1,,,,[photo_link],,,1\n"


		f.write(temp)

	f.close()


	# Google sheet
	fever_count = 0

	values = [
		[""] # skip one line
	]

	for i in range(postNumber + postNumber2):
		row = []
		today = str(datetime.today().strftime('%Y')) + "/" + str(datetime.today().strftime('%m')) + "/" + str(datetime.today().strftime('%d'))
		row.append(today) # e.g. "2021/9/30"
		row.append("a")

		end_date = datetime.today()+timedelta(days=1)
		end_date_str = str(end_date.strftime('%Y')) + "/" + str(end_date.strftime('%m')) + "/" + str(end_date.strftime('%d'))
		row.append(end_date_str)
		row.append("2100")

		if coral_code[i] == "None":
			row.append("")
		else:
			# CAA...
			row.append(coral_code[i][3:])

		row.append(str(coral_name[i]))
		if "frag" in str(coral_size[i]):
			row.append("")
		else:
			row.append(str(coral_size[i]).rstrip("吋頭"))

		row.append(str(base_price[i]))

		if fever_count < postNumber2:
			fever_count += 1
			row.append("發燒友拍賣區")
		else:
			row.append("環球水族HK - 拍賣區")

		row.append("y")

		wwalog.log("row " + str(i+1) + ": " + str(row))
		values.append(row)

	appendToGoogleSheet(values)



def appendToGoogleSheet(values):
	# If modifying these scopes, delete the file token.json.
	SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

	# The ID and range of a spreadsheet.
	SPREADSHEET_ID = '12ZPNlUc6bCrw3lXruTYev8Y0p3GYqXP-qGwmL8RhotE'
	RANGE_NAME = str(datetime.today().strftime('%Y')) + '年' + str(datetime.today().strftime('%m')).zfill(2) + '月 auction!A:AG'

	creds = None
	# The file token.json stores the user's access and refresh tokens, and is
	# created automatically when the authorization flow completes for the first
	# time.
	if os.path.exists(wwaglobal.src_path + 'token.json'):
		creds = Credentials.from_authorized_user_file(wwaglobal.src_path + 'token.json', SCOPES)
	# If there are no (valid) credentials available, let the user log in.
	if not creds or not creds.valid:
		if creds and creds.expired and creds.refresh_token:
			creds.refresh(Request())
		else:
			flow = InstalledAppFlow.from_client_secrets_file(wwaglobal.src_path + 'credentials.json', SCOPES)
			creds = flow.run_local_server(port=0)
		# Save the credentials for the next run
		with open(wwaglobal.src_path + 'token.json', 'w') as token:
			token.write(creds.to_json())

	service = build('sheets', 'v4', credentials=creds)

	# Call the Sheets API
	sheet = service.spreadsheets()

	body = {
		'values': values
	}
	value_input_option = 'USER_ENTERED'

	result = service.spreadsheets().values().append(
		spreadsheetId=SPREADSHEET_ID, range=RANGE_NAME,
		valueInputOption=value_input_option, body=body).execute()

	if result.get('updates').get('updatedCells') == None:
		print("Google sheet: 0 cells appended.")
		wwalog.log("Google sheet: 0 cells appended.")
	else:
		print("Google sheet: " + str(result.get('updates').get('updatedCells')) + ' cells appended.')
		wwalog.log("Google sheet: " + str(result.get('updates').get('updatedCells')) + ' cells appended.')


global data_folder_path
data_folder_path = wwaglobal.src_path + "history_data"
global picture_folder_path
picture_folder_path = wwaglobal.src_path + "temp_picture"
global auction_picture_path
auction_picture_path = wwaglobal.src_path + "auction_picture"

quited = False




class InputApp(tk.Tk):

	# __init__ function for class tkinterApp
	def __init__(self, *args, **kwargs):
		# __init__ function for class Tk
		tk.Toplevel.__init__(self, *args, **kwargs)
		self.geometry('700x500')
		self.title('INPUT')
		# creating a container
		container = tk.Frame(self)
		container.grid()

		container.grid_rowconfigure(0, weight=1)
		container.grid_columnconfigure(0, weight=1)

		# initializing frames to an empty array
		self.frames = {}

		# iterating through a tuple consisting
		# of the different page layouts
		for F in (input_page, combination_page):
			frame = F(container, self)

			# initializing frame of that object from
			# startpage, page1, page2 respectively with
			# for loop
			self.frames[F] = frame

			frame.grid(row=0, column=0, sticky="nsew")

		self.show_frame(input_page)
		#self.mainloop()

	# to display the current frame passed as
	# parameter
	def show_frame(self, cont):
		frame = self.frames[cont]
		frame.tkraise()


class input_page(tk.Frame):
	def __init__(self, parent, controller):
		tk.Frame.__init__(self, parent)

		# code_label = tk.Label(self, text='編號')
		# code_label.grid(row= 0, column=2)
		# name_label = tk.Label(self, text='品種')
		# name_label.grid(row=0, column=3)
		# photo_label = tk.Label(self, text='相片')
		# photo_label.grid(row=0, column=4)
		self.items = {}
		self.item_count = 0
		self.photo_count = {}
		self.photo_path = {}
		main_frame = tk.Frame(self, width=690, height=450)
		main_frame.grid(row=1, column=0, columnspan=4)
		main_frame.grid_propagate(0)
		canvas = tk.Canvas(main_frame, width=670, height=400, bg='gray')
		canvas.grid(row=0, column=0, sticky="NSEW")

		scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
		scrollbar.grid(row=0, column=2, sticky='ns')
		canvas.configure(yscrollcommand=scrollbar.set)
		scrollable_frame = ttk.Frame(canvas)
		scrollable_frame.grid(row=0, column=0, sticky='NSEW')
		canvas.create_window((0, 0), window=scrollable_frame, anchor='nw')

		# data_folder_path = r"C:\Users\user\PycharmProjects\infoInput\history_data"
		# picture_folder_path = "C:/Users/user/PycharmProjects/infoInput/picture/"
		# auction_picture_path = r"C:\Users\user\PycharmProjects\infoInput\auction_picture"

		def update_scroll():
			height_sum = 0
			for row in self.items.values():
				height_sum += row[2].winfo_height() + 10
			scrollable_frame.update_idletasks()
			scrollable_frame.config(width=700, height=height_sum)
			canvas.config(scrollregion=canvas.bbox("all"))

		def change_button_command():
			for key, row in self.items.items():
				row[0].config(command=lambda index=key: delete_row(index))

		def delete_row(row):
			print(f'{row=}')
			# print(self.items[row][2].get())
			for column in self.items[row]:
				column.grid_forget()

			# pprint(f'{self.items=}')
			# pprint(f'{self.item_count=}')
			# pprint(f'{self.photo_count=}')

			for column in self.items[row]:
				column.destroy()
			del self.items[row]
			del self.photo_path[row]
			# print(f'del {row}')
			del self.photo_count[row]

			# print(f'after del {self.photo_count=}')

			self.item_count -= 1
			# print(f'{row=} {self.item_count+1=}')
			for i in range(row, self.item_count + 1):
				# print(f'{i=}')
				self.items[i] = self.items[i + 1]
				# button = self.items[i][0]
				# print(button)
				# button.configure(text=i, command=lambda: test(i))
				# self.items[i][0].configure(text=i, command=lambda: delete_row(i))
				# self.items[i][0] = tk.Button(self, text=i, width=3, height=1, command=lambda: delete_row(i))
				self.items[i][0].grid(row=i, column=0, sticky=tk.NSEW, pady=3, padx=3)
				# print(f'row {i} now commmand is {i}')
				self.items[i][1].grid(row=i, column=1)
				self.items[i][2].grid(row=i, column=2, sticky=tk.NSEW, pady=3, padx=3)
				self.items[i][3].grid(row=i, column=3, sticky=tk.NSEW, pady=3, padx=3)
				self.items[i][4].grid(row=i, sticky=tk.NSEW, pady=3, padx=3)
				self.items[i][4].config(command=lambda index=i: add_photo(index))
				self.items[i][5].grid(row=i, column=4, sticky=tk.NSEW, pady=3, padx=3)

				self.photo_count[i] = self.photo_count[i + 1]

				self.photo_path[i] = self.photo_path[i + 1]
				if self.photo_count[i] > 0:
					for j in range(6, self.photo_count[i] + 6):
						self.items[i][j].grid(row=i)
			if row != self.item_count + 1:
				self.items.pop(self.item_count + 1)
				self.photo_count.pop(self.item_count + 1)
			change_button_command()
			# for i in range(1,self.item_count+1):
			#     self.items[i][0].configure(text='del' + str(i), command=lambda: delete_row(i))
			# pprint(f'{self.items=}')
			# pprint(f'{self.item_count=}')
			# pprint(f'{self.photo_count=}')
			update_scroll()

		# scrollable_frame.update_idletasks()
		# scrollable_frame.config(width=700, height=self.item_count * 50)
		# canvas.config(scrollregion=canvas.bbox("all"))

		def add_photo(row):
			# print(self.photo_count)
			pixels_x = 50
			pixels_y = 50
			pprint(self.photo_count)
			self.photo_count[row] += 1
			file_path = filedialog.askopenfilename(filetypes=[("image", ".jpg")])
			# print(file_path)
			if os.path.isfile(file_path):
				self.photo_path[row].append(file_path)
				image = Image.open(file_path)

				render = ImageTk.PhotoImage(image.resize((pixels_x, pixels_y)))
				img = tk.Label(scrollable_frame, image=render)
				img.image = render
				img.grid(row=row, column=self.photo_count[row] + 4)
				self.items[row].append(img)
				self.items[row][4].grid(row=row, column=self.photo_count[row] + 5)
			# pprint(self.photo_path)
			update_scroll()

		def add_entry_field(row):
			# print(f'{row=}')

			cancel_button = tk.Button(scrollable_frame, text='x', width=3, height=1, command=lambda: delete_row(row))
			cancel_button.grid(row=row, column=0, sticky=tk.NSEW, pady=3, padx=3)

			caa_label = tk.Label(scrollable_frame, text='CAA')
			caa_label.grid(row=row, column=1)
			code_entry_box = tk.Entry(scrollable_frame, relief=tk.RIDGE, width=4, justify='center')
			code_entry_box.grid(row=row, column=2, sticky=tk.NSEW, pady=3, padx=3)
			name_entry_box = tk.Entry(scrollable_frame, relief=tk.RIDGE, justify='center')
			name_entry_box.grid(row=row, column=3, sticky=tk.NSEW, pady=3, padx=3)
			size_entry_box = tk.Entry(scrollable_frame, relief=tk.RIDGE, width=4, justify='center')
			size_entry_box.grid(row=row, column=4, sticky=tk.NSEW, pady=3, padx=3)
			# file_path = filedialog.askopenfilename()
			# photo_entry_box = tk.Entry(self, relief=tk.RIDGE, justify='center')
			# photo_entry_box.grid(row=row+1, column=3, sticky=tk.NSEW, pady=3, padx=3)

			# photo_input = tk.Button(self, text='+', command=lambda: add_photo(row), image=pixel, width=3, height=0.5, compound="c")
			photo_input = tk.Button(scrollable_frame, text='+', command=lambda: add_photo(row), width=1, height=1)
			photo_input.grid(row=row, column=5, sticky=tk.NSEW, pady=3, padx=3)

			self.items[row] = [cancel_button, caa_label, code_entry_box, name_entry_box, photo_input, size_entry_box]

			# print(f'{row=} {self.items[row][0]=}')

			self.item_count += 1
			self.photo_count[self.item_count] = 0

			self.photo_path[row] = []

			update_scroll()
			# print(f'{self.photo_count=}')

		# scrollable_frame.update_idletasks()
		# scrollable_frame.config(width=700, height=(self.item_count+1) * 50)
		# canvas.config(scrollregion=canvas.bbox("all"))

		def save():
			for key, item in self.items.items():
				if item[2].get() == '' and item[3].get() == '':
					tk.messagebox.showerror(message=f'第{key}行未輸入編號')
					return
			folder_name = datetime.strftime(datetime.now(), "%d-%m-%Y_%H-%M-%S")
			photo_history_path = os.path.join(picture_folder_path, f'{folder_name}/')
			os.mkdir(photo_history_path)
			f = open(f'{data_folder_path}\\{folder_name}.csv', 'w', encoding='UTF8', newline='')
			writer = csv.writer(f)
			for key, item in sorted(self.items.items()):
				temp_list = []
				# pprint(item)
				pprint(f'{key=} {item[2].get()}')
				if item[2].get() == '':
					temp_list.append('None')
				else:
					temp_list.append(item[2].get())
				if item[3].get() == '':
					temp_list.append('None')
				else:
					temp_list.append(item[3].get())
				if item[5].get() == '':
					temp_list.append('可參考(frag座2.5cm)')
				elif item[2].get() == '':
					temp_list.append('約' + item[5].get() + '寸')
				else:
					temp_list.append(item[5].get() + '頭')
				# for path in self.photo_path[key]:
				#     temp_list.append(path)
				writer.writerow(temp_list)

				photo_folder_path = os.path.join(photo_history_path, str(key).zfill(3))
				# print(photo_folder_path)
				os.mkdir(photo_folder_path)
				for index, path in enumerate(self.photo_path[key]):
					copy(path, photo_folder_path)
					os.rename(photo_folder_path + '/' + os.path.basename(path),
							  photo_folder_path + '/' + f'save{str(index).zfill(2)}.jpg')
			f.close()
			list_of_files = glob.glob(data_folder_path + '\\*.csv')
			earliest_file = min(list_of_files, key=os.path.getctime)
			# print(f'{earliest_file=}')
			if len(list_of_files) > 30:
				os.remove(earliest_file)
			list_of_files = glob.glob(picture_folder_path)
			earliest_file = min(list_of_files, key=os.path.getctime)
			print(f'{earliest_file=}')
			if len(list_of_files) > 30:
				shutil.rmtree(earliest_file)
			return folder_name

		def open_save():
			file_path = filedialog.askopenfilename(initialdir=data_folder_path)
			if not file_path:
				return
			file_name = os.path.basename(file_path).split('.')[0]
			# save_folder = picture_folder_path + file_name
			save_folder = os.path.join(picture_folder_path, file_name)
			# print(picture_folder_path + file_name)
			photos_folder_path = os.listdir(save_folder)
			pprint(photos_folder_path)
			# print(file_path)
			file = open(file_path, 'r', encoding='UTF8')
			reader = csv.reader(file)

			for row in self.items.values():
				for column in row:
					column.destroy()
			self.items.clear()
			self.photo_count.clear()
			self.photo_path.clear()
			self.item_count = 0
			temp_content = {}
			row_count = 0
			for row in reader:
				row_count += 1
				# print(f'{row_count=} {row=}')
				temp_content[row_count] = row

			self.item_count = len(temp_content)
			for key, content in temp_content.items():

				cancel_button = tk.Button(scrollable_frame, text='x', width=1, height=1,
										  command=lambda k=key: delete_row(k))
				cancel_button.grid(row=key, column=0, sticky=tk.NSEW, pady=3, padx=3)

				caa_label = tk.Label(scrollable_frame, text='CAA')
				caa_label.grid(row=key, column=1)
				code_textEntry = tk.StringVar()
				if temp_content[key][0] != 'None':
					code_textEntry.set(temp_content[key][0])
				else:
					code_textEntry.set('')
				code_entry_box = tk.Entry(scrollable_frame, relief=tk.RIDGE, width=4, justify='center',
										  textvariable=code_textEntry)
				code_entry_box.grid(row=key, column=2, sticky=tk.NSEW, pady=3, padx=3)

				name_textEntry = tk.StringVar()
				if temp_content[key][1] != 'None':
					name_textEntry.set(temp_content[key][1])
				else:
					name_textEntry.set('')
				name_entry_box = tk.Entry(scrollable_frame, relief=tk.RIDGE, justify='center',
										  textvariable=name_textEntry)
				name_entry_box.grid(row=key, column=3, sticky=tk.NSEW, pady=3, padx=3)

				size_textEntry = tk.StringVar()
				# temp_content[key][2].replace('可參考(frag座2.5cm)', '').replace('約', '').replace('寸', '').replace('頭', '')
				size_textEntry.set(
					temp_content[key][2].replace('可參考(frag座2.5cm)', '').replace('約', '').replace('寸', '').replace('頭',
																												  ''))
				size_entry_box = tk.Entry(scrollable_frame, relief=tk.RIDGE, width=4, justify='center',
										  textvariable=size_textEntry)
				size_entry_box.grid(row=key, column=4, sticky=tk.NSEW, pady=3, padx=3)
				# file_path = filedialog.askopenfilename()
				# photo_entry_box = tk.Entry(self, relief=tk.RIDGE, justify='center')
				# photo_entry_box.grid(row=row+1, column=3, sticky=tk.NSEW, pady=3, padx=3)

				# photo_input = tk.Button(self, text='+', command=lambda: add_photo(row), image=pixel, width=3, height=0.5, compound="c")

				self.photo_path[key] = []
				#print(f'{save_folder + "/" + photos_folder_path[key - 1]}')
				corresponding_path = save_folder + "/" + photos_folder_path[key - 1]
				photos_path = os.listdir(corresponding_path)
				pprint(f'{photos_path=}')
				photo_input = tk.Button(scrollable_frame, text='+', command=lambda k=key: add_photo(k), width=1,
										height=1)
				photo_input.grid(row=key, column=6 + len(photos_path), sticky=tk.NSEW, pady=3, padx=3)
				self.items[key] = [cancel_button, caa_label, code_entry_box, name_entry_box, photo_input,
								   size_entry_box]
				for counter, photos in enumerate(photos_path):
					path = corresponding_path + '/' + photos
					self.photo_path[key].append(path)
					image = Image.open(path)

					render = ImageTk.PhotoImage(image.resize((50, 50)))
					img = tk.Label(scrollable_frame, image=render)
					img.image = render
					img.grid(row=key, column=counter + 1 + 4)
					self.items[key].append(img)
				self.photo_count[key] = len(photos_path)
				update_scroll()

		def input_photos():
			photos_path = filedialog.askopenfilenames()
			if photos_path:
				for path in photos_path:
					key = self.item_count + 1
					add_entry_field(key)
					self.photo_path[key].append(path)
					self.photo_count[key] += 1
					pixels_x = 50
					pixels_y = 50
					#pprint(self.photo_count)


					# print(file_path)
					if os.path.isfile(path):

						image = Image.open(path)

						render = ImageTk.PhotoImage(image.resize((pixels_x, pixels_y)))
						img = tk.Label(scrollable_frame, image=render)
						img.image = render
						img.grid(row=key, column=self.photo_count[key] + 4)
						self.items[key].append(img)
						self.items[key][4].grid(row=key, column=self.photo_count[key] + 5)
					# pprint(self.photo_path)
					update_scroll()
			pprint(f'{self.items=}')


		def finished_editing():

			# cv2.FONT_HERSHEY_SIMPLEX
			folder_name = save()
			for key, item in sorted(self.items.items()):
				if item[2].get() == '' and item[3].get() == '':
					return
			auction_picture_folder = os.listdir(auction_picture_path)
			for folder in auction_picture_folder:
				shutil.rmtree(os.path.join(auction_picture_path, folder))
			latest_folder_path = os.path.join(picture_folder_path, folder_name)
			temp_picture_folder = os.listdir(latest_folder_path)
			# pprint(temp_picture_folder)
			for temp_folder in temp_picture_folder:
				# print(f'{temp_folder=}')
				dst = os.path.join(auction_picture_path, temp_folder)
				# print(f'{dst=}')
				shutil.copytree(os.path.join(latest_folder_path, temp_folder), dst, copy_function=shutil.copy)
			# pprint(f'{auction_picture_path=}')
			auction_picture_folder = os.listdir(auction_picture_path)
			for dst_folder in auction_picture_folder:
				each_coral_path = os.path.join(auction_picture_path, dst_folder)
				photos = os.listdir(each_coral_path)
				# pprint(f'{photos=}')
				for photo in photos:
					photo_path = os.path.join(each_coral_path, photo)
					image = cv2.imread(photo_path)
					key = int(dst_folder.split('.')[0])
					# print('Original Dimensions : ', image.shape)
					font_path = wwaglobal.src_path + "JasonHandwriting4.ttf"
					font = ImageFont.truetype(font_path, 50, encoding="utf-8")
					image = cv2.resize(image, (500, 500))  # , interpolation=cv2.INTER_AREA
					cv2img = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
					imgPil = Image.fromarray(cv2img)
					draw = ImageDraw.Draw(imgPil)
					print(f'{self.items[key][2].get()=}')
					if self.items[key][2].get() != '':
						draw.text((10, 0), 'CAA' + self.items[key][2].get(), font=font, fill=(255, 0, 0))
					draw.text((10, 420), self.items[key][3].get(), font=font, fill=(255, 0, 0))
					# cv2.putText(image, 'CAA' + self.items[key][2].get(), (30, 30), font=font, 1, (0, 0, 255), 3)
					# cv2.putText(image, self.items[key][3].get(), (30, 60), font=font, 1, (0, 0, 255), 3)

					cv2charimg = cv2.cvtColor(np.array(imgPil), cv2.COLOR_RGB2BGR)
					# cv2.imshow('Image', cv2charimg)
					cv2.imwrite(os.path.join(each_coral_path, photo), cv2charimg)

			controller.show_frame(combination_page)

		# self.item_count = len(self.items)
		self.add_field_button = tk.Button(self, text='新增', command=lambda: add_entry_field(self.item_count + 1))
		self.complete_button = tk.Button(self, text='儲存', command=lambda: save())
		self.finished_button = tk.Button(self, text='完成', command=lambda: finished_editing())
		self.input_photo_button = tk.Button(self, text='批量加入相片', command=lambda: input_photos())
		self.input_photo_button.grid(row=0, column=1, sticky='W')
		self.add_field_button.grid(row=2, column=0, sticky='EW')
		self.complete_button.grid(row=2, column=1, sticky='EW')
		self.finished_button.grid(row=2, column=2, sticky='EW')
		add_entry_field(1)
		open_save_button = tk.Button(self, text='開啟舊檔', command=lambda: open_save())
		open_save_button.grid(row=0, column=0, sticky='W')


class combination_page(tk.Frame):
	def __init__(self, parent, controller):
		tk.Frame.__init__(self, parent)
		self.items = {}
		self.groups_count = 0
		self.grp_content = {}
		self.photo_count = {}

		back_button = tk.Button(self, text='返回', command=lambda: controller.show_frame(input_page), width=30, pady=5)
		back_button.grid(row=0, column=0, sticky='W')
		group_frame = tk.Frame(self, bg='Blue', width=680, height=400)
		group_frame.grid_propagate(0)
		group_frame.grid(row=1, column=0)
		group_frame.grid_columnconfigure(0, weight=1)
		main_canvas = tk.Canvas(group_frame, width=690, height=400)
		main_canvas.grid(row=0, column=0, sticky="NSEW")

		main_scrollbar = ttk.Scrollbar(group_frame, orient="vertical", command=main_canvas.yview)
		main_scrollbar.grid(row=0, column=2, sticky='ns')
		main_canvas.configure(yscrollcommand=main_scrollbar.set)
		main_scrollable_frame = ttk.Frame(main_canvas)
		main_scrollable_frame.grid(row=0, column=0, sticky='NSEW')
		# main_scrollable_frame.bind("<Configure>", lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all"), height=e.height))
		main_canvas.create_window((0, 0), window=main_scrollable_frame, anchor='nw')

		def del_grp(group_no):
			pprint(self.grp_content)
			self.items[group_no][0].destroy()
			self.grp_content.pop(group_no)
			for i in range(group_no, self.groups_count):
				self.items[i] = self.items[i + 1]
				self.items[i][1].config(command=lambda index=i: del_grp(index))
				self.items[i][4].config(command=lambda index=i: selection_menu(index))
				self.items[i][0].grid(row=i)
				self.items[i][2].config(text=f'組合{i}')
				self.grp_content[i] = self.grp_content[i + 1]
				self.grp_content.pop(i + 1)
			#self.grp_content.pop(self.grp_content[str(self.groups_count)])
			self.groups_count -= 1
			# self.grp_content[group_no].clear()
			pprint(self.grp_content)
			main_scrollable_frame.update_idletasks()
			main_scrollable_frame.config(width=700, height=self.groups_count * 60)
			main_canvas.config(scrollregion=main_canvas.bbox("all"))
			pprint(f'{self.grp_content=}')

		def add_to_grp(list_of_corals, target_group):
			# self.items[target_group][3] = the item frame

			for coral in list_of_corals:
				# print(f'{coral=} {target_group=}')
				path_of_folder = os.path.join(auction_picture_path, coral)
				self.grp_content[target_group].append(coral)
				for path in os.listdir(path_of_folder):
					self.photo_count[target_group] += 1
					row = self.photo_count[target_group] // 6
					if self.photo_count[target_group] < 6:
						column = self.photo_count[target_group] - 1
					else:
						column = self.photo_count[target_group] % 6
					complete_path = os.path.join(path_of_folder, path)
					image = Image.open(complete_path)

					render = ImageTk.PhotoImage(image.resize((50, 50)))
					img = tk.Label(self.items[target_group][3], image=render)
					img.image = render

					img.grid(row=row, column=column, sticky='NSEW', pady=1)

					if self.photo_count[target_group] % 5 == 0:

						row_button = row
					else:

						row_button = (self.photo_count[target_group] + 1) // 6

					if (self.photo_count[target_group] + 1) % 6 == 0:
						column_button = 6
					else:
						column_button = column + 1
					# column_button = (self.photo_count[target_group] + 1) % 6
					self.items[target_group][4].grid(row=row_button, column=column_button)
					self.items[target_group][3].config(
						height=(row + 1) * 60)  # (self.photo_count[target_group] + 1) // 6 * 60 + 60

		def selection_menu(row_no):
			def get_result(target, value_list):
				selected_items = []
				for value in value_list.values():
					#print(f'{value.get()=}')
					if value.get() != '000':
						selected_items.append(value.get())
				#print(f'{selected_items=}')
				if len(selected_items) > 0:
					add_to_grp(selected_items, target)
					menu.quit()
					menu.destroy()

			menu = tk.Toplevel()
			menu.title('選擇列表')
			menu.geometry('400x500')
			coral_list_frame = tk.Frame(menu, bg='gray', height=450, width=390)  # , height=450, width=290
			coral_list_frame.grid(row=0, column=0, padx=5, pady=5, sticky='NW')
			coral_list_frame.grid_rowconfigure(0, weight=1)
			coral_list_frame.grid_columnconfigure(0, weight=1)
			coral_list_frame.grid_propagate(0)

			# coral_list_frame.columnconfigure(0, weight=1)
			# coral_list_frame.columnconfigure(1, weight=1)

			canvas = tk.Canvas(coral_list_frame)
			canvas.grid(row=0, column=0, sticky="NSEW")

			scrollbar = ttk.Scrollbar(coral_list_frame, orient="vertical", command=canvas.yview)
			scrollbar.grid(row=0, column=2, sticky='ns')
			canvas.configure(yscrollcommand=scrollbar.set)
			scrollable_frame = ttk.Frame(canvas)
			scrollable_frame.grid(row=0, column=0, sticky='NSEW')
			canvas.create_window((0, 0), window=scrollable_frame, anchor='nw')

			button_frame = tk.Frame(menu, height=50, width=390, padx=5, pady=5)

			button_frame.grid(row=1, column=0)
			button_frame.grid_propagate(0)
			button_frame.columnconfigure((0, 1), weight=1)
			cancel_button = tk.Button(button_frame, text='取消', command=lambda: menu.destroy())
			cancel_button.grid(row=0, column=0, sticky='EW')

			list_of_files = glob.glob(data_folder_path + '\\*.csv')
			latest_file = max(list_of_files, key=os.path.getctime)

			file = open(latest_file, 'r', encoding='UTF8')
			reader = csv.reader(file)
			list_content = []
			for row in reader:
				list_content.append(row)

			file.close()
			scrollable_frame.grid_columnconfigure(0, weight=5)
			scrollable_frame.grid_columnconfigure(1, weight=6)
			var = {}
			skip = False
			for index, info in enumerate(list_content):
				pprint(self.grp_content)
				for content in self.grp_content.values():

					if str(index + 1).zfill(3) in content:
						skip = True
				if skip:
					skip = False
					continue
				var[index + 1] = tk.StringVar()
				if info[0] == 'None':
					checkbox = tk.Checkbutton(scrollable_frame, text=f'{info[1]}', variable=var[index + 1],
											  anchor='w',
											  width=20, onvalue=str(index + 1).zfill(3), offvalue='000')
				else:
					checkbox = tk.Checkbutton(scrollable_frame, text=f'CAA{info[0]} {info[1]}', variable=var[index + 1],
											  anchor='w',
											  width=20, onvalue=str(index + 1).zfill(3), offvalue='000')
				checkbox.deselect()
				checkbox.grid(row=index, column=0, sticky='EW')
				photo_frame = tk.Frame(scrollable_frame, height=60, width=200)
				photo_frame.grid(row=index, column=1, sticky='NSEW')
				photo_frame.grid_propagate(0)
				coral_pic_path = os.path.join(auction_picture_path, str(index + 1).zfill(3))
				photos = os.listdir(coral_pic_path)
				for photo_count, photo in enumerate(photos):
					complete_path = os.path.join(coral_pic_path, photo)
					image = Image.open(complete_path)

					render = ImageTk.PhotoImage(image.resize((50, 50)))
					img = tk.Label(photo_frame, image=render)
					img.image = render
					img.grid(row=index, column=photo_count + 1, sticky='NSEW', pady=1)

			scrollable_frame.update_idletasks()
			scrollable_frame.config(width=400 + scrollbar.winfo_width(), height=100 * len(list_content))
			canvas.config(scrollregion=canvas.bbox("all"))

			done_button = tk.Button(button_frame, text='完成', command=lambda: get_result(row_no, var))
			done_button.grid(row=0, column=1, sticky='EW')

			menu.mainloop()

		def add_group(row):
			self.groups_count += 1

			single_grp_frame = tk.Frame(main_scrollable_frame, bg='grey', pady=5, padx=5)
			single_grp_frame.grid(row=row + 1, column=0, sticky='EW')

			# single_grp_frame.grid_columnconfigure(0, weight=1)
			# single_grp_frame.grid_columnconfigure(1, weight=5)
			# single_grp_frame.grid_columnconfigure(2, weight=1)
			del_button = tk.Button(single_grp_frame, text='x', command=lambda: del_grp(row + 1))
			del_button.grid(row=0, column=0)
			group_label = tk.Label(single_grp_frame, text=f'組合{self.groups_count}', width=20)
			group_label.grid(row=0, column=1, sticky='NS', padx=5, pady=5)
			self.item_frame = tk.Frame(single_grp_frame, bg='Gray', pady=5, width=350, height=40)
			self.item_frame.grid_propagate(0)
			self.item_frame.grid(row=0, column=2, sticky='EW')

			group_add_button = tk.Button(self.item_frame, text='+', command=lambda: selection_menu(row + 1))
			group_add_button.grid(row=0, column=1, padx=5, sticky='NS', pady=1)
			group_base_entry = tk.Entry(single_grp_frame, width=20)
			group_base_entry.grid(row=0, column=4, sticky='NS', padx=5, pady=5)
			self.items[self.groups_count] = [single_grp_frame, del_button, group_label, self.item_frame,
											 group_add_button, group_base_entry]
			self.grp_content[self.groups_count] = []
			self.photo_count[self.groups_count] = 0

			main_scrollable_frame.update_idletasks()
			main_scrollable_frame.config(width=700, height=self.groups_count * single_grp_frame.winfo_height() * 2)
			main_canvas.config(scrollregion=main_canvas.bbox("all"))

		def create_format(fever):
			self.temp = 0
			if fever:
				# Enter number of posts
				fever_window = tk.Toplevel()

				def end():
					try:

						temp_number = int(entry.get())
						if temp_number > 3:
							tk.messagebox.showerror('貼文數量過大', '請輸入少於3的有效數字')
						fever_window.quit()
						self.temp = temp_number
					except:
						tk.messagebox.showerror('錯誤', '尚未輸入正確數量，請重新輸入')

				label = tk.Label(fever_window, text='請輸入發燒友拍賣數量:')
				entry = tk.Entry(fever_window)
				button = tk.Button(fever_window, text='輸入', command=lambda: end())
				label.grid(row=0, column=0, sticky='W')
				entry.grid(row=0, column=1, sticky='W')
				button.grid(row=0, column=2, sticky='W')

				fever_window.mainloop()
				fever_post_number = self.temp

			else:
				fever_post_number = 0
			#print(f'{fever_post_number}')
			coral_code = []
			coral_size = []
			base_price = []
			total = []
			list_of_files = glob.glob(data_folder_path + '\\*.csv')
			latest_file = max(list_of_files, key=os.path.getctime)

			file = open(latest_file, 'r', encoding='UTF8')
			reader = csv.reader(file)
			list_content = []
			for row in reader:
				list_content.append(row)
			for key, content in self.grp_content.items():
				temp_code_list = []
				temp_name_list = []
				temp_size_list = []
				total.append(len(content))
				for item in content:
					index = int(item) - 1
					if list_content[index][0] != 'None':
						temp_code_list.append(f'CAA{list_content[index][0]}')
					else:
						temp_code_list.append('')
					temp_name_list.append(list_content[index][1])
					temp_size_list.append(list_content[index][2])
				name = f"{' / '.join(temp_name_list)} {' '.join(temp_code_list)}"
				size = '-'.join(temp_size_list)
				#print(f'{key=} {name=} {size=}')
				coral_code.append(name)
				coral_size.append(size)
				try:
					base_price.append(int(self.items[key][5].get()))
					if int(self.items[key][5].get()) % 20 != 0:
						tk.messagebox.showerror('底價錯誤', f'組合{key}的底價不是20的倍數')
						return False
				except TypeError:
					print(f'組合{key}輸入錯誤數值')
					return False
				except ValueError:
					tk.messagebox.showerror('底價錯誤', f'{key}底價不能為空白')
					return False
				day = (datetime.today() + timedelta(days=1)).day
				date = datetime.strftime(datetime.today() + timedelta(days=1), '%Y-%m-%d')
				fever_date = datetime.strftime(datetime.today() + timedelta(days=3), '%Y-%m-%d')

			start_from = fever_post_number
			post = docx.Document()
			fever_guy_post = docx.Document()
			if fever:
				for i in range(fever_post_number):
					if base_price[i] < 1000:
						per_bid = 20
					else:
						per_bid = 50
					fever_guy_post_format = f'共 {total[i]} 件\n品種：{coral_code[i]}\n尺寸：{coral_size[i]}\n底價：${base_price[i]}\n每口價：${per_bid}\n截止時間：{date} 21:00\n\n交收時間及地點：\n星期二 太子金魚街萬寧 7-7:15pm\n星期日 太子金魚街萬寧4:30-4:45pm\n星期一 新蒲崗門市 5-8pm\n星期六 新蒲崗門市 5-8pm\n如最後5分鐘有人投標，拍賣將延遲5分鐘。\n如20:00結標，19:59有人投標 ，結標時間延遲為20:05。\n\n備注：中標者請在兩星期內完成交收'
					fever_guy_post.add_paragraph(fever_guy_post_format)
			for i in range(start_from, len(coral_code)):
				if base_price[i] < 1000:
					per_bid = 20
				else:
					per_bid = 50
				post_format = emoji.emojize(
					f':star:{day}號晚上9點完標:star:\n:red_exclamation_mark: 請注意交收時間、地點及拍賣規則，OK先好落標 :red_exclamation_mark:\n共 {total[i]} 件\n品種：{coral_code[i]}\n尺寸：{coral_size[i]}\n底價：${base_price[i]}\n每口價：${per_bid}\n截止時間：{date} 21:00\n\n交收時間及地點：\n星期二 太子金魚街萬寧 7-7:15pm\n星期日 太子金魚街萬寧4:30-4:45pm\n星期一 新蒲崗門市 5-8pm\n星期六 新蒲崗門市 5-8pm\n標準完標時間每晚9點，如最後15分鐘有人投標，拍賣將延遲15分鐘，無限次延遲，直到最後一個投標時間超過15分鐘為之結束。\n\n如21:00完標，20:50有人投標 ，完標時間延遲至21:05；21:03再有人投標，完標時間延遲至21:18；如21:17有人投標，完標時間延遲至21:32；直到21:32仍無人競標，則21:17的投標者中標。\n\n注意事項:\n1.中標者在我哋同事提供拍賣品中標連結後，請盡早落單\n2.中標者必須在兩星期內取貨\n3.如需hold貨超過兩星期，必需先支付貨款，不包生死，生死自付，請在拍賣前確保能在特定的時間及地點取貨\n4.如中標者無合理解釋下棄標或兩星期內沒有取貨，除會被踢出Band Group，取消其拍賣資格外，環球水族保留法律追討的權利\n5.相機影相始終或多或少點都有啲色差，我哋已盡量影接近實物顏色，介意勿拍')
				post.add_paragraph(post_format)
				# pprint(post_format)

			dir = wwaglobal.dropboxPath + str(
			    datetime.today().year) + '\\Auction Coral Retial\\'
			# fever_guy_path = wwaglobal.dropboxPath + str(datetime.today().year) + '\\Auction Coral Retial\\發燒友docx\\'
			# if os.path.isdir(path):
			#     # path already exists, remove it
			#     wwalog.log("Path \"" + path + "\" already exists, trying to remove it")
			#     shutil.rmtree(path)
			# os.mkdir(path)
			# os.mkdir(path + '\\Picture')
			# post.save(path + '\\品種.docx')
			# if fever:
			#     fever_guy_post.save(fever_guy_path + wwaglobal.today + '.docx')
			#dir = r"C:\Users\user\World Wide Aquarium Dropbox\WWA SHKS\PROGRAM\infoInput\test_template_folder"
			path = os.path.join(dir, datetime.strftime(datetime.today(), '%Y-%m-%d') + ' Auction Coral')
			fever_guy_path = os.path.join(dir, '發燒友docx') + '\\' + datetime.strftime(datetime.today(), '%Y-%m-%d') + '.docx'
			if os.path.isdir(path):
				# path already exists, remove it
				# wwalog.log("Path \"" + path + "\" already exists, trying to remove it")
				while 1:
					try:
						shutil.rmtree(path)
						break
					except PermissionError:
						continue
			if os.path.isdir(fever_guy_path):
				while 1:
					try:
						shutil.rmtree(fever_guy_path)
						break
					except PermissionError:
						continue
			os.mkdir(path)
			os.mkdir(path + '\\Picture')
			post.save(os.path.join(path + '/品種.docx'))
			if fever:
				fever_guy_post.save(fever_guy_path)

			auction_pictures = os.listdir(auction_picture_path)
			for folder in auction_pictures:
				# print(f'{temp_folder=}')
				src = os.path.join(auction_picture_path, folder)
				dst = os.path.join(path, 'Picture', folder)
				# print(f'{dst=}')
				shutil.copytree(src, dst, copy_function=shutil.copy)

			counter = 0
			f = open(path + "\\" + datetime.strftime(datetime.today(), '%Y-%m-%d') + " Auction Coral CSV.csv", "w+",
					 encoding="utf-8")
			f.write(
				"連結,商品編號,名稱,介紹,商品選項 1 - 種類,商品選項 1 - 名稱,商品選項 2 - 種類,商品選項 2 - 名稱,商品選項 3 - 種類,商品選項 3 - 名稱,商品選項相片連結,商品選項影片連結,類別,原價,折扣價,會員價,存貨,購買上限,重量 (公斤),SKU,相片網址,影片網址,Hashtags,上架狀態\n")
			if fever:
				counter = fever_post_number
			for i in range(counter, len(coral_code)):
				temp = ",,拍賣品 請勿購買  "
				if coral_code[i].find('CAA') == -1:
					# 海葵
					temp += coral_code[i] + ' ' + coral_size[i] + ",,,,,,,,,,海葵 / 拍賣品,"
				else:
					if coral_size[i].find("可參考(frag座2.5cm)") != -1:
						# frag座
						name = coral_code[i][:coral_code[i].find('CAA') - 1]
						temp += name + coral_size[i] + " " + coral_code[i][
															 coral_code[i].find('CAA') - 1:] + " ,,,,,,,,,,珊瑚 / 拍賣品,"
					else:
						# have size
						temp += coral_size[i] + ' ' + coral_code[i] + " ,,,,,,,,,,珊瑚 / 拍賣品,"

				temp += "[auction_price],,,1,,,,[photo_link],,,1\n"

				f.write(temp)

			f.close()

			pprint(f'{self.grp_content=}')
			grp_csv_path = os.path.join(path, datetime.strftime(datetime.today(), '%Y-%m-%d') + 'Photo Group.csv')
			file = open(grp_csv_path, 'w+', encoding='utf-8', newline='')
			writer = csv.writer(file)
			for key, content in self.grp_content.items():
				temp_list = [key]
				for item in content:
					temp_list.append(item)
				writer.writerow(temp_list)

			file.close()

			paragraph_number = len(post.paragraphs)
			#print(f'{paragraph_number}')

			fever_count = 0

			values = [
			    [""]  # skip one line
			]

			for i in range(self.groups_count):
				row = []
				today = str(datetime.today().strftime('%Y')) + "/" + str(datetime.today().strftime('%m')) + "/" + str(
					datetime.today().strftime('%d'))
				row.append(today)  # e.g. "2021/9/30"
				row.append("a")

				if fever and i <= fever_post_number - 1:
					end_date = datetime.today() + timedelta(days=3)
				else:
					end_date = datetime.today() + timedelta(days=1)
				end_date_str = str(end_date.strftime('%Y')) + "/" + str(end_date.strftime('%m')) + "/" + str(
					end_date.strftime('%d'))
				row.append(end_date_str)
				row.append("2100")

				if coral_code[i].find('CAA') == -1:
					row.append("")
					row.append(coral_code[i])
				else:
					# CAA...
					row.append(coral_code[i][coral_code[i].find('CAA'):])

					row.append(str(coral_code[i][:coral_code[i].find('CAA')-1]))
				if "frag" in str(coral_size[i]):
					row.append("")
				else:
					row.append(str(coral_size[i]).rstrip("寸頭").lstrip('約'))

				row.append(str(base_price[i]))

				if fever and i <= fever_post_number - 1:
					row.append("發燒友拍賣區")
				else:
					row.append("環球水族HK - 拍賣區")

				row.append("y")

				wwalog.log("row " + str(i + 1) + ": " + str(row))
				values.append(row)

			appendToGoogleSheet(values)

			return True

		def complete():
			pprint(f'{self.grp_content=}')
			# validation
			for key, item in self.grp_content.items():
				if len(item) == 0:
					tk.messagebox.showerror('有空白組合', f'組合{key}沒有物件')
					return


			messagebox = tk.messagebox.askyesno('發燒友', '請問要出發燒友拍賣嗎?')
			if create_format(messagebox):
				global quited
				quited = True
				#controller.quit_myself()

				#quit()
				# self.destroy()

		def print_content():
			pprint(f'{self.grp_content=}')

		def create_check_script():
			coral_code = []
			coral_size = []
			base_price = []
			total = []
			list_of_files = glob.glob(data_folder_path + '\\*.csv')
			latest_file = max(list_of_files, key=os.path.getctime)

			file = open(latest_file, 'r', encoding='UTF8')
			reader = csv.reader(file)
			list_content = []
			for row in reader:
				list_content.append(row)
			for key, content in self.grp_content.items():
				temp_code_list = []
				temp_name_list = []
				temp_size_list = []
				total.append(len(content))
				for item in content:
					index = int(item) - 1
					if list_content[index][0] != 'None':
						temp_code_list.append(f'CAA{list_content[index][0]}')
					else:
						temp_code_list.append('')
					temp_name_list.append(list_content[index][1])
					temp_size_list.append(list_content[index][2])
				name = f"{' / '.join(temp_name_list)} {' '.join(temp_code_list)}"
				size = '-'.join(temp_size_list)
				#print(f'{key=} {name=} {size=}')
				coral_code.append(name)
				coral_size.append(size)
				try:
					base_price.append(int(self.items[key][5].get()))
					if int(self.items[key][5].get()) % 20 != 0:
						tk.messagebox.showerror('底價錯誤', f'組合{key}的底價不是20的倍數')
						return False
				except TypeError:
					print(f'組合{key}輸入錯誤數值')
					return False
				except ValueError:
					tk.messagebox.showerror('底價錯誤', f'{key}底價不能為空白')
					return False

			temp_dict = {'編號': coral_code, '尺寸': coral_size, '總數': total, '底價': base_price}
			pprint(temp_dict)
			df = pd.DataFrame(temp_dict)

			dfShow = tk.Toplevel()

			# dfShow.geometry("1500x500")  # set the root dimensions
			w = 1200
			h = 500
			ws = dfShow.winfo_screenwidth()  # width of the screen
			hs = dfShow.winfo_screenheight()  # height of the screen

			# calculate x and y coordinates for the Tk root window
			x = (ws / 2) - (w / 2)
			y = (hs / 2) - (h / 2)

			# set the dimensions of the screen
			# and where it is placed
			dfShow.geometry('%dx%d+%d+%d' % (w, h, x, y))
			dfShow.pack_propagate(False)  # tells the root to not let the widgets inside it determine its size.
			dfShow.resizable(0, 0)  # makes the root window fixed in size.

			# Frame for TreeView
			frame1 = tk.LabelFrame(dfShow, text="Data")
			frame1.place(height=500, width=1200)

			copy_button = tk.Button(dfShow, text='Copy', command=lambda: copy())
			copy_button.pack()

			## Treeview Widget
			tv1 = ttk.Treeview(frame1)
			tv1.place(relheight=1,
					  relwidth=1)  # set the height and width of the widget to 100% of its container (frame1).

			treescrolly = tk.Scrollbar(frame1, orient="vertical",
									   command=tv1.yview)  # command means update the yaxis view of the widget
			treescrollx = tk.Scrollbar(frame1, orient="horizontal",
									   command=tv1.xview)  # command means update the xaxis view of the widget
			tv1.configure(xscrollcommand=treescrollx.set,
						  yscrollcommand=treescrolly.set)  # assign the scrollbars to the Treeview Widget
			treescrollx.pack(side="bottom", fill="x")  # make the scrollbar fill the x axis of the Treeview widget
			treescrolly.pack(side="right", fill="y")  # make the scrollbar fill the y axis of the Treeview widget

			def clear_data():
				tv1.delete(*tv1.get_children())
				return None

			clear_data()
			tv1["column"] = list(df.columns)
			tv1["show"] = "headings"
			for column in tv1["columns"]:
				tv1.heading(column, text=column)  # let the column heading = column name

			df_rows = df.to_numpy().tolist()  # turns the dataframe into a list of lists
			for row in df_rows:
				tv1.insert("", "end",
						   values=row)  # inserts each list into the treeview.

			tv1.column(0, width=700)
			tv1.column(1, width=250)
			tv1.column(2, width=25)
			tv1.column(3, width=25)
			def copy():
				image = ImageGrab.grab((x, y, x+w, y+h))
				output = BytesIO()
				image.convert('RGB').save(output, 'BMP')
				data = output.getvalue()[14:]
				output.close()
				clip.OpenClipboard()
				clip.EmptyClipboard()
				clip.SetClipboardData(win32con.CF_DIB, data)
				clip.CloseClipboard()
				dfShow.quit()
				return None
			dfShow.mainloop()


		add_group(self.groups_count)
		add_button = tk.Button(self, text='新增', command=lambda: add_group(self.groups_count), pady=5)
		add_button.grid(row=2, column=0, sticky='EW')
		complete_button = tk.Button(self, text='完成', pady=5, command=lambda: complete())
		complete_button.grid(row=3, column=0, sticky='EW')
		test_button = tk.Button(self, text='test', command=lambda: create_check_script())
		test_button.grid(row=4, column=0, sticky='EW')

def main():
	global quited
	global input_window
	#pprint(f'{quited=}')
	input_window = InputApp()
	while not quited:
		input_window.update_idletasks()
		input_window.update()
		#pprint(f'{quited=}')
	#input_window.mainloop()
	input_window.quit()
	#input_window.destroy()
	print('ended')
