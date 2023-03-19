import os
import win32api
import mouse
import docx
from datetime import datetime
from datetime import timedelta
import pyautogui
from keyboard import press
from bs4 import BeautifulSoup
from pynput.keyboard import Key, Controller
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
import datetime
import re
import wwaglobal
import wwalog
from selenium.webdriver.chrome.options import Options
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import wwautils




#check is the bidding valid
import wwautils


def checkBidding(bidPrice, reservedPrice):
	if bidPrice < int(reservedPrice) or bidPrice % 20 != 0:
		return False
	else:
		return True

#check the bid price higher than last time or not
# def checkLastbid(lastbid, thisbid):
#     if thisbid>lastbid:
#         return True
#     else: return False

#check if the comment time is over 2100. If yes, return True.
def checkTimeOvered(commentTime,date_now,endTime):

	now = datetime.datetime.now()
	return (commentTime > endTime)

#change comment time format
def commentTimeFormatting(comment):
	temp = re.findall(r'\d+', comment)
	res = list(map(int, temp))
	dateFormat = '%Y-%M-%d %H:%m:%S'
	year = res[0]
	month = res[1]
	day = res[2]
	hour = res[3]
	if comment.split('日')[1].split('午')[0].strip() == '下':
		if hour != 12:
			hour += 12
	minutes = res[4]
	if month <10:
		monthstr = '0' + str(month)
	else: monthstr = str(month)
	if day < 10:
		daystr = '0' + str(day)
	else: daystr = str(day)
	if hour < 10:
		hourstr = '0' + str(hour)
	else: hourstr = str(hour)
	if minutes < 10:
		minutesstr = '0' + str(minutes)
	else: minutesstr = str(minutes)


	formattedTime = str(year) + "-" + monthstr + "-" + daystr + ' ' + hourstr + ":" + minutesstr + ':' +'00'
	timeFormat = '%Y-%m-%d %H:%M:%S'
	#print(formattedTime)
	return datetime.datetime.strptime(formattedTime, timeFormat)

	#print(datetime.datetime.strptime(formattedTime, dateFormat))



def findbidder(date):
	wwaglobal.init()

	# date = input("Date (yyyy-mm-dd): ")
	date_now = datetime.datetime.strptime(date, '%Y-%m-%d')
	#date = datetime.datetime.strftime(date_now, '%Y-%m-%d')
	postformatPath = wwaglobal.dropboxPath + str(date_now.year)+"\\Auction Coral Retial\\" + date + " Auction Coral\\品種.docx"
	postDocument = docx.Document(postformatPath)

	endTime = datetime.datetime.strptime(postDocument.paragraphs[0].text.split('截止時間：')[1].split('\n')[0], '%Y-%m-%d %H:%M')
	print(endTime)
	#print(date_now + timedelta(days=3))
	#modifier
	documentPath = wwaglobal.dropboxPath + str(date_now.year)+"\\Auction Coral Retial\\" + date + " Auction Coral\\AuctionDoc.docx"
	webdriverPath = wwaglobal.chromeDriverPath
	commentFormatSavingPath = wwaglobal.dropboxPath + str(date_now.year)+ "\\Auction Coral Retial\\" + date + " Auction Coral\\comment.docx"



	#find bidder and bid price
	linkDocument = docx.Document(documentPath)
	link = []
	startbid = []
	finalBidder = []
	finalPrice = []

	now = datetime.datetime.now()

	#newTime = now.replace(day=date_now.day, hour=23, minute=0,second=0,microsecond=0)
	#print(str(newTime))

	#print(time.strftime("%Y{y}%m{m}%d{d} %H{h}%M", newTime).format(y="年", m="月", d="日", h=":"))

	for i in range(2, len(linkDocument.paragraphs)-2, 3):
		if 'https://band.us/band/75420559/post/' in linkDocument.paragraphs[i].text.split('拍賣連結：')[1].strip():
			continue
		link.append(linkDocument.paragraphs[i].text.split('拍賣連結：')[1].strip())
	for i in range(1, len(linkDocument.paragraphs)-2, 3):
		if 'https://band.us/band/75420559/post/' in linkDocument.paragraphs[i+1].text.split('拍賣連結：')[1].strip():
			continue
		startbid.append(linkDocument.paragraphs[i].text.split('起標價：$')[1].strip())
	print(startbid)
	commentText = []
	commentName = []
	commentTime = []

	options = Options()
	options.headless = False

	#print("Copying Chrome profile.")
	#wwalog.log("Copying Chrome profile.")
	#os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")

	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
	original_end_time = endTime
	for i in range(len(link)):
		import time
		endTime = original_end_time
		url = link[i]
		driver.get(url)
		time.sleep(10)
		commentText.clear()
		commentTime.clear()
		#get all comments' bidding and time

		soup = BeautifulSoup(driver.page_source, "lxml")
		for el in soup.find_all('p', attrs={'class': 'txt _commentContent'}):
			commentText.append(el.get_text())

		for time in soup.find_all('time', attrs={'class': 'time'}):
			commentTime.append(str(time).split('title="')[1].split('"')[0])

		#print(commentTime)

		#print(checkTimeOvered(commentTimeFormatting(commentTime[0])))


		#print(datetime.datetime.strptime(commentTime[0], '%Y年%M月%d日 下午%H:%m'))
		# commentText.append(soup.find("p", {"class": "txt _commentContent"}).getText())
		# commentText.append(soup.findNext("p", {"class": "txt _commentContent"}).getText())
		bidder = "No one"
		bidTemp = 0
		for name in soup.find_all('strong', attrs={'class': 'name'}):
			commentName.append(name.get_text())
		for j in range(len(commentText)):
			#print([int(i) for i in commentText[j].split() if i.isdigit()])
			temp = re.findall(r'\d+', commentText[j])
			commentBid = list(map(int, temp))
			if not len(commentBid)<=0:
				if endTime-timedelta(minutes=15) <= commentTimeFormatting(commentTime[j]) <= endTime:
					endTime += timedelta(minutes=15)
					print(f'End Time Extended for {link[i]}')
				# print('comment ' + str(j+1))
				# print('time: ')
				#print(checkTimeOvered(commentTimeFormatting(commentTime[j])))
				#print(commentBid[0])
				# print('check bid: ')
				#print(checkBidding(commentBid[0], startbid[i]))
				if (not (checkTimeOvered(commentTimeFormatting(commentTime[j]), date_now, endTime)) and checkBidding(commentBid[0], startbid[i])):
					if bidTemp < commentBid[0]:
						bidder = commentName[j]
						bidTemp = commentBid[0]
					#print(res)
		finalBidder.append(bidder)
		finalPrice.append(bidTemp)
		print(bidder + ' ' + str(bidTemp))

		commentName.clear()
	driver.quit()

	commentFormat = docx.Document()
	bidderVerificationFormat = '恭喜師兄[name]你以(price)中標，為確保交收/自取順利，請先點擊以下連結進入小店網站落單購買拍賣品，落單時有得揀返交收或自取時間😊 \n拍賣品連結：'
	appreciationMessage= '多謝各位師兄師姐嘅參與，由於band要大約過咗24~48小時先睇到確實嘅留言日期同時間，我哋會盡快確認中標者🙏'
	try:
		for i in range(len(link)):
			if finalPrice[i] != 0:
				try:
					commentFormat.add_paragraph((bidderVerificationFormat.split('(price)')[0] + '$' + str(finalPrice[i]) + bidderVerificationFormat.split('(price)')[1] + '\n').replace('[name]',finalBidder[i]))
				except:
					commentFormat.add_paragraph((bidderVerificationFormat.split('(price)')[0] + '$' + str(finalPrice[i]) +
												 bidderVerificationFormat.split('(price)')[1] + '\n'))
			else:
				commentFormat.add_paragraph('無人投標\n')
		commentFormat.save(commentFormatSavingPath)

	except:
		print('error message: \ncant append value in comment doc')
	try:
		wwautils.appendGoogleAuction(date_now, finalBidder, finalPrice)
	except:
		print('cant append to Auction List')

	return finalBidder




#find the link of the auction item
# driver = webdriver.Chrome(webdriverPath)
# for i in range(1):
#     import time
#     url = 'https://wwa.boutir.com/?q=' + itemName
#     driver.get(url)
#     WebDriverWait(driver, 10).until(
#         EC.presence_of_element_located((By.CLASS_NAME, 'search-result-title'))
#     )
#     driver.find_element_by_class_name('clickable').click()
#     WebDriverWait(driver, 10).until(
#         EC.presence_of_element_located((By.CLASS_NAME, 'addthis_button_more'))
#     )
#
