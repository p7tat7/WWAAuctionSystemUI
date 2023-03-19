import csv

import wwautils
import wwaglobal
from datetime import datetime
from datetime import timedelta
import os
from pprint import pprint
import gspread
import wwaformatcreate
#from gspread.models import Cell
from oauth2client.service_account import ServiceAccountCredentials
import getpass
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import TimeoutException
import pickle
import random
from bs4 import BeautifulSoup
import wwaformatcreate
import wwaglobal
import wwalog
import os
import docx
from distutils.dir_util import copy_tree
import shutil
import dropbox
import re
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from collections import Counter
from operator import itemgetter
from datetime import datetime
from datetime import timedelta
import xlrd
from pprint import pprint
#from Google import Create_Service
from selenium.webdriver.remote.webelement import WebElement
import wwatkmenu
#from gspread.models import Cell
from openpyxl import load_workbook
from openpyxl.styles import Border, Side
from dropbox import common
#from gspread.models import Cell
from ppadb.client import Client as AdbClient

def test_post_auction():

    # postNumber = int(input("貼文數量 (環球水族HK - 拍賣區): "))
    # postNumber2 = int(input("貼文數量 (發燒友拍賣區): "))
    # testing

    wwaformatcreate.main()
    time.sleep(5)

    # check if today's document is ready
    while True:
        try:
            if not os.path.isfile(wwaglobal.auctionDocumentPath):
                # no document
                print("Today's auction document not found: " + wwaglobal.auctionDocumentPath + "\nexitting...")
                wwalog.log("Today's auction document not found: " + wwaglobal.auctionDocumentPath + "\nexitting...")
                temp = input('wait')
            else:
                break
        except PermissionError:
            continue

    time.sleep(5)


    if os.path.isfile(wwaglobal.dropboxPath + str(
            datetime.today().year) + "\\Auction Coral Retial\\發燒友docx\\" + wwaglobal.today + ".docx"):
        auctionDocument2 = docx.Document(wwaglobal.dropboxPath + str(
            datetime.today().year) + "\\Auction Coral Retial\\發燒友docx\\" + wwaglobal.today + ".docx")
        postNumber2 = len(auctionDocument2.paragraphs)
    else:
        postNumber2 = 0

    auctionDocument = docx.Document(wwaglobal.auctionDocumentPath)
    postNumber = len(auctionDocument.paragraphs)


    # path = wwaglobal.dropboxPath + str(
    # 	datetime.today().year) + '\\Auction Coral Retial\\' + wwaglobal.today + ' Auction Coral'
    while 1:
        photoList = os.listdir(wwaglobal.photoPath)
        pprint(photoList)
        if len(photoList) > 0:
            break
        else:
            print('Loading Photos')
            temp = input('wait')
            # os.system("Xcopy \"Picture\" \"" + (path + '\\Picture') + "\" /E/H/C/I/Y/Q > nul 2>&1")
            time.sleep(5)

    time.sleep(1)

    picture_folders = os.listdir(wwaglobal.photoPath)
    for index, folder in enumerate(picture_folders):
        full_path = os.path.join(wwaglobal.photoPath, folder)
        if os.path.isdir(full_path):
            picture_folders[index] = full_path
    pprint(f'{picture_folders=}')
            
    file = open(wwaglobal.grp_content_csv_path, 'r', encoding='utf-8')
    reader = csv.reader(file)
    grp_photo_path = {}
    for row in reader:
        grp_photo_path[row[0]] = []
        #print(f'{row=}')
        for index, item in enumerate(row):
            if index == 0:
                continue
            #print(f'{item=}')
            grp_photo_path[row[0]].append(picture_folders[int(item)-1])
    file.close()

    time.sleep(5)


    options = Options()
    options.headless = False

    # print("Copying Chrome profile.")
    # wwalog.log("Copying Chrome profile.")
    # os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")

    options = webdriver.ChromeOptions()
    options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
    options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

    driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
    driver.get("https://band.us/")


    # check if login band success
    while 1:
        try:
            bt = WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.CLASS_NAME, "bandCreate")))
            wwalog.log("login band success.")
            break
        except TimeoutException:
            pass

    # Save cookies
    # pickle.dump(driver.get_cookies(), open(cookie_file_name, "wb"))

    coralType = []
    startBid = []

    # 發燒友拍賣區
    auctionDocPath = wwaglobal.dropboxPath + str(
        datetime.today().year) + "\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\AuctionDoc.docx"
    auctionDocMessage = ""
    if postNumber2 > 0:
        # 獲取最後一個post的編號 (發燒友拍賣區（珊瑚、海水魚、器材用品）)
        if wwaglobal.debug:
            driver.get("")
        else:
            driver.get("https://band.us/band/75420559/")
            # driver.get("https://band.us/band/86026020/")

        firstPostCode2 = -1
        while firstPostCode2 == -1:
            try:
                soup = BeautifulSoup(driver.page_source, "lxml")
                a_tags = soup.find_all('a')
                link = ""
                for tag in a_tags:
                    if tag.get('href').startswith('https://band.us/band/75420559/post/'):
                    # if tag.get('href').startswith('https://band.us/band/86026020/post/'):
                        link = tag.get('href')
                        break
                firstPostCode2 = int(link.split('post/')[1])
            except IndexError:
                time.sleep(0.5)

        # auctionDocument2 = docx.Document(wwaglobal.dropboxPath + str(
        #     datetime.today().year) + "\\Auction Coral Retial\\發燒友docx\\" + wwaglobal.today + ".docx")
        # auctionDocument2 = docx.Document(wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\發燒友doc\\2021-09-30.docx")

        for i in range(postNumber2):
            coralType.append(
                auctionDocument2.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
            startBid.append(auctionDocument2.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())


            time.sleep(2)

            try:
                # click create post button
                driver.find_element_by_css_selector(
                    "[class*='cPostWriteEventWrapper'][class*='_btnOpenWriteLayer']").click()
            except:
                time.sleep(0.5)
            time.sleep(2)

            # find write post input
            while 1:
                try:
                    postTextInput = driver.find_element_by_css_selector(
                        "div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")
                    break
                except:
                    time.sleep(1)

            # write post text
            postTextInput.send_keys(auctionDocument2.paragraphs[i].text)

            time.sleep(0.5)

            # find file input
            postFileInput = driver.find_element_by_css_selector("input[type='file'][id^=postPhotoInput_view]")

            # upload file
            # print(wwaglobal.photoPath)
            # print(i)
            temp_picture_list = []
            for path in grp_photo_path[str(i+1)]:
                for photo in os.listdir(path):
                    temp_picture_list.append(os.path.join(path, photo))
            path_string = '\n'.join(temp_picture_list)
            postFileInput.send_keys(path_string)
            while 1:
                try:
                    bt = WebDriverWait(driver, 1).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
                    # click attach button
                    # bt.click()
                    time.sleep(1)
                    bt.send_keys(Keys.RETURN)
                    break
                except TimeoutException:
                    pass
            #wwalog.log("發燒友拍賣區: Uploaded file " + wwaglobal.photoPath + '\\' + photoList[i])
            time.sleep(len(temp_picture_list) * 2)
            # find attach button

            # find post button
            while 1:
                try:
                    bt = WebDriverWait(driver, 2).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton._btnSubmitPost.-confirm")))
                    time.sleep(1)

                    # click post button
                    # bt.click()
                    bt.send_keys(Keys.RETURN)
                    wwalog.log("Finished 發燒友拍賣區 post " + str(i + 1) + ".")
                    time.sleep(2)
                    break
                except TimeoutException:
                    try:
                        bt = WebDriverWait(driver, 2).until(
                            EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
                        # click attach button
                        # bt_temp = driver.find_element_by_xpath('//*[@id="wrap"]/div[2]/div[2]/section/div/footer/button[2]')
                        time.sleep(1)
                        # bt_temp.click()

                        bt.send_keys(Keys.RETURN)
                    except:
                        pass

            # temp_wait = input('press post button')
            wwalog.log("\n")
    time.sleep(1)
    # 獲取最後一個post的編號
    driver.get("https://band.us/band/78427905/")
    # driver.get("https://band.us/band/86026020/")

    firstPostCode = -1
    while firstPostCode == -1:
        try:
            soup = BeautifulSoup(driver.page_source, "lxml")
            a_tags = soup.find_all('a')
            link = ""
            for tag in a_tags:
                if tag.get('href').startswith('https://band.us/band/78427905/post/'):
                # if tag.get('href').startswith('https://band.us/band/86026020/post/'):
                    link = tag.get('href')
                    break
            firstPostCode = int(link.split('post/')[1])
        except IndexError:
            time.sleep(0.5)
    # firstPostCode =2212
    wwalog.log("firstPostCode: " + str(firstPostCode) + "\n")
    auctionDocMessage += '❗最新珊瑚拍賣❗\n'
    time.sleep(0.1)
    auctionDocument = docx.Document(wwaglobal.auctionDocumentPath)
    print(f'{postNumber2=} {postNumber=}')
    try:
        # create a file for sharing post
        f = open(
            wwaglobal.dropboxPath + str(
                datetime.today().year) + "\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\today_auction_links.txt",
            "w+", encoding="utf-8")
        for i in range(postNumber2):
            auctionDocMessage += coralType[i] + " 起標價：" + startBid[i] + "\n"
            auctionDocMessage += "拍賣連結： https://band.us/band/75420559/post/" + str(
                int(firstPostCode2) + i + 1) + "\n\n"
            # auctionDocMessage += "拍賣連結： https://band.us/band/86026020/post/" + str(
            #     int(firstPostCode2) + i + 1) + "\n\n"
            time.sleep(0.1)

            f.write("https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")
            # f.write("https://band.us/band/86026020/post/" + str(int(firstPostCode2) + i + 1) + "\n")

        for i in range(postNumber):
            # each post info
            coralType.append(
                auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
            startBid.append(auctionDocument.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())
            auctionDocMessage += coralType[postNumber2 + i] + " 起標價：" + startBid[postNumber2 + i] + "\n"
            auctionDocMessage += "拍賣連結： https://band.us/band/78427905/post/" + str(
                int(firstPostCode) + i + 1) + "\n\n"
            # auctionDocMessage += "拍賣連結： https://band.us/band/86026020/post/" + str(
            #     int(firstPostCode) + i + 1) + "\n\n"
            time.sleep(0.1)

            f.write("https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1) + "\n")
            # f.write("https://band.us/band/86026020/post/" + str(int(firstPostCode) + i + 1) + "\n")

        f.close()
    except:
        print('create auction share post unsuccessfully')
    f = open(
        wwaglobal.dropboxPath + str(
            datetime.today().year) + "\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\today_auction_links.txt",
        "w+", encoding="utf-8")
    # try:
    for i in range(postNumber2):
        f.write("https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")
        # f.write("https://band.us/band/86026020/post/" + str(int(firstPostCode2) + i + 1) + "\n")
    for i in range(postNumber):
        f.write("https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1) + "\n")
        # f.write("https://band.us/band/86026020/post/" + str(int(firstPostCode) + i + 1) + "\n")
    # except:
    #     print('create link file unsuccessfully')
    f.close()
    auctionDoc = docx.Document()
    auctionDocLines = auctionDocMessage.split("\n")
    for i in range(len(auctionDocLines)):
        auctionDoc.add_paragraph(auctionDocLines[i])
    auctionDoc.save(auctionDocPath)

    temp_count = 0
    # print(photoList)
    for i in range(postNumber):
        # coralType.append(auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
        # coralAll = auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[
        #     0].strip().split(' / ')
        # no_of_coral = len(coralAll)

        # startBid.append(auctionDocument.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())

        while 1:
            try:
                # click create post button
                driver.find_element_by_css_selector(
                    "[class*='cPostWriteEventWrapper'][class*='_btnOpenWriteLayer']").click()
                # time.sleep(2)
                # postTextInput = driver.find_element_by_css_selector(
                # "div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")
                break
            except:
                time.sleep(0.5)
        time.sleep(4)

        postTextInput = driver.find_element_by_css_selector(
            "div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")
        time.sleep(0.5)
        postTextInput.send_keys(auctionDocument.paragraphs[i].text)
        time.sleep(0.5)

        # find file input
        postFileInput = driver.find_element_by_css_selector("input[type='file'][id^=postPhotoInput_view]")

        # upload file
        # print(temp_count)
        temp_index = str(i+1)

        temp_picture_list = []
        for path in grp_photo_path[str(i + 1 + postNumber2)]:
            for photo in os.listdir(path):
                temp_picture_list.append(os.path.join(path, photo))
        #print(f'{temp_picture_list=}')
        path_string = '\n'.join(temp_picture_list)


        postFileInput.send_keys(path_string)


        wwalog.log("Uploaded file " + wwaglobal.photoPath + '\\' + photoList[postNumber2 + i])
        clicked = False
        # find attach button
        while 1:
            try:
                bt = WebDriverWait(driver, 1).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
                # click attach button
                # bt.click()
                time.sleep(1)
                bt.send_keys(Keys.RETURN)
                break
            except TimeoutException:
                pass
        # temp_wait = input('press attach button')

        time.sleep(len(temp_picture_list) * 2)
        # find post button
        while 1:
            try:
                bt = WebDriverWait(driver, 2).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton._btnSubmitPost.-confirm")))
                time.sleep(1)

                # click post button
                # bt.click()
                bt.send_keys(Keys.RETURN)
                #wwalog.log("Finished 發燒友拍賣區 post " + str(i + 1) + ".")
                time.sleep(2)
                break
            except TimeoutException:
                try:
                    bt = WebDriverWait(driver, 2).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
                    # click attach button
                    # bt_temp = driver.find_element_by_xpath('//*[@id="wrap"]/div[2]/div[2]/section/div/footer/button[2]')
                    time.sleep(1)
                    # bt_temp.click()

                    bt.send_keys(Keys.RETURN)
                except:
                    pass
        time.sleep(5)


    time.sleep(1)



    # open web whatsapp
    wwalog.log("Finished posting to band.\nOpening https://web.whatsapp.com/")
    print("Finished posting to band.\nOpening https://web.whatsapp.com/")
    driver.get("https://web.whatsapp.com/")

    # check if login whatsapp success
    while 1:
        try:
            WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
            wwalog.log("login whatsapp success.")
            break
        except TimeoutException:
            pass

    grpList = ['拍賣通知專區']
    # grpList = ['63393947']
    # grpList = ['93258078'] #DEBUG
    counter = 0
    time.sleep(2)

    auctionDocPath = wwaglobal.dropboxPath + str(
        datetime.today().year) + "\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\AuctionDoc.docx"
    auctionDocMessage = ""

    for grpName in grpList:
        # find group
        wwalog.log("Finding group \"" + grpName + "\"")
        while 1:
            try:
                driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(grpName)
                break
            except:
                wwalog.log("Cannot access group name input box of web Whatsapp, trying again...")
                time.sleep(1)
        # driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys('🌊環球水族 WWA🌊會員通知區')
        time.sleep(1)
        driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(Keys.RETURN)
        time.sleep(2)

        # type message
        wwalog.log("Finding input box.")
        inputbox = driver.find_element_by_css_selector("div[role='textbox'][spellcheck='true']")
        # inputbox.send_keys('❗最新珊瑚拍賣❗')
        inputbox.send_keys(':bell')
        inputbox.send_keys(Keys.RETURN)
        inputbox.send_keys('最新拍賣')
        inputbox.send_keys(':bell')
        inputbox.send_keys(Keys.RETURN)

        inputbox.send_keys(':fire')
        inputbox.send_keys(Keys.RETURN)
        time.sleep(0.5)
        inputbox.send_keys(Keys.ARROW_LEFT)
        time.sleep(0.5)
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        time.sleep(0.5)
        inputbox.send_keys(Keys.ARROW_RIGHT)
        time.sleep(0.5)
        inputbox.send_keys('標準完標時間每晚9點')
        inputbox.send_keys(':fire')
        inputbox.send_keys(Keys.RETURN)
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys('多款珊瑚、魚種任你拍賣，最終價錢由你話事！!')
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys('快啲加入我地拍賣Band Page參加拍賣啦：')
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys('https://band.us/n/ada260kfmfib4')
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys('想第一時間知道最新拍賣內容？咁就要加入我地既拍賣whatsapp群組啦：')
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys('https://chat.whatsapp.com/EyGWGzUlNqIKiIzk5lXOxu')
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)

        wwalog.log("\n❗最新珊瑚拍賣❗")

        # auctionDocMessage += '❗最新珊瑚拍賣❗\n'
        time.sleep(0.1)
        inputbox.send_keys(Keys.SHIFT + Keys.RETURN)

        # create a file for sharing post
        # f = open(wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\today_auction_links.txt", "w+", encoding = "utf-8")
        for i in range(postNumber2):
            wwalog.log("Generating msg for post " + str(i + 1) + ".")
            time.sleep(0.1)
            inputbox.send_keys(coralType[i] + " 起標價：" + startBid[i])
            wwalog.log(coralType[i] + " 起標價：" + startBid[i])
            # auctionDocMessage += coralType[i] + " 起標價：" + startBid[i] + "\n"
            inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
            inputbox.send_keys("拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1))
            wwalog.log("拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1))
            # auctionDocMessage += "拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n\n"
            inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
            inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
            time.sleep(0.1)

        for i in range(postNumber):
            # each post info
            wwalog.log("Generating msg for post " + str(i + postNumber2 + 1) + ".")
            time.sleep(0.1)
            inputbox.send_keys(coralType[postNumber2 + i] + " 起標價：" + startBid[postNumber2 + i])
            wwalog.log(coralType[postNumber2 + i] + " 起標價：" + startBid[postNumber2 + i])
            auctionDocMessage += coralType[postNumber2 + i] + " 起標價：" + startBid[postNumber2 + i] + "\n"
            inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
            inputbox.send_keys("拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1))
            wwalog.log("拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1))
            auctionDocMessage += "拍賣連結： https://band.us/band/78427905/post/" + str(
                int(firstPostCode) + i + 1) + "\n\n"
            inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
            inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
            time.sleep(0.1)

        # send message
        driver.find_element_by_css_selector("span[data-testid='send'][data-icon='send']").find_element_by_xpath(
            "..").send_keys(Keys.RETURN)
        wwalog.log("Sent message.")
        time.sleep(0.5)

        try:
            time.sleep(1)

            while 1:
                try:
                    WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
                    wwalog.log("login whatsapp success.")
                    break
                except TimeoutException:
                    time.sleep(0.5)
            # print(len(photoList))
            photoList = []
            for paths in grp_photo_path.values():
                #print(f'{paths=}')
                for path in paths:
                    #print(f'{path=}')
                    for photo in os.listdir(path):
                        #print(f'{os.path.join(path, photo)=}')
                        photoList.append(os.path.join(path, photo))
            #pprint(f'{photoList}')
            #temp = input('wait')
            uploaded = 0
            for i in range(len(photoList) // 30 + 1):
                temp_photo_list = ''

                if (len(photoList) - uploaded) // 30 >= 1:
                    photo_limit = 30
                else:
                    photo_limit = len(photoList) % 30
                # print(photo_limit)
                for j in range(photo_limit):
                    temp_photo_list += photoList[i * 30 + j]

                    temp_photo_list += '\n'
                    uploaded += 1

                pprint(temp_photo_list[:len(temp_photo_list) - 1])
                WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "span[data-testid='clip'][data-icon='clip']")))
                counter_attach = 0
                while 1:
                    try:
                        driver.find_element_by_css_selector(
                            "span[data-testid='clip'][data-icon='clip']").find_element_by_xpath(
                            "..").send_keys(Keys.RETURN)
                        break
                    except:
                        print('Try Clicking Button')
                        counter_attach += 1
                        time.sleep(1)
                    if counter_attach == 5:
                        break
                time.sleep(1)

                driver.find_element_by_xpath(
                    "//input[@accept='image/*,video/mp4,video/3gpp,video/quicktime']").find_element_by_xpath(
                    "../input").send_keys(temp_photo_list[:len(temp_photo_list) - 1])

                while 1:
                    try:
                        WebDriverWait(driver, 20).until(
                            EC.presence_of_element_located((By.XPATH, "//span[@data-testid='send']")))
                        break
                    except TimeoutException:
                        time.sleep(1)
                time.sleep(1)
                driver.find_element_by_xpath("//div[@role='button' and @class='_165_h _2HL9j']").send_keys(
                    Keys.RETURN)
                time.sleep(3)
            wwalog.log("Sent photos.")
        except:
            print('Error in sending images, please do it on hand.')
        time.sleep(5)




    driver.quit()

    print("Finished.")

def test_whatsappPhoto():
    # auction_list = ['1234','2345','3456']
    # placement_list = [['01234', '1234', '越南菇'],['04567', '4567', 'Zoas']]
    # area_b = [['1234', '彩頭綠火柴'],['4567','超熒光綠豬腰']]
    # wwatkmenu.placement_window(auction_list,placement_list,area_b)
    #auction_placement(datetime.strptime('2021/12/08', '%Y/%m/%d'))
    while 1:
        photoList = os.listdir(wwaglobal.photoPath)
        if len(photoList) > 0:
            break
        else:
            print('Loading Photos')
            time.sleep(1)
    pprint(photoList)
    options = Options()
    options.headless = False
    options = webdriver.ChromeOptions()
    options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
    options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

    driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
    driver.get('https://web.whatsapp.com')
    while 1:
        try:
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
            wwalog.log("login whatsapp success.")
            break
        except TimeoutException:
            time.sleep(0.5)
    #print(len(photoList))

    uploaded=0
    for i in range(len(photoList)//30 + 1):
        temp_photo_list = ''

        if (len(photoList)-uploaded)//30 >=1:
            photo_limit =30
        else:
            photo_limit = len(photoList)%30
        #print(photo_limit)
        for j in range(photo_limit):
            temp_photo_list += wwaglobal.photoPath + '\\' + photoList[i*30+j]

            temp_photo_list += '\n'
            uploaded +=1
        pprint(temp_photo_list[:len(temp_photo_list)-1])
        WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "span[data-testid='clip'][data-icon='clip']")))
        counter_attach = 0
        while 1:
            try:
                driver.find_element_by_css_selector("span[data-testid='clip'][data-icon='clip']").find_element_by_xpath(
                    "..").send_keys(Keys.RETURN)
                break
            except:
                print('Try Clicking Button')
                counter_attach += 1
                time.sleep(1)
            if counter_attach == 5:
                break
        time.sleep(0.5)
        temp = input('wait')

        driver.find_element_by_xpath(
            "//input[@accept='image/*,video/mp4,video/3gpp,video/quicktime']").send_keys(temp_photo_list[:len(temp_photo_list)-1])

        while 1:
            try:
                WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, "//span[@data-testid='send']")))
                break
            except TimeoutException:
                time.sleep(1)
        time.sleep(1)
        driver.find_element_by_xpath("//div[@role='button' and @class='_165_h _2HL9j']").send_keys(Keys.RETURN)
        time.sleep(3)

def dropbox_test():
    f = open("dropbox_token.txt", "r", encoding="utf-8")
    temp = f.read().split("\n")
    f.close()
    DROPBOX_TOKEN = temp[0]

    wwalog.log("Trying to connect Dropbox api.")
    # app_key = "suxigrgh0rh70sg"
    # app_secret = "5vrzdcd01zu6yua"
    #
    # # build the authorization URL:
    # authorization_url = "https://www.dropbox.com/oauth2/authorize?client_id=%s&response_type=code" % app_key
    #
    # # send the user to the authorization URL:
    # print('Go to the following URL and allow access:')
    # print(authorization_url)
    #
    # # get the authorization code from the user:
    # authorization_code = raw_input('Enter the code:\n')
    #
    # # exchange the authorization code for an access token:
    # token_url = "https://api.dropboxapi.com/oauth2/token"
    # params = {
    # 	"code": authorization_code,
    # 	"grant_type": "authorization_code",
    # 	"client_id": app_key,
    # 	"client_secret": app_secret
    # }

    dbx = dropbox.Dropbox(DROPBOX_TOKEN).with_path_root(common.PathRoot.namespace_id('4'))
    #dbx.users_get_current_account()
    #.as_user('SHKS WWA')
    #dbx = common.PathRoot.namespace_id('4')
    #dbx.with_path_root(dropbox.common.PathRoot.root(namespace=4)).files_list_folder("")
    #print(dbx.sharing_create_shared_link_with_settings('/Temporary Store for Posting/01.JPG'))
    # try:
    print(dbx.sharing_create_shared_link_with_settings('/Coral Data/Picture/CAA4407 2021-12-07.JPG').as_user('SHKS WWA'))
    # except:
    # 	print('1 unable')
    try:
        print(dbx.sharing_create_shared_link_with_settings('/../WWA SHKS/Coral Data/Picture/CAA4407 2021-12-07.JPG'))
    except:
        print('2 unable')

    # dbx.Dropbox.files_get_temporary_link(r'C:\Users\p7tat7\World Wide Aquarium Dropbox\WWA SHKS\Coral Data\Picture\CAA4207 2021-12-07.JPG')
    pass

def append_google_meeting_test():
    delivery = False
    date = datetime.strptime('2022/1/8', '%Y/%m/%d')
    value = ['', '', '5151841261322240', '黎偉豪', '12345678', '門市']

    on_date = False
    if delivery:
        range_names = '送貨'
    else:
        range_names = str(date.year) + '年' + str(date.month).zfill(2) + "月"
    scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
             "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
    creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
    client = gspread.authorize(creds)
    sheetname = range_names
    sheet = client.open('交收list ').worksheet(sheetname)
    data = sheet.get_all_records()
    if not delivery:

        for index, row in enumerate(data):
            if row['日期'] != '':
                if datetime.strptime(row['日期'], '%Y/%m/%d') == date:
                    on_date = True
            if on_date:
                if row['客人名稱'].lower() == value[3].lower():
                    value[3] = ''
                    value[4] = ''
                    sheet.insert_row(value, index + 3)
                    notification_cell = 'J' + str(index + 3)
                    sheet.format(notification_cell, {
                        "backgroundColor": {
                            "red": 0.09375,
                            "green": 0.7109,
                            "blue": 0.5468
                        }})

                    red_cell = 'A' + str(index + 3)
                    sheet.format(red_cell, {
                        "backgroundColor": {
                            "red": 1,
                            "green": 0,
                            "blue": 0
                        }})
                    break
            if row['日期'] != '':
                if datetime.strptime(row['日期'], '%Y/%m/%d') == date + timedelta(days=1):
                    sheet.insert_row(value, index + 2)

                    notification_cell = 'J' + str(index + 2)
                    sheet.format(notification_cell, {
                        "backgroundColor": {
                            "red": 0.09375,
                            "green": 0.7109,
                            "blue": 0.5468
                        }})

                    red_cell = 'A' + str(index + 2)
                    sheet.format(red_cell, {
                        "backgroundColor": {
                            "red": 1,
                            "green": 0,
                            "blue": 0
                        }})



    else:
        sheet.insert_row(value, len(data) + 2)

        notification_cell = 'K' + str(len(data) + 2)
        sheet.format(notification_cell, {
            "backgroundColor": {
                "red": 0.09375,
                "green": 0.7109,
                "blue": 0.5468
            }})

        red_cell = 'A' + str(len(data) + 2)
        sheet.format(red_cell, {
            "backgroundColor": {
                "red": 1,
                "green": 0,
                "blue": 0
            }})

def auction_place():
    date = str(input('Date: (dd/mm/yyyy)'))
    wwautils.auction_placement(datetime.strptime(date , '%d/%m/%Y'))




        #sheet.insert_row(value, 1)
    # try:
    #     os.system("Xcopy \"C:\\Users\\world\\World Wide Aquarium Dropbox\\WWA SHKS\\PROGRAM\\AuctionSystem UI\\src\\Picture\" \""+temp_path+"\" /E/H/C/I/Y/Q > nul 2>&1")
    # except:
    #     print('bypass2')