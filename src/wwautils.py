# -*- coding: UTF-8 -*-

import getpass
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import TimeoutException
import pickle
import random
from bs4 import BeautifulSoup
import wwaformatcreate
import wwaglobal
import wwalog
import os
import docx
from distutils.dir_util import copy_tree
import shutil
import dropbox
import re
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from collections import Counter
from operator import itemgetter
from datetime import datetime
from datetime import timedelta
import xlrd
from pprint import pprint
#from Google import Create_Service
from selenium.webdriver.remote.webelement import WebElement
import wwatkmenu
#from gspread.models import Cell
from openpyxl import load_workbook
from openpyxl.styles import Border, Side
from dropbox import common
try:
	from gspread.models import Cell
except:
	pass
from ppadb.client import Client as AdbClient


def postAuction(postNumber, postNumber2):
	#postNumber = int(input("貼文數量 (環球水族HK - 拍賣區): "))
	#postNumber2 = int(input("貼文數量 (發燒友拍賣區): "))
	#testing
	wwaformatcreate.formatCreate(postNumber, postNumber2)

	# check if today's document is ready
	if not os.path.isfile(wwaglobal.auctionDocumentPath):
		# no document
		print("Today's auction document not found: " + wwaglobal.auctionDocumentPath + "\nexitting...")
		wwalog.log("Today's auction document not found: " + wwaglobal.auctionDocumentPath + "\nexitting...")
		return
	total = 0
	if postNumber2>0:
		auctionDocument2 = docx.Document(wwaglobal.dropboxPath + str(
			datetime.today().year) + "\\Auction Coral Retial\\發燒友docx\\" + wwaglobal.today + ".docx")
		for i in range(postNumber2):
			coralAll = auctionDocument2.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[
				0].strip().split(' / ')
			total += len(coralAll)
	auctionDocument = docx.Document(wwaglobal.auctionDocumentPath)

	for i in range(postNumber):
		print(f'{i=}')
		coralAll = auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[
			0].strip().split(' / ')
		total += len(coralAll)

	# path = wwaglobal.dropboxPath + str(
	# 	datetime.today().year) + '\\Auction Coral Retial\\' + wwaglobal.today + ' Auction Coral'
	while 1:
		photoList = os.listdir(wwaglobal.photoPath)
		pprint(photoList)
		if len(photoList) > 0:
			break
		else:
			print('Loading Photos')
			temp = input('wait')
			#os.system("Xcopy \"Picture\" \"" + (path + '\\Picture') + "\" /E/H/C/I/Y/Q > nul 2>&1")
			time.sleep(5)
	try:
		photoList.remove('desktop.ini')
	except:
		pass
	# del photoList[len(photoList)-1]
	# del photoList[0]
	time.sleep(5)
	if (total) != len(photoList):
		print(f'{total=} {len(photoList)=}')
		wwatkmenu.error_box('Photo number and coral number not match')
		return


	# 環球水族HK - 拍賣區: https://band.us/band/78427905/
	# 發燒友拍賣區（珊瑚、海水魚、器材用品）: https://band.us/band/75420559


	options = Options()
	options.headless = False

	# print("Copying Chrome profile.")
	# wwalog.log("Copying Chrome profile.")
	# os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")

	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
	driver.get("https://band.us/")

	"""
	#gmail_address = input("Enter your Gmail address:\t")
	#gmail_password = input("Enter your password:\t")
	f = open("account.txt", "r")
	temp = f.read().split("\n")
	f.close()
	gmail_address = temp[0]
	gmail_password = temp[1]

	# login band via Google
	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, chrome_options=options)

	url = "https://accounts.google.com/o/oauth2/v2/auth/identifier?response_type=code&client_id=297053617361-2jq02c7vlo14302ppvfvb3di2r9ca4bq.apps.googleusercontent.com&scope=email%20profile&state=ZQDGTYK3M2LNPUBCG7L4LICPUCQL5JVNCMA3QHAZFKDTJWESRW5ZIFDOUKUJXF7O3VL3QV7GGIU5A%3D%3D%3D&prompt=consent&redirect_uri=https%3A%2F%2Fauth.band.us%2Fexternal_account_login%3Ftype%3Dgoogle&flowName=GeneralOAuthFlow"
	driver.get(url)
	emailinput = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.NAME, "identifier")))
	emailinput.send_keys(gmail_address)
	time.sleep(random.randint(1, 10))
	nextbutton = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.ID, "identifierNext")))
	ActionChains(driver).move_to_element(nextbutton).click().perform()
	time.sleep(random.randint(1, 10))

	passwordinput = WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.NAME, "password")))
	passwordinput.send_keys(gmail_password)
	time.sleep(random.randint(1, 10))
	nextbutton = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.ID, "passwordNext")))
	ActionChains(driver).move_to_element(nextbutton).click().perform()
	time.sleep(random.randint(1, 10))

	#print(colored("[P] ","red")+colored(f"You logged in using {gmail_address} ... ", "blue"))


	# check if there is alert for login error
	try:
		WebDriverWait(driver, 3).until(EC.alert_is_present())

		alert = driver.switch_to.alert
		alert.accept()
		wwalog.log("alert accepted")
		loginWithGoogleBt = driver.find_element_by_partial_link_text("Log in with Google")
		loginWithGoogleBt.click()
	except TimeoutException:
		wwalog.log("no alert")
	"""

	# check if login band success
	while 1:
		try:
			bt = WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.CLASS_NAME, "bandCreate")))
			wwalog.log("login band success.")
			break
		except TimeoutException:
			pass

	#Save cookies
	#pickle.dump(driver.get_cookies(), open(cookie_file_name, "wb"))

	coralType = []
	startBid = []
	while 1:
		photoList = os.listdir(wwaglobal.photoPath)
		if len(photoList) > 0:
			break
		else:
			print('Loading Photos')
			time.sleep(1)

	print(photoList)
	# 發燒友拍賣區
	auctionDocPath = wwaglobal.dropboxPath +str(datetime.today().year)+"\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\AuctionDoc.docx"
	auctionDocMessage = ""
	if postNumber2 > 0:
		# 獲取最後一個post的編號 (發燒友拍賣區（珊瑚、海水魚、器材用品）)
		if wwaglobal.debug:
			driver.get("")
		else:
			driver.get("https://band.us/band/75420559/")

		firstPostCode2 = -1
		while firstPostCode2 == -1:
			try:
				soup = BeautifulSoup(driver.page_source, "lxml")
				a_tags = soup.find_all('a')
				link = ""
				for tag in a_tags:
					if tag.get('href').startswith('https://band.us/band/75420559/post/'):
						link = tag.get('href')
						break
				firstPostCode2 = int(link.split('post/')[1])
			except IndexError:
				time.sleep(0.5)

		auctionDocument2 = docx.Document(wwaglobal.dropboxPath + str(datetime.today().year)+"\\Auction Coral Retial\\發燒友docx\\" + wwaglobal.today + ".docx")
		#auctionDocument2 = docx.Document(wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\發燒友doc\\2021-09-30.docx")


		coralType = []
		startBid = []

		for i in range(postNumber2):
			coralType.append(auctionDocument2.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
			startBid.append(auctionDocument2.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())
			wwalog.log("[發燒友拍賣區 Post " + str(i + 1) + "]\n" + coralType[i] + " 起標價：" + startBid[i] + "\n拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")

			time.sleep(2)

			try:
				# click create post button
				driver.find_element_by_css_selector(
					"[class*='cPostWriteEventWrapper'][class*='_btnOpenWriteLayer']").click()
			except:
				time.sleep(0.5)
			time.sleep(2)

			# find write post input
			postTextInput = driver.find_element_by_css_selector("div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")

			# write post text
			postTextInput.send_keys(auctionDocument2.paragraphs[i].text)


			time.sleep(0.5)

			# find file input
			postFileInput = driver.find_element_by_css_selector("input[type='file'][id^=postPhotoInput_view]")

			# upload file
			#print(wwaglobal.photoPath)
			#print(i)
			postFileInput.send_keys(wwaglobal.photoPath + '\\' + photoList[i])
			wwalog.log("發燒友拍賣區: Uploaded file " + wwaglobal.photoPath + '\\' + photoList[i])
			clicked = False
			# find attach button
			while 1:
				try:
					bt = WebDriverWait(driver, 2).until(
						EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
					# click attach button
					#bt_temp = driver.find_element_by_xpath('//*[@id="wrap"]/div[2]/div[2]/section/div/footer/button[2]')
					time.sleep(1)
					#bt_temp.click()

					bt.send_keys(Keys.RETURN)
					break
				except TimeoutException:
					pass
			#temp_wait = input('press attach button')
			time.sleep(1)

			#find post button
			while 1:
				try:
					bt = WebDriverWait(driver, 2).until(
						EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton._btnSubmitPost.-confirm")))
					time.sleep(1)

					# click post button
					#bt.click()
					bt.send_keys(Keys.RETURN)
					wwalog.log("Finished 發燒友拍賣區 post " + str(i + 1) + ".")
					time.sleep(2)
					break
				except TimeoutException:
					pass
			#temp_wait = input('press post button')
			wwalog.log("\n")
	time.sleep(1)
	#獲取最後一個post的編號
	driver.get("https://band.us/band/78427905/")

	firstPostCode = -1
	while firstPostCode == -1:
		try:
			soup = BeautifulSoup(driver.page_source, "lxml")
			a_tags = soup.find_all('a')
			link = ""
			for tag in a_tags:
				if tag.get('href').startswith('https://band.us/band/78427905/post/'):
					link = tag.get('href')
					break
			firstPostCode = int(link.split('post/')[1])
		except IndexError:
			time.sleep(0.5)
	#firstPostCode =2212
	wwalog.log("firstPostCode: " + str(firstPostCode) + "\n")
	auctionDocMessage += '❗最新珊瑚拍賣❗\n'
	time.sleep(0.1)
	auctionDocument = docx.Document(wwaglobal.auctionDocumentPath)

	try:
		# create a file for sharing post
		f = open(
			wwaglobal.dropboxPath + str(datetime.today().year)+"\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\today_auction_links.txt",
			"w+", encoding="utf-8")
		for i in range(postNumber2):

			auctionDocMessage += coralType[i] + " 起標價：" + startBid[i] + "\n"
			auctionDocMessage += "拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n\n"
			time.sleep(0.1)

			f.write("https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")

		for i in range(postNumber):
			# each post info
			coralType.append(
				auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
			startBid.append(auctionDocument.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())
			auctionDocMessage += coralType[postNumber2 + i] + " 起標價：" + startBid[postNumber2 + i] + "\n"
			auctionDocMessage += "拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1) + "\n\n"
			time.sleep(0.1)

			f.write("https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1) + "\n")

		f.close()
	except:
		print('create auction share post unsuccessfully')
	f = open(
		wwaglobal.dropboxPath + str(datetime.today().year)+"\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\today_auction_links.txt",
		"w+", encoding="utf-8")
	try:
		for i in range(postNumber2):
			f.write("https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")
		for i in range(postNumber):
			f.write("https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1) + "\n")
	except:
		print('create link file unsuccessfully')
	f.close()
	auctionDoc = docx.Document()
	auctionDocLines = auctionDocMessage.split("\n")
	for i in range(len(auctionDocLines)):
		auctionDoc.add_paragraph(auctionDocLines[i])
	auctionDoc.save(auctionDocPath)

	temp_count = 0
	#print(photoList)
	for i in range(postNumber):
		#coralType.append(auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip())
		coralAll = auctionDocument.paragraphs[i].text.split('品種：')[1].split('CAA')[0].strip().split('\n')[0].strip().split(' / ')
		no_of_coral = len(coralAll)

		#startBid.append(auctionDocument.paragraphs[i].text.split('底價：')[1].split('\n')[0].strip())
		wwalog.log("[Post " + str(i + 1) + "]\n" + coralType[i + postNumber2] + " 起標價：" + startBid[i + postNumber2] + "\n拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode) + i + 1) + "\n")

		time.sleep(2)
		while 1:
			try:
				# click create post button
				driver.find_element_by_css_selector("[class*='cPostWriteEventWrapper'][class*='_btnOpenWriteLayer']").click()
				#time.sleep(2)
				#postTextInput = driver.find_element_by_css_selector(
					#"div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")
				break
			except:
				time.sleep(0.5)
		time.sleep(4)

		postTextInput = driver.find_element_by_css_selector(
			"div[class*='cke_editable'][class*='cke_editable_inline'][class*='_richEditor']")
		time.sleep(0.5)
		postTextInput.send_keys(auctionDocument.paragraphs[i].text)
		time.sleep(0.5)

		# find file input
		postFileInput = driver.find_element_by_css_selector("input[type='file'][id^=postPhotoInput_view]")


		# upload file
		#print(temp_count)
		temp_photo_list = ''
		for i in range(no_of_coral):
			temp_photo_list += wwaglobal.photoPath + '\\' + photoList[postNumber2 + temp_count]
			temp_count += 1
			if i != no_of_coral - 1:
				temp_photo_list += '\n'

		time.sleep(1)
		postFileInput.send_keys(temp_photo_list)

		# postFileInput.send_keys(wwaglobal.photoPath + '\\' + photoList[postNumber2 + temp_count])
		# temp_count += 1
		# if no_of_coral > 1:
		# 	for j in range(no_of_coral-1):
		# 		# addBtn = WebDriverWait(driver, 15).until(
		# 		# 	EC.visibility_of_element_located((By.XPATH, '//*[@id="wrap"]/div[2]/div[2]/section/div/div/div/div[1]/span')))
		# 		# moreFotoInput = driver.find_element_by_xpath(
		# 		# 	'//*[@id="wrap"]/div[2]/div[2]/section/div/div/div/div[1]/span/input')
		# 		while 1:
		# 			try:
		#
		# 				moreFotoInput = driver.find_element_by_xpath(
		# 				"(//input[@type='file' and @accept='image/*'])[2]")
		# 				break
		# 			except:
		# 				time.sleep(1)
		# 				#print('cannot attach more photo')
		# 		moreFotoInput.send_keys(wwaglobal.photoPath + '\\' + photoList[postNumber2 + temp_count])
		# 		time.sleep(5)
		#
		# 		temp_count += 1


		wwalog.log("Uploaded file " + wwaglobal.photoPath + '\\' + photoList[postNumber2 + i])
		clicked = False
		# find attach button
		while 1:
			try:
				bt = WebDriverWait(driver, 1).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton.-confirm._submitBtn")))
				# click attach button
				#bt.click()
				time.sleep(1)
				bt.send_keys(Keys.RETURN)
				#bt_temp = driver.find_element_by_xpath('//*[@id="wrap"]/div[2]/div[2]/section/div/footer/button[2]')

				# while clicked == False:
				# 	try:
				# 		bt.click()
				# 		clicked = True
				# 	except:
				# 		time.sleep(1)
				# 		print('try')
				#bt_temp.click()
				break
			except TimeoutException:
				pass
		#temp_wait = input('press attach button')
		time.sleep(2 * no_of_coral)

		#find post button
		while 1:
			try:
				bt = WebDriverWait(driver, 1).until(EC.element_to_be_clickable((By.CSS_SELECTOR, "button.uButton._btnSubmitPost.-confirm")))
				time.sleep(1)
				# click post button
				while 1:
					try:
						#bt.click()
						bt.send_keys(Keys.RETURN)
						time.sleep(1)
						driver.find_element_by_css_selector(
							"[class*='cPostWriteEventWrapper'][class*='_btnOpenWriteLayer']")
						break
					except KeyboardInterrupt:
						break
					except:
						print("Post button cannot be clicked. Trying again in 1 seconds... Or after clicking the button press Ctrl+C")
						#time.sleep(1)
				wwalog.log("Finished post " + str(i + 1) + ".")
				time.sleep(2)
				break
			except TimeoutException:
				pass
		time.sleep(5)
		#temp_wait = input('press post button')
		wwalog.log("\n")



	time.sleep(1)


	# open web whatsapp
	wwalog.log("Finished posting to band.\nOpening https://web.whatsapp.com/")
	print("Finished posting to band.\nOpening https://web.whatsapp.com/")
	driver.get("https://web.whatsapp.com/")

	# check if login whatsapp success
	while 1:
		try:
			WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
			wwalog.log("login whatsapp success.")
			break
		except TimeoutException:
			pass

	grpList = ['拍賣通知專區']
	#grpList = ['93258078'] #DEBUG
	counter = 0
	time.sleep(2)

	auctionDocPath = wwaglobal.dropboxPath + str(datetime.today().year)+"\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\AuctionDoc.docx"
	auctionDocMessage = ""

	for grpName in grpList:
		# find group
		wwalog.log("Finding group \"" + grpName + "\"")
		while 1:
			try:
				driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(grpName)
				break
			except:
				wwalog.log("Cannot access group name input box of web Whatsapp, trying again...")
				time.sleep(1)
		#driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys('🌊環球水族 WWA🌊會員通知區')
		time.sleep(1)
		driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(Keys.RETURN)
		time.sleep(2)

		# type message
		wwalog.log("Finding input box.")
		inputbox = driver.find_element_by_css_selector("div[role='textbox'][spellcheck='true']")
		#inputbox.send_keys('❗最新珊瑚拍賣❗')
		inputbox.send_keys(':bell')
		inputbox.send_keys(Keys.RETURN)
		inputbox.send_keys('最新拍賣')
		inputbox.send_keys(':bell')
		inputbox.send_keys(Keys.RETURN)

		inputbox.send_keys(':fire')
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys(Keys.ARROW_LEFT)
		time.sleep(0.5)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys(Keys.ARROW_RIGHT)
		time.sleep(0.5)
		inputbox.send_keys('標準完標時間每晚9點')
		inputbox.send_keys(':fire')
		inputbox.send_keys(Keys.RETURN)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys('多款珊瑚、魚種任你拍賣，最終價錢由你話事！!')
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys('快啲加入我地拍賣Band Page參加拍賣啦：')
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys('https://band.us/n/ada260kfmfib4')
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys('想第一時間知道最新拍賣內容？咁就要加入我地既拍賣whatsapp群組啦：')
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys('https://chat.whatsapp.com/EyGWGzUlNqIKiIzk5lXOxu')
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)

		wwalog.log("\n❗最新珊瑚拍賣❗")


		#auctionDocMessage += '❗最新珊瑚拍賣❗\n'
		time.sleep(0.1)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)

		# create a file for sharing post
		#f = open(wwaglobal.dropboxPath + "2021\\Auction Coral Retial\\" + wwaglobal.today + " Auction Coral\\today_auction_links.txt", "w+", encoding = "utf-8")
		for i in range(postNumber2):
			wwalog.log("Generating msg for post " + str(i + 1) + ".")
			time.sleep(0.1)
			inputbox.send_keys(coralType[i] + " 起標價：" + startBid[i])
			wwalog.log(coralType[i] + " 起標價：" + startBid[i])
			#auctionDocMessage += coralType[i] + " 起標價：" + startBid[i] + "\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys("拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1))
			wwalog.log("拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1))
			#auctionDocMessage += "拍賣連結： https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			time.sleep(0.1)

			#f.write("https://band.us/band/75420559/post/" + str(int(firstPostCode2) + i + 1) + "\n")

		for i in range(postNumber):
			# each post info
			wwalog.log("Generating msg for post " + str(i + postNumber2 + 1) + ".")
			time.sleep(0.1)
			inputbox.send_keys(coralType[postNumber2+i] + " 起標價：" + startBid[postNumber2+i])
			wwalog.log(coralType[postNumber2+i] + " 起標價：" + startBid[postNumber2+i])
			auctionDocMessage += coralType[postNumber2+i] + " 起標價：" + startBid[postNumber2+i] + "\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys("拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1))
			wwalog.log("拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1))
			auctionDocMessage += "拍賣連結： https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1) + "\n\n"
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
			time.sleep(0.1)

			#f.write("https://band.us/band/78427905/post/" + str(int(firstPostCode)+i+1) + "\n")

		#f.close()

		# auctionDoc = docx.Document()
		# auctionDocLines = auctionDocMessage.split("\n")
		# for i in range(len(auctionDocLines)):
		# 	auctionDoc.add_paragraph(auctionDocLines[i])
		# auctionDoc.save(auctionDocPath)


		#temp = input('If finished editing (dont send it), press <enter>')
		# send message
		driver.find_element_by_css_selector("span[data-testid='send'][data-icon='send']").find_element_by_xpath("..").send_keys(Keys.RETURN)
		wwalog.log("Sent message.")
		time.sleep(0.5)

		try:
			time.sleep(1)
			
			while 1:
				try:
					WebDriverWait(driver, 10).until(
						EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
					wwalog.log("login whatsapp success.")
					break
				except TimeoutException:
					time.sleep(0.5)
			# print(len(photoList))

			uploaded = 0
			for i in range(len(photoList) // 30 + 1):
				temp_photo_list=''

				if (len(photoList) - uploaded) // 30 >= 1:
					photo_limit = 30
				else:
					photo_limit = len(photoList) % 30
				# print(photo_limit)
				for j in range(photo_limit):
					temp_photo_list += wwaglobal.photoPath + '\\' + photoList[i * 30 + j]

					temp_photo_list += '\n'
					uploaded += 1

				pprint(temp_photo_list[:len(temp_photo_list) - 1])
				WebDriverWait(driver, 20).until(
					EC.presence_of_element_located((By.CSS_SELECTOR, "span[data-testid='clip'][data-icon='clip']")))
				counter_attach = 0
				while 1:
					try:
						driver.find_element_by_css_selector(
							"span[data-testid='clip'][data-icon='clip']").find_element_by_xpath(
							"..").send_keys(Keys.RETURN)
						break
					except:
						print('Try Clicking Button')
						counter_attach += 1
						time.sleep(1)
					if counter_attach == 5:
						break
				time.sleep(0.5)

				driver.find_element_by_xpath(
					"//input[@accept='image/*,video/mp4,video/3gpp,video/quicktime']").find_element_by_xpath(
					"../input").send_keys(temp_photo_list[:len(temp_photo_list) - 1])

				while 1:
					try:
						WebDriverWait(driver, 20).until(
							EC.presence_of_element_located((By.XPATH, "//span[@data-testid='send']")))
						break
					except TimeoutException:
						time.sleep(1)
				time.sleep(1)
				driver.find_element_by_xpath("//div[@role='button' and @class='_165_h _2HL9j']").send_keys(Keys.RETURN)
				time.sleep(3)
			wwalog.log("Sent photos.")
		except:
			print('Error in sending images, please do it on hand.')
		time.sleep(5)
	# time.sleep(0.5)
	# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys('🌊環球水族🌊到貨及優惠通知專區')
	# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(Keys.RETURN)
	# time.sleep(0.5)
	# driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys("❗最新珊瑚拍賣❗")
	# driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys(Keys.RETURN)




	time.sleep(5)
	driver.quit()

	print("Finished.")







def endAuction(date):
	#date = input("Date (yyyy-mm-dd): ")
	auctionFolder = wwaglobal.dropboxPath + str(datetime.strptime(date,'%Y-%m-%d').year)+"\\Auction Coral Retial\\" + date + " Auction Coral\\"

	tempFolder = wwaglobal.personalDropboxPath + "Temporary Store for Posting"
	if not os.path.isdir(auctionFolder):
		wwatkmenu.error_box("Path does not exist: " + auctionFolder)
		print("Path does not exist: " + auctionFolder)
		wwalog.log("Path does not exist: " + auctionFolder)
		return

	photoFolder = auctionFolder + "Picture"
	if not os.path.isdir(photoFolder):
		wwatkmenu.error_box("Path does not exist: " + photoFolder)
		print("Path does not exist: " + photoFolder)
		wwalog.log("Path does not exist: " + photoFolder)
		return

	commentDocumentPath = wwaglobal.dropboxPath + str(datetime.strptime(date,'%Y-%m-%d').year)+"\\Auction Coral Retial\\" + date + " Auction Coral\\comment.docx"
	if not os.path.isfile(commentDocumentPath):
		wwatkmenu.error_box("Comment document does not exist: " + commentDocumentPath)
		print("Comment document does not exist: " + commentDocumentPath)
		wwalog.log("Comment document does not exist: " + commentDocumentPath)
		return

	wwalog.log("commentDocumentPath: " + commentDocumentPath)
	commentDocument = docx.Document(commentDocumentPath)

	"""
	# remove temp folder from dropbox root folder
	if os.path.isdir(tempFolder):
		if not wwaglobal.debug:
			shutil.rmtree(tempFolder)
			wwalog.log("Removed folder: " + tempFolder)

			# create it again
			os.mkdir(tempFolder)
			wwalog.log("Created folder: " + tempFolder)


	# move photos from date's picture folder to root folder
	wwalog.log("Trying to copy photo folder.")
	print("Trying to copy photo folder. Please wait about several minutes.")
	if not wwaglobal.debug:
		os.system("Xcopy \"" + photoFolder + "\" \"" + tempFolder + "\" /E/H/C/I/Y/Q > nul 2>&1")
	"""
	# get the dropbox shared link of those pictures

	wwalog.log("Connected Dropbox api.")
	dropboxApiPhotoPath = "/Temporary Store for Posting/"
	#photoList = os.listdir(tempFolder)
	photoLink = []


	options = Options()
	options.headless = False

	print("Copying Chrome profile.")
	wwalog.log("Copying Chrome profile.")
	os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")
	#shutil.copytree(wwaglobal.chromeProfilePath, wwaglobal.chromeProfilePath1, dirs_exist_ok=True)
	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)


	f = open(auctionFolder + "today_auction_links.txt", "r", encoding = "utf-8")
	temp = f.read().split("\n")
	f.close()
	fever_guy = False
	for i in range(len(temp)):
		if not "http" in temp[i]:
			continue
		if 'https://band.us/band/75420559/post/' in temp[i]:
			fever_guy = True
			continue
		driver.get(temp[i])

		while 1:
			try:
				elem = driver.find_element_by_css_selector("[class*='collageItem']").find_element_by_xpath("a/img").get_attribute('src')
				break
			except:
				time.sleep(1)

		photoLink.append(elem)
		#wwalog.log(photoLink[i])



	"""
	for i in range(len(photoList)):
		photoLink.append(dbx.sharing_get_shared_links(dropboxApiPhotoPath + photoList[i]).links)
		if len(photoLink[i]) == 0:
			# generate shared link
			while 1:
				try:
					dbx.sharing_create_shared_link_with_settings(dropboxApiPhotoPath + photoList[i])
					break
				except:
					wwalog.log("Path in dropbox \"" + dropboxApiPhotoPath + photoList[i] + "\" not found. Try again in 5 seconds...")
					print("Path in dropbox \"" + dropboxApiPhotoPath + photoList[i] + "\" not found. Try again in 5 seconds...")
					time.sleep(5)

		photoLink[i] = str(dbx.sharing_get_shared_links(dropboxApiPhotoPath + photoList[i]).links[0].url)
		wwalog.log("Link of photo " + str(i + 1) + ": " + photoLink[i])
	"""

	bid_price = []
	for i in range(len(commentDocument.paragraphs)):

		try:
			bid_price.append(int(commentDocument.paragraphs[i].text.split('$')[1].split('中標')[0].strip()))
		except:
			# no one bid
			bid_price.append(0)
		wwalog.log("Bid " + str(i + 1) + ": $" + str(bid_price[i]))

	f = open(auctionFolder + date + " Auction Coral CSV.csv", "r", encoding = "utf-8")
	lines = f.read().split("\n")
	f.close()
	line_to_del = []

	for i in range(1, len(lines)-1):
		if lines[i] != "":
			#print(i)
			if str(bid_price[i-1]) == '0':
				# del lines[i]
				line_to_del.append(i)

			lines[i] = lines[i].replace("[auction_price]", str(bid_price[i-1])).replace("[photo_link]", photoLink[i-1])
	temp_counter = 0
	for line_del in line_to_del:
		del lines[line_del-temp_counter]
		temp_counter += 1

	f = open(auctionFolder + date + " Auction Coral CSV.csv", "w+", encoding = "utf-8")
	for i in range(len(lines)):
		if lines[i] != "":
			f.write(lines[i] + "\n")
	f.close()
	#wait_temp = input('wait')






	driver.get("https://www.boutir.com/user_cms/upload-products")


	bt = None
	# find upload input
	while 1:
		try:
			bt = WebDriverWait(driver, 1).until(EC.presence_of_element_located((By.CSS_SELECTOR, "input.csvCapture")))
			break
		except TimeoutException:
			pass

	time.sleep(2)

	bt.send_keys(auctionFolder + date + " Auction Coral CSV.csv")

	time.sleep(5)
	#//button[contains(text(), '上載')]
	#if not wwaglobal.debug:
	#	WebDriverWait(driver, 1).until(EC.visibility_of_element_located((By.CSS_SELECTOR, "button.btn.btn-primary"))).click()
	try:
		WebDriverWait(driver, 2).until(
			EC.visibility_of_element_located((By.XPATH, "//button[contains(text(), '上載')]"))).click()
	except:
		time.sleep(1)
	print("Uploading CSV...")
	wwalog.log("Uploading CSV...")
	end_input = False
	print("When upload finished, press <Ctrl+C> to continue.")
	while not end_input:
		try:
			if driver.find_element_by_xpath("//*[@class='sr-only']").text == '100%':
				end_input = True
				break
		except KeyboardInterrupt:
			end_input = True
			break
		except:
			time.sleep(1)
	wwalog.log("User pressed <enter> to continue.")
	print("Gonna wait for 2 minutes to let the online store fully completed")
	time.sleep(120)


	f = open(auctionFolder + date + " Auction Coral CSV.csv", "r", encoding = "utf-8")
	lines = f.read().split("\n")
	f.close()

	itemLinks = []
	if wwaglobal.debug:
		for i in range(1, len(lines)):
			if lines[i] == "":
				continue
			print("link" + str(i))
			itemLinks.append("link" + str(i))
	seller_id = "5690381968343040"
	# find itemLinks
	if not wwaglobal.debug:
		for i in range(1, len(lines)):
			if lines[i] == "":
				continue
			query = lines[i].split(",")[2].strip()
			url = "https://wwa.boutir.com/?q=" + query

			while 1:
				try:
					driver.get(url)
					WebDriverWait(driver, 20).until(
						EC.presence_of_element_located((By.CLASS_NAME, 'search-result-title'))
					)
					break
				except TimeoutException:
					time.sleep(0.5)
			driver.find_element_by_class_name('clickable').click()
			while 1:
				try:
					WebDriverWait(driver, 10).until(
						EC.presence_of_element_located((By.CLASS_NAME, 'addthis_button_more'))
					)
					break
				except:
					time.sleep(0.5)
			time.sleep(1)
			item_id = driver.current_url.split("?q=")[0].split("/")[-1]

			driver.get("https://wwa.boutir.com/apis/get_cache_item_details?is_shorten_url=0&item_id=" + item_id + "&item_preview=&seller_id=" + seller_id)
			time.sleep(1)
			itemLinks.append(driver.page_source.split("\"shorten_item_url\": \"")[1].split("\"")[0])
			print(itemLinks[i-1])
			wwalog.log(itemLinks[i-1])




	linkDocument = docx.Document(auctionFolder + "AuctionDoc.docx")

	links = []

	for i in range(2, len(linkDocument.paragraphs), 3):
		links.append(linkDocument.paragraphs[i].text.split('拍賣連結：')[1].strip())
	for i in range(len(links)-1,-1,-1):
		if 'https://band.us/band/75420559/post/' in links[i]:
			del links[i]


	nameList = []
	f = open(wwaglobal.src_path + "nameList.txt", "r", encoding = "utf-8")
	nameList = f.read().split("\n")
	f.close()
	counter = 0
	for i in range(len(links)):
		if bid_price[i] == 0:
			continue
		wwalog.log("Opening: " + links[i])
		driver.get(links[i])
		time.sleep(5)

		commentText = []

		soup = BeautifulSoup(driver.page_source, "lxml")
		for el in soup.find_all('p', attrs={'class': 'txt _commentContent'}):
			commentText.append(el.get_text())

		names = []
		for j in range(len(commentText)):
			try:
				names.append(WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="content"]/div/section/div/div/div[5]/div/div/div[' + str(j+1) + ']/div[1]/div[1]/div[2]/button/strong'))).text)
			except:
				continue

		print(names)

		index = 0

		if nameList[i] == "No one":
			# no one bid
			continue

		for j in range(len(names) - 1, -1, -1):
			if names[j] == nameList[i]:
				index = j
				break
		time.sleep(0.5)
		#print("i = " + str(i) + "; index = " + str(index))
		try:
			driver.execute_script('window.scrollTo(0, document.body.scrollHeight)')
		except:
			pass
		time.sleep(0.5)
		while 1:
			try:
				bt = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, f"(//*[contains(text(), '{bid_price[i]}') and contains(@class, 'txt _commentContent')])[1]//parent::div//child::button[@class='reply _replyBtn']")))
				break
			except:
				print('Reply button unable to find')
				temp = input('try again (y/n)?')
				if temp == 'n':
					break
		# click reply buttons
		time.sleep(0.5)

		while 1:
			try:

				bt.send_keys(Keys.RETURN)
				break
			except:
				time.sleep(1)
		time.sleep(1)
		#//*[@id="content"]/div/section/div/div/div[5]/div/div/div[1]/div[1]/div[1]/div[3]/div[3]/button[2]
		#//*[@id="content"]/div/section/div/div/div[5]/div/div/div[2]/div[1]/div[1]/div[3]/div[3]/button[2]
		inputBox = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, 'textarea.commentWrite._use_keyup_event._messageTextArea')))
		JS_ADD_TEXT_TO_INPUT = """
			var elm = arguments[0], txt = arguments[1];
			elm.value += txt;
			elm.dispatchEvent(new Event('change'));
			"""

		elem = driver.find_element_by_css_selector('textarea.commentWrite._use_keyup_event._messageTextArea')
		text = commentDocument.paragraphs[i].text.replace("拍賣品連結：", "拍賣品連結：" + itemLinks[counter])
		counter+=1
		driver.execute_script(JS_ADD_TEXT_TO_INPUT, elem, text)

		# send
		if not wwaglobal.debug:
			inputBox.send_keys(Keys.CONTROL + Keys.RETURN)



		time.sleep(5)
	try:
		if datetime.strptime(date,'%Y-%m-%d') > datetime.strptime('2022-01-31', '%Y-%m-%d'):
			auction_placement(datetime.strptime(date,'%Y-%m-%d'))
	except:
		print('Auction Placement Error')
	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	sheetname = str(datetime.strptime(date, '%Y-%m-%d').year)+'年'+str(datetime.strptime(date, '%Y-%m-%d').month).zfill(2) + '月 auction'
	sheet = client.open('Auction List').worksheet(sheetname)
	data = sheet.get_all_records()
	date_got = False
	cell_value = []
	for i in range(len(data) - 1, 1, -1):

		if data[i]['Upload Date'] == datetime.strftime(datetime.strptime(date, '%Y-%m-%d'), '%Y/%m/%d'):
			# print(len(data)-1-i)
			if not date_got:
				temp = i+2
				date_got = True
			last_temp = i+2
			try:
				if data[i]['中標者名稱'] == 'n':

					cell_value.append('n')


				else:

					cell_value.append('y')

			except:
				pass
		if data[i]['Upload Date'] is None:
			break
			#sheet.update_cell(i+2,2, 'f')
	#print(f'{temp=}{last_temp=}')
	# print(cell_value)
	cell_value = cell_value[::-1]
	#print(cell_value)
	cell_list = sheet.range(f'L{temp}:O{last_temp}')

	for i, val in enumerate(cell_value):
		if last_temp + i > temp:
			break
		for j in range(12, 16):
			cell_list.append(Cell(row=last_temp + i, col=j, value=val))
	
	#print(cell_list)
	#wait_temp = input('wait')
	# cell_list[i].value = val
	sheet.update_cells(cell_list)
	cell_list = sheet.range(f'B{last_temp}:B{temp}')
	for val in cell_list:
		val.value = 'f'
	sheet.update_cells(cell_list)
	driver.quit()


def checkOrderList():
	options = Options()
	options.headless = False

	# print("Copying Chrome profile.")
	# wwalog.log("Copying Chrome profile.")
	# os.system("Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")
	#
	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)
	options.add_argument('--disable-browser-side-navigation')
	#
	options.add_experimental_option('excludeSwitches', ['enable-logging'])
	# driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)

	#print('pass1 794')
	finished = 0
	while finished == 0:
		try:
			driver.get("https://www.boutir.com/user_cms/orders/")
			finished = 1
		except:
			time.sleep(10)
			print('try again')
	#driver.get("https://www.boutir.com/user_cms/orders/")
	#print('pass1 796')
	elemList = []
	dateList = []
	months = set()
	sheetname = set()
	googleSheetOrders = []

	while len(elemList) == 0 or len(dateList) == 0:
		time.sleep(5)
		#print('pass1 803')
		elemList = driver.find_elements_by_css_selector("[href*='/user_cms/edit-order/']")
		dateList = driver.find_elements_by_css_selector("[class*='purchase_date']")
		getdateList = driver.find_elements_by_css_selector("[class*='item_shipping_desc']")

	#print(elemList[1].find_element_by_xpath("../..").find_element_by_css_selector("[class='delivery_method']").find_element_by_xpath("div[1]/span[2]").text)
	#print('pass2 809')
	"""
	for i in range(len(elemList)):
		elemList[i] = elemList[i].find_element_by_xpath("../..")

	print(elemList[0].find_element_by_xpath("td[3]/a").text)
	"""
	#temp = input('wait')
	i = 0

	while i < len(dateList):
		if "日期" in dateList[i].text:
			dateList.pop(i)
		else:
			#print(dateList[i].text)
			try:
				months.add(int(getdateList[i].text.split("年")[1].split('月')[0]))
				if str(getdateList[i].text.split('年')[0].split('\n')[0])+'年'+str(getdateList[i].text.split("年")[1].split('月')[0]).zfill(2)+'月' == '2021年01月':
					i += 1
					continue
				sheetname.add(str(getdateList[i].text.split('年')[0].split('\n')[0])+'年'+str(getdateList[i].text.split("年")[1].split('月')[0]).zfill(2)+'月')
			except:
				pass
		i += 1

	#print(months) # {9, 10}
	#print(len(elemList))
	#print(len(dateList))
	#print(dateList[0].text)

	# dict of boolean
	wrote = {}
	temp = ""
	for i in range(len(elemList)):
		temp = str(elemList[i].text).strip(" \n")
		#print('temp: ',temp)
		wrote[temp] = False
		if "已發貨" in elemList[i].find_element_by_xpath("../..").find_element_by_css_selector("[class='status']").find_element_by_xpath("div[1]/span").text:
			wrote[temp] = True
		elif "無效訂單" in elemList[i].find_element_by_xpath("../..").find_element_by_css_selector("[class='status']").find_element_by_xpath("div[1]/span").text:
			wrote[temp] = True
		elif "等待收款確認" in elemList[i].find_element_by_xpath("../..").find_element_by_css_selector("[class='status']").find_element_by_xpath("div[1]/span").text:
			wrote[temp] = True
		wwalog.log("[CMS " + str(i+1) + "]: " + temp)
	#print('pass3 848')
	# If modifying these scopes, delete the file token.json.
	SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

	# The ID and range of a spreadsheet.
	SPREADSHEET_ID = '1VDfkinOq_SkezxMwmF4Fhi1pe0qTtYEbCxp1CQxtxbY'
	range_names = ["送貨"]
	for sheetnames in sheetname:
		range_names.append(sheetnames)
	print(range_names)
	#range_names.append(str(month+1)+'月')

	wwalog.log("Range names: " + ", ".join(range_names))

	"""Shows basic usage of the Sheets API.
	Prints values from a sample spreadsheet.
	"""
	creds = None
	# The file token.json stores the user's access and refresh tokens, and is
	# created automatically when the authorization flow completes for the first
	# time.
	if os.path.exists(wwaglobal.src_path + 'token.json'):
		creds = Credentials.from_authorized_user_file(wwaglobal.src_path + 'token.json', SCOPES)
	# If there are no (valid) credentials available, let the user log in.
	if not creds or not creds.valid:
		if creds and creds.expired and creds.refresh_token:
			creds.refresh(Request())
		else:
			flow = InstalledAppFlow.from_client_secrets_file(wwaglobal.src_path + 'credentials.json', SCOPES)
			creds = flow.run_local_server(port=0)
		# Save the credentials for the next run
		with open(wwaglobal.src_path + 'token.json', 'w') as token:
			token.write(creds.to_json())

	service = build('sheets', 'v4', credentials=creds)

	# Call the Sheets API
	sheet = service.spreadsheets()

	for i in range(len(range_names)):
		result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=range_names[i], valueRenderOption="FORMULA").execute()
		values = result.get('values', [])
		#print(values)

		if not values:
			print("Google sheet " + range_names[i] + ": No data found.")
			wwalog.log("Google sheet " + range_names[i] + ": No data found.")
		else:
			for row in values:
				# column A = row[0], B = row[1], C = row[2]...
				try:
					# format: 16-digit
					if len(re.findall("\d{16}", str(row[2]))) >= 1:
						googleSheetOrders.append(re.findall("\d{16}", str(row[2]))[0])
						wwalog.log("Google sheet: " + re.findall("\d{16}", str(row[2]))[0])
					else:
						wwalog.log("Not 16-digit: " + row[2])
				except:
					wwalog.log("Row: " + str(row))

	for key in wrote.keys():
		if key in googleSheetOrders:
			wrote[key] = True

	valuesToAdd = []
	content=[]
	phone_no = []
	temp = ""
	counter = 0
	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	sheetname = str(datetime.today().year)+'年'+str(datetime.today().month).zfill(2) + '月 auction'
	sheet = client.open('Auction List').worksheet(sheetname)
	# sheet = client.open('Auction List').worksheet('11月 auction')
	data = sheet.get_all_records()
	if datetime.today().month == 1:
		sheet2 = client.open('Auction List').worksheet(
			str(datetime.today().year-1) + '年12月 auction')
	else:
		sheet2 = client.open('Auction List').worksheet(str(datetime.today().year)+'年'+str(datetime.today().month-1) + '月 auction')
	data2 = sheet2.get_all_records()
	# pprint(data[251])

	for key in wrote.keys():
		if not wrote[key]:
			counter += 1
			temp += str(counter) + ": " + key + "\n"

			driver.get("https://www.boutir.com/user_cms/edit-order/" + key)
			tempRow = ["", "", key]
			table_content = {'訂單編號': key,
							 '客人名稱': '',
							 '聯絡電話': '',
							 '交收方法': '',
							 '付款方式': '',
							 '商品': [],
							 '總數': ''
							 }
			table = None
			while 1:
				try:
					table = driver.find_element_by_css_selector("table[class*='tool-tip-table']")
					break
				except:
					time.sleep(1)

			# check if need to add
			if "已發貨" in driver.find_element_by_css_selector("[class*='status-label']").text:
				continue
			elif "無效訂單" in driver.find_element_by_css_selector("[class*='status-label']").text:
				continue
			elif "等待收款確認" in driver.find_element_by_css_selector("[class*='status-label']").text:
				continue
			elif "等待付款" in driver.find_element_by_css_selector("[class*='status-label']").text:
				continue


			# name
			table_content['客人名稱'] = table.find_element_by_xpath("tr[1]/td[2]").find_element_by_css_selector(
				"[class*='bold-class']").text
			tempRow.append(table.find_element_by_xpath("tr[1]/td[2]").find_element_by_css_selector("[class*='bold-class']").text)

			# phone
			if table.find_element_by_xpath("tr[2]/td[2]").find_element_by_css_selector("[class*='bold-class']").text[4] == "1":
				continue
			tempRow.append(table.find_element_by_xpath("tr[2]/td[2]").find_element_by_css_selector("[class*='bold-class']").text[4:])
			table_content['聯絡電話'] = table.find_element_by_xpath("tr[2]/td[2]").find_element_by_css_selector(
				"[class*='bold-class']").text[
									4:]

			# location
			table_content['交收方法'] = table.find_element_by_xpath("tr[3]/td[2]").find_element_by_css_selector(
				"[class*='bold-class']").text

			if "門市" in table.find_element_by_xpath("tr[3]/td[2]").find_element_by_css_selector("[class*='bold-class']").text:
				tempRow.append("門市")
			elif "魚街" in table.find_element_by_xpath("tr[3]/td[2]").find_element_by_css_selector("[class*='bold-class']").text:
				tempRow.append("魚街")
			elif "送貨" in table.find_element_by_xpath("tr[3]/td[2]").find_element_by_css_selector("[class*='bold-class']").text:
				tempRow.append("送貨")
			else:
				tempRow.append("")

			# items
			tempElements = driver.find_elements_by_css_selector("[class*='item-title-cell']")
			tempItems = []
			for i in range(len(tempElements)):
				tempElements[i] = tempElements[i].find_element_by_xpath("span")
				table_content['商品'].append(tempElements[i].text)
				combination_count = tempElements[i].text.count('CAA')
				if combination_count == 0:
					tempItems.append(tempElements[i].text)
				for j in range(1,combination_count+1):

					if "CAA" in tempElements[i].text:
						tempItems.append(tempElements[i].text.split("CAA")[j].strip())
					# elif "caa" in tempElements[i].text:
					# 	tempItems.append(tempElements[i].text.split("caa")[j])
					else:
						tempItems.append(tempElements[i].text)
			for items in tempItems:
				cell_row = searchItemsGoogle(data, items)
				if cell_row == -1:

					cell_row2 = searchItemsGoogle(data2, items)
				try:
					if cell_row > 1:
						sheet.update_cell(cell_row, 18, 'y')
					if cell_row2 > 1:
						sheet2.update_cell(cell_row2, 18, 'y')
				except:

					print('Cannot append to Auction List')



			# pprint(data[251])
			try:
				for items in tempItems:
					cell = sheet.findall(items, in_column=5)
					sheet.update_cell(cell[len(cell) - 1].row, 18, 'y')
			except:
				try:
					for items in tempItems:
						cell = sheet2.findall(items, in_column=5)
						sheet2.update_cell(cell[len(cell)-1].row, 18, 'y')
				except:
					print('Could not sychronize with Auction List')

			#print(tempItems)
			if wwaglobal.isInt(tempItems[0]):
				# add a '#' before it
				tempRow.append("#" + " ".join(tempItems))
			else:
				#print('pass 1017')
				tempRow.append(" ".join(tempItems))

			# money
			tempRow.append(driver.find_element_by_css_selector("[class*='dollars']").find_element_by_css_selector("[class*='bold-class']").text[4:])
			table_content['總數'] = driver.find_element_by_css_selector(
				"[class*='dollars']").find_element_by_css_selector(
				"[class*='bold-class']").text[4:].split('.')[0]
			# payment method
			if "payme" in driver.find_element_by_css_selector("[class*='payment-method']").text.lower():
				tempRow.append("payme")
			elif "銀行轉帳" in driver.find_element_by_css_selector("[class*='payment-method']").text:
				tempRow.append("bank")
			else:
				tempRow.append("")
			table_content['付款方式'] = driver.find_element_by_css_selector("[class*='payment-method']").text
			#date
			if "本地取貨" in driver.find_element_by_xpath('//*[@id="wrapper"]/div[3]/div[1]/div[1]/div[2]/div/div[1]/table/tr[3]/td[2]/div/span').text:
				getdate = driver.find_element_by_xpath('//*[@id="wrapper"]/div[3]/div[1]/div[1]/div[2]/div/div[1]/table/tr[3]/td[2]/div/span').text.split('本地取貨 - ')[1].split(' ')[0]
				print(f'{getdate=}')
				if getdate == '2021年1月8日':
					datecollected = datetime.strptime('2022年1月8日', '%Y年%m月%d日')
				else:
					datecollected = datetime.strptime(getdate, '%Y年%m月%d日')
				print(f'{datecollected=}')
				if datecollected.month == 1 & datecollected.year == 2021:
					datecollected.replace(year=2022)
				#tempRow[1] = datetime.strftime(datecollected, '%Y/%m/%d')
			#(//*[@class="details"])[4]//child::span[@class='bold-class']
			elif "本地送貨" in driver.find_element_by_xpath('//*[@id="wrapper"]/div[3]/div[1]/div[1]/div[2]/div/div[1]/table/tr[3]/td[2]/div/span').text:
				datecollected = datetime.today()
				tempRow[1] = datetime.strftime(datecollected, '%Y/%m/%d')
				tempRow.append(driver.find_element_by_xpath("(//*[@class='details'])[4]//child::span[@class='bold-class']").text)

			else:
				datecollected = datetime.today()
				tempRow[1] = datetime.strftime(datecollected, '%Y/%m/%d')
				tempRow.append(driver.find_element_by_xpath(
					'//*[@id="wrapper"]/div[3]/div[1]/div[1]/div[2]/div/div[1]/table/tr[3]/td[2]/div/span').text)

			#print(datecollected.strftime('%Y/%m/%d'))
			listtoinput = tempRow
			pprint(f'{listtoinput=}')

			phone_no.append([table_content['聯絡電話'], table_content['客人名稱']])
			content.append(table_content)
			#print(tempRow)
			if "本地取貨" in driver.find_element_by_xpath(
					'//*[@id="wrapper"]/div[3]/div[1]/div[1]/div[2]/div/div[1]/table/tr[3]/td[2]/div/span').text:
				print('add')
				addIntoMeetingList(datecollected, listtoinput, False)
			elif "本地送貨" in driver.find_element_by_xpath(
					'//*[@id="wrapper"]/div[3]/div[1]/div[1]/div[2]/div/div[1]/table/tr[3]/td[2]/div/span').text:
				addIntoMeetingList(datecollected, listtoinput, True)
			#tempRow.append(driver.find_element_by_xpath('//*[@id="wrapper"]/div[3]/div[1]/div[1]/div[2]/div/div[1]/table/tr[3]/td[2]/div/span').text)
			#wait = input('wait')
			valuesToAdd.append(tempRow)


	driver.get('https://web.whatsapp.com')
	WhatsappOrder(content, phone_no, driver)


	valuesToAdd = sorted(valuesToAdd, key=itemgetter(4))

	# add to google sheet
	# body = {
	# 	'values': valuesToAdd
	# }
	# value_input_option = 'USER_ENTERED'
	#
	# result = service.spreadsheets().values().append(
	# 	spreadsheetId=SPREADSHEET_ID, range=range_names[0],
	# 	valueInputOption=value_input_option, body=body).execute()
	#
	# if result.get('updates').get('updatedCells') == None:
	# 	print("Google sheet: 0 cells appended.")
	# 	wwalog.log("Google sheet: 0 cells appended.")
	# 	print("Do not have new order.")
	# 	wwalog.log("Do not have new order.")
	# else:
	# 	# updated
	# 	print("Google sheet: " + str(result.get('updates').get('updatedCells')) + ' cells appended.')
	# 	wwalog.log("Google sheet: " + str(result.get('updates').get('updatedCells')) + ' cells appended.')
	# 	print("Have new order, please check the google sheet \"" + range_names[0] + "\" tab.")
	# 	wwalog.log("Have new order, please check the google sheet \"" + range_names[0] + "\" tab.")




	# wwalog.log("\nOrders that did not put into Google sheet (Total " + str(counter) + "):")
	# wwalog.log(temp)
	# print("\nOrders that did not put into Google sheet (Total " + str(counter) + "):")
	# print(temp)

	# duplicatedItems = []
	# Counter(googleSheetOrders)
	# for key in Counter(googleSheetOrders).keys():
	# 	if Counter(googleSheetOrders)[key]>1:
	# 		duplicatedItems.append(key)
	#
	# if len(duplicatedItems) > 0:
	# 	print("Duplicated items in Google sheet: " + ", ".join(duplicatedItems) + ".\n")
	# 	wwalog.log("Duplicated items in Google sheet: " + ", ".join(duplicatedItems) + ".\n")

	driver.quit()

def searchItemsGoogle(data, target):
	for i in range(len(data)-1,1,-1):
		if str(data[i]['Serial Number']).find(target) != -1:
			return i+2
	return -1



def addIntoMeetingList(date, value, delivery):
	# range_id = {
	# 	'送貨': '766588160',
	# 	'2021年10月': '943347126',
	# 	'2021年11月': '813817967',
	# 	'2021年12月': '1982121589',
	# 	'2022年01月': '461242197'
	# }
	# SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
	# 
	# # The ID and range of a spreadsheet.
	# SPREADSHEET_ID = '1VDfkinOq_SkezxMwmF4Fhi1pe0qTtYEbCxp1CQxtxbY'
	# if delivery:
	# 	range_names = '送貨'
	# else:
	# 	range_names =  str(date.year)+'年'+str(date.month).zfill(2)+"月"
	# if not delivery:
	# 	creds = None
	# 	# The file token.json stores the user's access and refresh tokens, and is
	# 	# created automatically when the authorization flow completes for the first
	# 	# time.
	# 	if os.path.exists(wwaglobal.src_path + 'token.json'):
	# 		creds = Credentials.from_authorized_user_file(wwaglobal.src_path + 'token.json', SCOPES)
	# 	# If there are no (valid) credentials available, let the user log in.
	# 	if not creds or not creds.valid:
	# 		if creds and creds.expired and creds.refresh_token:
	# 			creds.refresh(Request())
	# 		else:
	# 			flow = InstalledAppFlow.from_client_secrets_file(wwaglobal.src_path + 'credentials.json', SCOPES)
	# 			creds = flow.run_local_server(port=0)
	# 		# Save the credentials for the next run
	# 		with open(wwaglobal.src_path + 'token.json', 'w') as token:
	# 			token.write(creds.to_json())
	# 
	# 	service = build('sheets', 'v4', credentials=creds)
	# 	#service = Create_Service('token.json', 'sheets', 'v4', SCOPES)
	# 	# Call the Sheets API
	# 
	# 	sheet_id = range_id[range_names]
	# 	# list = [["valuea1"], ["valuea2"], ["valuea3"]]
	# 	# resource = {
	# 	# 	"majorDimension": "ROWS",
	# 	# 	"values": list
	# 	# }
	# 	#
	# 	# range = "to_add!A:A";
	# 	# service.spreadsheets().values().append(
	# 	# 	spreadsheetId=SPREADSHEET_ID,
	# 	# 	range=range,
	# 	# 	body=resource,
	# 	# 	valueInputOption="USER_ENTERED"
	# 	# ).execute()
	# 	sheet = service.spreadsheets()
	# 	result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=range_names,
	# 								valueRenderOption="FORMULA").execute()
	# 	values = result.get('values', [])
	# 	#print((date + timedelta(days=1)).strftime('%Y/%m/%d'))
	# 	start_row = len(values)
	# 	ondate = False
	# 	for row in values:
	# 
	# 		try:
	# 			if xlrd.xldate_as_datetime(row[1],0).date().strftime('%Y/%m/%d') == (date).strftime('%Y/%m/%d'):
	# 				ondate = True
	# 				start_row = int(values.index(row))+1
	# 
	# 		except:
	# 			True
	# 		try:
	# 
	# 
	# 			if str(row[3]) == value[3] and ondate:
	# 				#print('pass 1118')
	# 				start_row = int(values.index(row))+1
	# 				value[3] = ''
	# 				value[4] = ''
	# 				break
	# 			#print('1122 ' + xlrd.xldate_as_datetime(row[1], 0).date().strftime('%Y/%m/%d'))
	# 
	# 			#print(ondate)
	# 			if xlrd.xldate_as_datetime(row[1], 0).date().strftime('%Y/%m/%d') == (date + timedelta(days=1)).strftime('%Y/%m/%d'):
	# 				#start_row = int(values.index(row))
	# 
	# 				break
	# 		except:
	# 			True
	# 
	# 	value[1] = ''
	# 	value_to_input = ' ,'.join(value)
	# 
	# 	batch_update_spreadsheet_request_body = {
	# 		"requests": [
	# 			{
	# 				"insertRange": {
	# 					"range": {
	# 						"sheetId": sheet_id,
	# 						"startRowIndex": start_row,
	# 						"endRowIndex": start_row + 1
	# 					},
	# 					"shiftDimension": "ROWS"
	# 				}
	# 			},
	# 			{
	# 				"pasteData": {
	# 					"data": value_to_input,
	# 					"type": "PASTE_NORMAL",
	# 					"delimiter": ",",
	# 					"coordinate": {
	# 						"sheetId": sheet_id,
	# 						"rowIndex": start_row,
	# 						"columnIndex": 0
	# 					}
	# 				}
	# 			},
	# 			{
	# 				"updateCells": {
	# 					"range": {
	# 						"sheetId": sheet_id,
	# 						"startRowIndex": start_row,
	# 						"endRowIndex": start_row + 1,
	# 						"startColumnIndex": 0,
	# 						"endColumnIndex": 15
	# 					},
	# 					"rows": [
	# 						{
	# 							"values": [
	# 								{
	# 									"userEnteredFormat": {
	# 										"backgroundColor": {
	# 											"red": 1
	# 										}
	# 									}
	# 								}
	# 							]
	# 						}
	# 					],
	# 					"fields": "userEnteredFormat.backgroundColor"
	# 				}
	# 			}
	# 		]
	# 	}
	# 	request = service.spreadsheets().batchUpdate(spreadsheetId=SPREADSHEET_ID,
	# 												 body=batch_update_spreadsheet_request_body)
	# 	response = request.execute()
	# 	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
	# 			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	# 	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	# 	client = gspread.authorize(creds)
	# 	sheetname = str(date.year)+'年'+str(date.month).zfill(2) + '月'
	# 	sheet = client.open('交收list ').worksheet(sheetname)
	# 	cell = 'J' + str(start_row + 1)
	# 	sheet.format(cell, {
	# 		"backgroundColor": {
	# 			"red": 0.09375,
	# 			"green": 0.7109,
	# 			"blue": 0.5468
	# 		}})
	on_date = False
	if delivery:
		range_names = '送貨'
	else:
		range_names = str(date.year) + '年' + str(date.month).zfill(2) + "月"
	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	print(f'{range_names=}')
	sheetname = range_names
	sheet = client.open('交收list ').worksheet(sheetname)
	data = sheet.get_all_records()
	
	if not delivery:
		#print('not delivery')
		for index, row in enumerate(data):
			#print(f"{row['日期']=}")
			if row['日期'] != '':
				if datetime.strptime(row['日期'], '%Y/%m/%d') == date:
					on_date = True
			if on_date:
				if row['客人名稱'].lower() == value[3].lower():
					value[3] = ''
					value[4] = ''
					#pprint(f'{value=} {index+3=}')
					sheet.insert_row(value, index + 3)
					notification_cell = 'J' + str(index + 3)
					sheet.format(notification_cell, {
						"backgroundColor": {
							"red": 0.09375,
							"green": 0.7109,
							"blue": 0.5468
						}})

					red_cell = 'A' + str(index + 3)
					sheet.format(red_cell, {
						"backgroundColor": {
							"red": 1,
							"green": 0,
							"blue": 0
						}})
					break
			if row['日期'] != '':
				if datetime.strptime(row['日期'], '%Y/%m/%d') == date + timedelta(days=1):
					sheet.insert_row(value, index + 2)

					notification_cell = 'J' + str(index + 2)
					sheet.format(notification_cell, {
						"backgroundColor": {
							"red": 0.09375,
							"green": 0.7109,
							"blue": 0.5468
						}})

					red_cell = 'A' + str(index + 2)
					sheet.format(red_cell, {
						"backgroundColor": {
							"red": 1,
							"green": 0,
							"blue": 0
						}})
					break
	else:
		sheet.insert_row(value, len(data) + 2)

		notification_cell = 'K' + str(len(data) + 2)
		sheet.format(notification_cell, {
			"backgroundColor": {
				"red": 0.09375,
				"green": 0.7109,
				"blue": 0.5468
			}})

		red_cell = 'A' + str(len(data) + 2)
		sheet.format(red_cell, {
			"backgroundColor": {
				"red": 1,
				"green": 0,
				"blue": 0
			}})




def meetingNotification(staff):
	#ask for which staff
	# while 1:
	# 	staff = input('員工登入:\n------\n1.\tGeorge\n2.\tHarry\n3.\tRyan\n')
	# 	if int(staff) == 1 or int(staff) == 2 or int(staff) == 3:
	# 		break
	# 	else: print('請輸入正確數字')
	if int(staff) == 1:
		staff_name = 'g'
	elif int(staff) == 2:
		staff_name = 'h'
	elif int(staff) == 3:
		staff_name = 'r'
	staff_message = '/' + staff_name

	# If modifying these scopes, delete the file token.json.
	SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

	# The ID and range of a spreadsheet.
	SPREADSHEET_ID = '1VDfkinOq_SkezxMwmF4Fhi1pe0qTtYEbCxp1CQxtxbY'
	range_names = str(datetime.today().year)+'年'+str(datetime.today().month).zfill(2) + "月"
	creds = None
	# The file token.json stores the user's access and refresh tokens, and is
	# created automatically when the authorization flow completes for the first
	# time.
	if os.path.exists(wwaglobal.src_path + 'token.json'):
		creds = Credentials.from_authorized_user_file(wwaglobal.src_path + 'token.json', SCOPES)
	# If there are no (valid) credentials available, let the user log in.
	if not creds or not creds.valid:
		if creds and creds.expired and creds.refresh_token:
			creds.refresh(Request())
		else:
			flow = InstalledAppFlow.from_client_secrets_file(wwaglobal.src_path + 'credentials.json', SCOPES)
			creds = flow.run_local_server(port=0)
		# Save the credentials for the next run
		with open(wwaglobal.src_path + 'token.json', 'w') as token:
			token.write(creds.to_json())

	service = build('sheets', 'v4', credentials=creds)
	sheet = service.spreadsheets()
	result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=range_names,
								valueRenderOption="FORMULA").execute()
	values = result.get('values', [])



	#whatsapp


	message_store = '/門市交收'
	message_fs = ''
	if datetime.today().weekday() == 6:

		message_fs= '/星期日交收當日通知'

	elif datetime.today().weekday() == 1:
		message_fs= '/星期二交收當日通知'

	print(message_fs)
	#temp = input("wait")
	today_bool = False
	phone_no_fs = []
	phone_no_store = []
	phone_no_delivery = []
	for row in values:
		try:

			if xlrd.xldate_as_datetime(row[1],0).date().strftime('%Y/%m/%d') == (datetime.today()).strftime('%Y/%m/%d'):
				today_bool = True

			if xlrd.xldate_as_datetime(row[1],0).date().strftime('%Y/%m/%d') == (datetime.today() + timedelta(1)).strftime('%Y/%m/%d'):
				today_bool = False


		except:
			#print('error')
			pass

		try:
			if today_bool and len(str(row[4])) ==8:
				#print(values.index(row))

				if str(row[5]) == '魚街':
					phone_no_fs.append(str(row[4]))
				elif str(row[5]) == '門市':
					phone_no_store.append(str(row[4]))
				elif str(row[5]) == '送貨':
					phone_no_delivery.append(str(row[4]))
				scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
						 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
				creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
				client = gspread.authorize(creds)
				sheetname = str(datetime.today().year)+'年'+str(datetime.today().month).zfill(2) + '月'
				sheet = client.open('交收list ').worksheet(sheetname)
				cell = 'M' + str(values.index(row) + 1)
				sheet.format(cell, {
					"backgroundColor": {
						"red": 0.09375,
						"green": 0.7109,
						"blue": 0.5468
					}})
		except:
			pass
		# if today_bool and len(str(row[4])) == 8:
		# 	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
		# 			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
		# 	creds = ServiceAccountCredentials.from_json_keyfile_name('creds.json', scope)
		# 	client = gspread.authorize(creds)
		# 	sheetname = str(datetime.today().month) + '月'
		# 	sheet = client.open('交收list ').worksheet(sheetname)
		# 	cell = 'M' + str(values.index(row)+1)
		# 	sheet.format(cell, {
		# 		"backgroundColor": {
		# 			"red": 0.09375,
		# 			"green": 0.7109,
		# 			"blue": 0.5468
		# 		}})
	print(phone_no_fs)
	print(phone_no_store)

	#temp = input('wait')

	options = Options()
	options.headless = False

	# print("Copying Chrome profile.")
	# wwalog.log("Copying Chrome profile.")
	# os.system(
	# 	"Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")

	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)

	driver.get('https://web.whatsapp.com')

	while 1:
		try:
			WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
			wwalog.log("login whatsapp success.")
			break
		except TimeoutException:
			pass
	textbox_xpath = '//*[@id="main"]/footer/div[1]/div/div/div[2]/div[1]/div/div[2]'

	searchbox_xpath = '//*[@id="side"]/div[1]/div/label/div/div[2]'

	for phone in phone_no_fs:
		driver.find_element_by_xpath(searchbox_xpath).send_keys(phone)
		# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys("63393947")
		driver.find_element_by_xpath(searchbox_xpath).send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox = driver.find_element_by_css_selector("div[role='textbox'][spellcheck='true']")
		inputbox.send_keys(staff_message)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(1)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys(message_fs)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)

	for phone in phone_no_store:
		driver.find_element_by_xpath(searchbox_xpath).send_keys(phone)
		# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys("63393947")
		driver.find_element_by_xpath(searchbox_xpath).send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox = driver.find_element_by_css_selector("div[role='textbox'][spellcheck='true']")
		inputbox.send_keys(staff_message)
		time.sleep(0.5)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys(message_store)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)

	driver.quit()
	print('Done')
	print('注意：有' + str(len(phone_no_delivery)) + '位客人需要送貨\n---------------------------------------')
	for phone in phone_no_delivery:
		print(phone)
	print('\n')

def WhatsappOrder(content,phone,driver):
	print('\n')
	phone_numbers = list(set(tuple(element) for element in phone))
	#phone = set(phone)
	pprint(phone_numbers)
	try:
		wwatkmenu.phone_window(phone_numbers)
	except:
		print('Window not shown')
	#temp=input('請確保以上電話號碼已加入聯絡人 確認後請按Enter以繼續')




	while 1:
		try:
			WebDriverWait(driver, 20).until(
				EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
			wwalog.log("login whatsapp success.")
			break
		except TimeoutException:
			pass
	textbox_xpath = '//*[@id="main"]/footer/div[1]/div/div/div[2]/div[1]/div/div[2]'

	searchbox_xpath = '//*[@id="side"]/div[1]/div/label/div/div[2]'
	for i in range(len(content)):
		driver.find_element_by_xpath(searchbox_xpath).send_keys(content[i]['聯絡電話'])
		# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys("63393947")
		driver.find_element_by_xpath(searchbox_xpath).send_keys(Keys.RETURN)
		time.sleep(0.5)
		#WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, "div[role='textbox'][spellcheck='true']")))
		inputbox = driver.find_element_by_css_selector("div[role='textbox'][spellcheck='true']")
		inputbox.send_keys('多謝師兄支持')
		time.sleep(0.5)
		inputbox.send_keys(':smiling face with smiling')
		time.sleep(1)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys('我地係環球水族，請確認送貨時間同交收地點是否正確')
		time.sleep(0.5)
		inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		for keys,values in content[i].items():
			inputbox.send_keys(keys)
			inputbox.send_keys(':  ')
			inputbox.send_keys(values)
			inputbox.send_keys(Keys.SHIFT + Keys.RETURN)
		time.sleep(0.5)
		#temp= input('wait')
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
	driver.quit()

def PostOrder():
	# If modifying these scopes, delete the file token.json.
	SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

	# The ID and range of a spreadsheet.
	SPREADSHEET_ID = '1VDfkinOq_SkezxMwmF4Fhi1pe0qTtYEbCxp1CQxtxbY'
	range_names = str(datetime.today().year)+'年'+str(datetime.today().month).zfill(2) + "月"
	creds = None
	# The file token.json stores the user's access and refresh tokens, and is
	# created automatically when the authorization flow completes for the first
	# time.
	if os.path.exists(wwaglobal.src_path + 'token.json'):
		creds = Credentials.from_authorized_user_file(wwaglobal.src_path + 'token.json', SCOPES)
	# If there are no (valid) credentials available, let the user log in.
	if not creds or not creds.valid:
		if creds and creds.expired and creds.refresh_token:
			creds.refresh(Request())
		else:
			flow = InstalledAppFlow.from_client_secrets_file(wwaglobal.src_path + 'credentials.json', SCOPES)
			creds = flow.run_local_server(port=0)
		# Save the credentials for the next run
		with open(wwaglobal.src_path + 'token.json', 'w') as token:
			token.write(creds.to_json())

	service = build('sheets', 'v4', credentials=creds)
	sheet = service.spreadsheets()
	result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=range_names,
								valueRenderOption="FORMULA").execute()
	values = result.get('values', [])
	#pprint(values)
	phone =[]
	today_bool = False
	for row in values:
		try:
			if xlrd.xldate_as_datetime(row[1],0).date().strftime('%Y/%m/%d') == (datetime.today() - timedelta(days=2)).strftime('%Y/%m/%d'):
				today_bool = True
			if xlrd.xldate_as_datetime(row[1],0).date().strftime('%Y/%m/%d') == (datetime.today() - timedelta(days=1)).strftime('%Y/%m/%d'):
				today_bool = False


		except:
			wwalog.log("Row: " + str(row))
		try:
			if today_bool and len(str(row[4])) ==8:
				phone.append(str(row[4]))
				scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
						 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
				creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
				client = gspread.authorize(creds)
				sheetname = str(datetime.today().year)+'年'+str(datetime.today().month).zfill(2) + '月'
				sheet = client.open('交收list ').worksheet(sheetname)
				cell = 'S' + str(values.index(row) + 1)
				sheet.format(cell, {
					"backgroundColor": {
						"red": 0.09375,
						"green": 0.7109,
						"blue": 0.5468
					}})
		except:
			pass


	print(phone)



	# print("Copying Chrome profile.")
	# wwalog.log("Copying Chrome profile.")
	# os.system(
	# 	"Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")
	options = Options()
	options.headless = False
	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)

	driver.get('https://web.whatsapp.com')


	while 1:
		try:
			WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
			wwalog.log("login whatsapp success.")
			break
		except TimeoutException:
			pass
	textbox_xpath = '//*[@id="main"]/footer/div[1]/div/div/div[2]/div[1]/div/div[2]'

	searchbox_xpath = '//*[@id="side"]/div[1]/div/label/div/div[2]'
	for phones in phone:
		driver.find_element_by_xpath(searchbox_xpath).send_keys(phones)
		# driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys("63393947")
		driver.find_element_by_xpath(searchbox_xpath).send_keys(Keys.RETURN)
		time.sleep(0.5)

		inputbox = driver.find_element_by_css_selector("div[role='textbox'][spellcheck='true']")
		inputbox.send_keys('/售後慰問1')
		time.sleep(0.5)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
		inputbox.send_keys(Keys.RETURN)
		time.sleep(0.5)
	driver.quit()

def urgeOrder():
	limit_date = datetime.today() - timedelta(days=14)

	if datetime.today().day <=14:
		on_two_month = True
	else:
		on_two_month = False
	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	sheetname = str(datetime.today().year)+'年'+str(datetime.today().month).zfill(2) + '月 auction'
	sheet = client.open('Auction List').worksheet(sheetname)
	data = sheet.get_all_records()
	if on_two_month:
		if str((datetime.today()).month) != '1':
			sheetname = str(datetime.today().year)+'年'+str((datetime.today()).month -1).zfill(2) + '月 auction'
		else:
			sheetname = str(datetime.today().year-1) + '年12月 auction'
		sheet = client.open('Auction List').worksheet(sheetname)
		data2 = sheet.get_all_records()

	item_not_yet_order = []

	for i in range(len(data)-1, 1, -1):
		if not on_two_month and data[i]['Upload Date'] == datetime.strftime(limit_date, '%Y/%m/%d'):
			break
		if data[i]['確認中標者已落單'] != 'y' and len(data[i]['中標者名稱']) != 0 and data[i]['中標者名稱'] != 'n' and len(str(data[i]['Serial Number']))>0:
			if data[i]['確認中標者已落單'] != 'n':
				item_not_yet_order.append([str(data[i]['Serial Number']),data[i]['價錢']])
	if on_two_month:
		for i in range(len(data2)-1, 1, -1):
			if data2[i]['Upload Date'] == datetime.strftime(limit_date, '%Y/%m/%d'):
				break
			if data2[i]['確認中標者已落單'] != 'y' and len(data2[i]['中標者名稱']) != 0 and data2[i]['中標者名稱'] != 'n' and len(
					str(data2[i]['Serial Number'])) > 0:
				item_not_yet_order.append([str(data2[i]['Serial Number']), data2[i]['價錢']])
	pprint(item_not_yet_order)
	# try:
	# 	temp = input('wait')
	# except KeyboardInterrupt:
	# 	return
	#https://band.us/my/bands/posts/
	#pprint(item_not_yet_order)
	#temp=input('wait')
	print('Copying Chrome Profile....')
	os.system(
		"Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")
	#shutil.copytree(wwaglobal.chromeProfilePath, wwaglobal.chromeProfilePath1, dirs_exist_ok=True)
	options = Options()
	options.headless = False
	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)

	for items in item_not_yet_order:
		driver.get('https://band.us/band/78427905/board/search/' + items[0][:8])
		try:
			WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR,".postBody")))
		except TimeoutException:
			print('Item CAA'+items[0]+ ' Unable to Find')
			time.sleep(1)
			continue
		time.sleep(1)
		while 1:
			try:
				WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, ".postBody"))).click()
				break
			except:
				#driver.find_element_by_css_selector('[class="postBody"]').click()
				time.sleep(0.5)



		#element = driver.find_element_by_id("my-id")

		#driver.find_element_by_css_selector("[class='activeLayerView']").send_keys(Keys.END)

		# time.sleep(1)
		# try:
		# 	driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
		# except:
		# 	print('cannot scroll')
		# try:

		#reply=driver.find_element_by_xpath("//*[contains(text(), '恭喜師兄')]//parent::div//parent::div//parent::div//parent::div//parent::div//parent::div//parent::div//child::button[@class='reply _replyBtn']")
		time.sleep(1)
		# try:
		# 	reply = driver.find_element_by_xpath(
		# 		"//*[contains(text(), '恭喜師兄')]//parent::div//parent::div//parent::div//parent::div//parent::div//parent::div//parent::div//child::button[@class='reply _replyBtn']")
		# except:
		# 	continue
		#print(reply)
		#(//*[contains(text(), '160')])[1]//parent::div//child::button[@class='reply _replyBtn']
		counter = 0
		time_error = False
		# time_of_sending = driver.find_element_by_xpath(
		# 	"(//*[contains(text(), '恭喜師兄')]//parent::div)[1]//child::time").text
		while not time_error:
			try:
				time_of_sending = driver.find_element_by_xpath("(//*[contains(text(), '恭喜師兄')]//parent::div)[1]//child::time").text
				break

			except:
				print(f'CAA{items} Time collection error')
				counter+=1
				if counter ==5 :
					time_error=True
				time.sleep(1)
		if time_error:
			continue
		if time_of_sending.find('小時前')!=-1 or time_of_sending.find('分鐘前') != -1:
			continue
		try:
			reply = driver.find_element_by_xpath(
				f"(//*[contains(text(), '{items[1]}') and contains(@class, 'txt _commentContent')])[1]//parent::div//child::button[@class='reply _replyBtn']")
		except:
			continue

		#print(reply)
		reply.send_keys(Keys.RETURN)
		time.sleep(0.5)
		try:
			inputBox = WebDriverWait(driver, 20).until(EC.presence_of_element_located(
				(By.CSS_SELECTOR, 'textarea.commentWrite._use_keyup_event._messageTextArea')))
			JS_ADD_TEXT_TO_INPUT = """
						var elm = arguments[0], txt = arguments[1];
						elm.value += txt;
						elm.dispatchEvent(new Event('change'));
						"""

			elem = driver.find_element_by_css_selector('textarea.commentWrite._use_keyup_event._messageTextArea')
			text = '師兄你好，根據記錄，你中標的珊瑚仍未下單，請盡快落單，並於兩星期內進行交收'
			driver.execute_script(JS_ADD_TEXT_TO_INPUT, elem, text)
		except:
			print('got some problem in sending message')
			break
		time.sleep(0.5)
		try:
			submit_btn = driver.find_element_by_xpath("//*[@type='submit' and @class='writeSubmit uButton _sendMessageButton -active']")
			submit_btn.send_keys(Keys.RETURN)
		except:
			print('Submittion Error')
			break

		# replyField = driver.find_element_by_xpath("//*[contains(text(), '恭喜師兄')]//parent::div//parent::div//parent::div//parent::div//parent::div//parent::div//parent::div//child::div//*[@class='cCommentWrite _replyInputRegion replyCommentWrite'][1]")
		# #cCommentWrite _replyInputRegion replyCommentWrite
		# time.sleep(0.5)
		# replyField.send_keys('hi')

		time.sleep(0.5)
		# except:
		# 	print('cant')
		# 	pass
	driver.quit()

def auction_placement(date):
	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	sheetname = f'{str(date.year)}年{str(date.month).zfill(2)}月 auction'
	sheet = client.open('Auction List').worksheet(sheetname)
	data = sheet.get_all_records()
	on_date = False
	data.reverse()
	wait_to_append = []
	areaB_list = []
	for items in data:
		try:
			if items['Upload Date'] == datetime.strftime(date, '%Y/%m/%d'):
				on_date = True
		except:
			continue

		if on_date:
			if items['Upload Date'] == '':
				break
			if items['中標者名稱'] != '' and items['中標者名稱'] != 'n' and items['Serial Number'] != '':
				if '豬腰' in items['Name']:
					areaB_list.append([str(items['Serial Number']), str(items['Name'])])
					continue
				elif '火柴' in items['Name']:
					areaB_list.append([str(items['Serial Number']), str(items['Name'])])
					continue
				elif '啫喱' in items['Name']:
					areaB_list.append([str(items['Serial Number']), str(items['Name'])])
					continue
				elif '提子' in items['Name']:
					areaB_list.append([str(items['Serial Number']), str(items['Name'])])
					continue
				print(f"{str(items['Serial Number']).replace('CAA', '')=}")
				print(f"{str(items['Serial Number']).replace('CAA', '').split(' ')}")
				wait_to_append.append([str(items['Serial Number']).replace('CAA', '').split(' '), str(items['Name']).split(' / ')])
	pprint(f'{wait_to_append=}')
	splited_list = []
	for item in wait_to_append:
		for index, combination in enumerate(item[0]):
			#print(f'{str(combination).strip()}')
			splited_list.append([str(combination).strip(), str(item[1][index]).strip()])
	splited_list.reverse()
	pprint(f'中標珊瑚 (非豬腰/火柴): {splited_list}')
	total_auction = list(map(lambda name:name,splited_list))
	#wait = input('wait')
	sheetname = 'list'
	sheet = client.open('格仔板').worksheet(sheetname)
	data = {}
	data = sheet.get_all_records()
	#pprint(f'{data=}')
	spaces = []
	placed = []
	for index, board in enumerate(data):
		#pprint(f'{index=}')
		#pprint(f'{board=}')
		for keys, coordinate in board.items():
			if coordinate == '' and keys != '板':
				complete_coordinate = f'{str(board["板"]).zfill(3)}{str(keys).zfill(2)}'
				spaces.append([complete_coordinate])
			elif coordinate != '' and keys != '板':
				complete_coordinate = f'{str(board["板"]).zfill(3)}{str(keys).zfill(2)}'
				placed.append([str(coordinate).zfill(4), complete_coordinate])
	#pprint(spaces)
	#pprint(placed)

	del_no = 0
	for index, space in enumerate(spaces):
		try:
			spaces[index].append(splited_list[index-del_no][0])
			spaces[index].append(splited_list[index - del_no][1])
			splited_list.remove(splited_list[index-del_no])
			del_no += 1
		except:
			break

	for space in spaces:
		#print(f'{space=}')
		if len(space)>1:
			print(space)

			for index, board in enumerate(data):
				if str(board['板']).zfill(3) == space[0][:3]:
					row = index+2
					column = int(space[0][3:5]) + 1
					#pprint(f'{row=}{column=}')
					break
			sheet.update_cell(row,column,str(space[1]))
			wwalog.log(f'Updated cell{row=} {column=}: {space[1]}')
	#pprint(spaces)
	if len(splited_list)>0:
		pprint(f'not yet input: {splited_list}')
	else:
		print('All are appended')

	pprint(f'這些為豬腰或火柴，請放於B區: {areaB_list}')
	pprint(f'{total_auction=} {spaces=} {areaB_list}')

	wwatkmenu.placement_window(total_auction, spaces, areaB_list, splited_list)
	# if '豬腰' in item['Name'] or '火柴' in item['Name']:
	# 	areaB_list.append(item['備註'])
def check_auction_item():
	print('Loading...')
	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	sheetname = str(datetime.today().year)+'年'+str(datetime.today().month).zfill(2) + '月'
	sheet = client.open('交收list ').worksheet(sheetname)
	data = sheet.get_all_records()
	today_item_list = []
	temp_list = []
	areaB_list = []
	today = False
	#pprint(data)
	for item in data:
		try:

			if item['日期'] == datetime.strftime(datetime.today(), '%Y/%m/%d'):
				#print('on date')
				today = True
		except:
			if not today:
				continue
			pass
		if today:

			temp_list = item['備註'][1:len(item['備註'])].split(' ')
			for temp_item in temp_list:
				if len(temp_item)==4:
					try:
						temp = int(temp_item)
					except:
						continue
					today_item_list.append(str(temp_item))
			try:
				if item['日期'] == datetime.strftime(datetime.today()+timedelta(days=1), '%Y/%m/%d'):
					break
			except:
				pass
			temp_list.clear()
	sheetname = 'list'
	sheet = client.open('格仔板').worksheet(sheetname)
	data = list(sheet.get_all_records())

	coordinate_list = []
	not_found_list = []
	for index, item in enumerate(today_item_list):
		found = False
		for index2, list_item in enumerate(data):
			for keys,code in list_item.items():
				if keys == '板':
					continue
				if str(code) == item:
					coordinate_list.append([index, item, f'{str(list_item["板"]).zfill(3)}{str(keys).zfill(2)}'])
					sheet.update_cell(index2+2, int(keys)+1, '')
					#print(f'row={index2+2} column={int(keys)+1}')
					found = True
		if not found:
			not_found_list.append(item)
	# for items in coordinate_list:
	# 	if items[1] != '':
	# 		print(items)
	pprint(coordinate_list)
	pprint(f'這些珊瑚仍未編位 {not_found_list}')
	wwatkmenu.check_placement_window(coordinate_list,not_found_list)

	#pprint(today_item_list)
def finished_meeting():

	date = datetime.today()
	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	sheetname = str(date.year) + '年' + str(date.month).zfill(2) + '月'
	print(f'{sheetname}')
	sheet = client.open('交收list ').worksheet(sheetname)
	# sheet = client.open('Auction List').worksheet('11月 auction')
	data = sheet.get_all_records()
	today_orderid = []
	phone = set()
	date_found = False
	for i in range(len(data) - 1, 1, -1):
		#print(f'{data[i]["日期"]}')
		if data[i]['日期'] == datetime.strftime(date+timedelta(days=1), '%Y/%m/%d'):
			#print('find')
			date_found = True
		if date_found:
			try:
				if len(str(data[i]['聯絡電話'])) == 8:
					phone.add(data[i]['聯絡電話'])
			except:
				pass
			try:
				if len(str(data[i]['訂單編號'])) == 16:
					#print(data[i]['訂單編號'])
					today_orderid.append(int(str(data[i]['訂單編號'])))
			except:
				print(f'Invalid order id in row {i+2}')
				continue
			cell_finished_meeting = 'N' + str(i + 2)
			sheet.format(cell_finished_meeting, {
				"backgroundColor": {
					"red": 0.09375,
					"green": 0.7109,
					"blue": 0.5468
				}})
			cell_thankyou_msg = 'O' + str(i + 2)
			sheet.format(cell_thankyou_msg, {
				"backgroundColor": {
					"red": 0.09375,
					"green": 0.7109,
					"blue": 0.5468
				}})
			try:
				if data[i]['日期'] == datetime.strftime(date, '%Y/%m/%d'):
					break
			except:
				pass
	pprint(today_orderid)
	#temp = input('wait')
	# os.system(
	# 	"Xcopy \"" + wwaglobal.chromeProfilePath + "\" \"" + wwaglobal.chromeProfilePath1 + "\" /E/H/C/I/Y/Q > nul 2>&1")
	options = Options()
	options.headless = False
	options = webdriver.ChromeOptions()
	options.add_experimental_option("excludeSwitches", ['enable-automation', 'enable-logging'])
	options.add_argument("user-data-dir=" + wwaglobal.chromeProfilePath1)

	driver = webdriver.Chrome(executable_path=wwaglobal.chromeDriverPath, options=options)
	driver.get('https://www.boutir.com/user_cms/orders')
	while 1:
		try:
			WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//*[@class='order-table']")))
			break
		except TimeoutException:
			print('Loading...')
			time.sleep(1)
	unfound_id = []
	for id in today_orderid:
		statue=''
		try:
			check_box = driver.find_element_by_xpath(f"((//*[contains(text(), '{id}')])//parent::td//parent::tr//child::input[@type='checkbox'])")
			status = driver.find_element_by_xpath(f"((//*[contains(text(), '{id}')])//parent::td//parent::tr//child::td[@class='status']//child::span)").text
			time.sleep(0.5)
			if status != '待客取貨' or status != '等待發貨':
				unfound_id.append([id,'invalid status'])
				print(f'Order {id} 已發貨或未確認付款')
				time.sleep(2)
				continue
		except:
			print(f'Check Box of order {id} can\'t find')
			unfound_id.append([id,'not found'])
			time.sleep(2)
			continue
		driver.execute_script("arguments[0].scrollIntoView();", check_box)
		time.sleep(1)
		check_box.click()
		#check_box.send_keys(Keys.RETURN)
		time.sleep(2)
	time.sleep(1)
	driver.find_element_by_xpath("//button[@id='actionMenu']").click()
	time.sleep(0.5)
	driver.find_element_by_xpath("//li[contains(text(), '確認送貨')]").click()
	time.sleep(1)
	while 1:
		try:
			WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//button[contains(text(), '是')]")))
			break
		except TimeoutException:
			time.sleep(1)
			print('Loading for confirm button...')
	try:
		driver.find_element_by_xpath("//button[contains(text(), '是')]").click()
	except:
		print('Error occured. Please check the order status.')
		return
	time.sleep(5)
	driver.get('https://web.whatsapp.com')
	while 1:
		try:
			WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div')))
			wwalog.log("login whatsapp success.")
			break
		except TimeoutException:
			time.sleep(0.5)
	textbox_xpath = "//div[@title='輸入訊息']"

	searchbox_xpath = "(//div[@role='textbox'])[1]"
	for phone_no in phone:
		driver.find_element_by_xpath(searchbox_xpath).send_keys(phone_no)
		time.sleep(0.5)
		driver.find_element_by_xpath(searchbox_xpath).send_keys(Keys.RETURN)
		time.sleep(0.5)
		driver.find_element_by_xpath(textbox_xpath).send_keys('/交收完畢1')
		time.sleep(0.5)
		driver.find_element_by_xpath(textbox_xpath).send_keys(Keys.RETURN)
		time.sleep(0.5)
		driver.find_element_by_xpath(textbox_xpath).send_keys(Keys.RETURN)
		time.sleep(0.5)

	driver.quit()
	append_excel(date)
	wwatkmenu.order_id_window(unfound_id)
	#狀態異常或查找不到訂單 unfound_id


def appendGoogleAuction(date,temp_namelist, temp_price):
	#date =datetime.strptime(input("Date (yyyy-mm-dd): "), '%Y-%m-%d')

	pprint(f'{date=}')
	pprint(f'{temp_namelist=}')
	pprint(f'{temp_price=}')
	temp_namelist.reverse()
	temp_price.reverse()
	#date = datetime.datetime.strptime(date, '%Y-%m-%d')
	scope = ["https://spreadsheets.google.com/feeds",'https://www.googleapis.com/auth/spreadsheets',"https://www.googleapis.com/auth/drive.file","https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json',scope)
	client = gspread.authorize(creds)
	sheetname = str(date.year)+'年'+str(date.month).zfill(2) + '月 auction'
	print(f'{sheetname}')
	sheet = client.open('Auction List').worksheet(sheetname)
	#sheet = client.open('Auction List').worksheet('11月 auction')
	data = sheet.get_all_records()
	#pprint(data[251])
	pprint(f'{temp_namelist=}')
	for i in range(len(data)-1,1,-1):


		if data[i]['Upload Date'] == datetime.strftime(date, '%Y/%m/%d'):
			#print(len(data)-1-i)
			for index, items in enumerate(temp_namelist):
				try:

					if items == 'No one':
						sheet.update_cell(i + 2 - index, 16, 'n')


					else:

						sheet.update_cell(i + 2 - index, 16, items)
						sheet.update_cell(i + 2 - index, 17, temp_price[index])

				except:
					print('1st Could not append to Auction List')
			break

			# print('on date')
			# pprint(f'{date[len(data)-1-i]=}')
			# pprint(f'{temp_namelist[len(data)-1-i]=}')
			# pprint(f'{temp_price[len(data)-1-i]=}')
			# try:
			#
			# 	if temp_namelist[len(data)-1-i] == 'No one':
			# 		sheet.update_cell(i + 2, 16, 'n')
			#
			#
			# 	else:
			#
			# 		sheet.update_cell(i+2,16,temp_namelist[len(data)-1-i])
			# 		sheet.update_cell(i+2,17,temp_price[len(data)-1-i])
			#
			# except:
			# 	print('2nd Could not append to Auction List')

	temp_namelist.reverse()
	temp_price.reverse()


def append_excel(date):


	scope = ["https://spreadsheets.google.com/feeds", 'https://www.googleapis.com/auth/spreadsheets',
			 "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
	creds = ServiceAccountCredentials.from_json_keyfile_name(wwaglobal.src_path + 'creds.json', scope)
	client = gspread.authorize(creds)
	sheetname = str(date.year) + '年' + str(date.month).zfill(2) + '月'
	#print(f'{sheetname}')
	sheet = client.open('交收list ').worksheet(sheetname)
	# sheet = client.open('Auction List').worksheet('11月 auction')
	data = sheet.get_all_records()
	today_cash_order = []
	conditions = ('payme', 'bank')
	date_found = False
	for i in range(len(data) - 1, 1, -1):
		if data[i]['日期'] == datetime.strftime(date + timedelta(days=1), '%Y/%m/%d'):
			date_found = True
		if date_found:
			try:
				if len(str(data[i]['訂單編號'])) == 16 and data[i]['付款方式'] not in conditions:
					# print(data[i]['訂單編號'])
					today_cash_order.append([int(str(data[i]['訂單編號'])), int(data[i]['金額'])])
			except:
				print(f'Invalid order id in row {i + 2}')
				continue
			try:
				if data[i]['日期'] == datetime.strftime(date, '%Y/%m/%d'):
					break
			except:
				pass
	pprint(today_cash_order)
	insert_row = 0
	if len(today_cash_order)>10:
		insert_row = len(today_cash_order)-10
	sheet_name = f'{datetime.today().year}{datetime.strftime(datetime.today(), "%b")}'
	file_name = wwaglobal.dropboxPath + '日常記錄 Daily Record.xlsx'
	wb = load_workbook(filename = file_name)
	sheet = wb[sheet_name]
	for i in range(insert_row):
		sheet.insert_rows(7)
		sheet['D7'].border = Border(left=Side(border_style='medium', color='000000'))
		sheet['E7'].border = Border(right=Side(border_style='medium', color='000000'))
		sheet['G7'].border = Border(left=Side(border_style='medium', color='000000'))
		sheet['H7'].border = Border(right=Side(border_style='medium', color='000000'))
	for row, info in enumerate(today_cash_order):
		sheet.cell(column=4, row=row+2, value="'"+str(info[0]))
		sheet.cell(column=5, row=row + 2, value=info[1])
	# sheet['B1'].value = '12/12/2021'
	wb.save(filename= file_name)

def move_objects_cms(number_of_items):
	def connect():
		client = AdbClient(host="127.0.0.1", port=5037)  # Default is "127.0.0.1" and 5037

		devices = client.devices()

		if len(devices) == 0:
			#print('No devices')
			#wwatkmenu.question_box('Device Not Found', 'Device Not Found. Try Again?')
			#wwatkmenu.error_box('No Devices Connected')
			return(None, None)

		device = devices[0]

		return device, client


	seconds = 60 # 31secs
	scrolling_time = 10  # scroll for how many times

	#number_of_items = int(input('Move how many objects? '))
	while 1:
		device, client = connect()
		#print(device)
		if device is None:
			if wwatkmenu.question_box('Device Not Found','Device Not Found. Try Again?') == 'no':
				return
		else: break
	# device.shell('input touchscreen swipe 1000 2025 1000 2000 500')
	# time.sleep(2)
	for items in range(number_of_items):
		# move the second object to the bottom of the screen
		#device.shell('input touchscreen swipe 1000 700 1000 2025 500')
		#time.sleep(1)
		print('moving')
		#device.shell("sendevent /dev/input/event~ 1 330 0")

		device.shell("sendevent /dev/input/event1 1 330 1")
		device.shell("sendevent /dev/input/event1 3 57 10")
		device.shell("sendevent /dev/input/event1 3 53 1000")
		device.shell("sendevent /dev/input/event1 3 54 700")
		device.shell("sendevent /dev/input/event1 0 0 0")
		device.shell("sendevent /dev/input/event1 3 53 1000")
		device.shell("sendevent /dev/input/event1 3 54 2025")
		device.shell("sendevent /dev/input/event1 0 0 0")

		time.sleep(seconds)
		print('lift up')
		device.shell("sendevent /dev/input/event1 1 330 0")
		device.shell("sendevent /dev/input/event1 0 0 0")
		# holding the object at the bottom
		# device.shell(f'input touchscreen swipe 1000 1950 1000 2000 100')
		# time.sleep(0.5)
		#device.shell(f'input touchscreen swipe 1000 2025 1000 2025 {seconds}')
		time.sleep(0.5)

		# scroll to the top
		for scroll in range(scrolling_time):
			device.shell('input touchscreen swipe 540 700 540 2000 50')
			#device.shell("sendevent /dev/input/event1 0 0 0")
			time.sleep(0.5)
		print('scrolled')
		# device.shell("sendevent /dev/input/event~ 1 330 1")  # Puts down finger
		# device.shell("sendevent /dev/input/event~ 3 57 10")  # Sets pressure
		# device.shell("sendevent /dev/input/event~ 3 53 100")  # Sets X to 100
		# device.shell("sendevent /dev/input/event~ 3 54 230")  # Sets Y to 230
		# device.shell("sendevent /dev/input/event~ 0 0 0")  # "0 0 0" (its called a SYN_REPORT)
		# device.shell("sendevent /dev/input/event~ 1 330 0")  # Lift up finger
		# device.shell("sendevent /dev/input/event~ 0 0 0")
		time.sleep(1)
